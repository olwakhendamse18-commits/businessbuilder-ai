from flask import Flask, request, jsonify, render_template, redirect, session, send_file
from openai import OpenAI
from dotenv import load_dotenv
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from werkzeug.exceptions import HTTPException
from pypdf import PdfReader
from docx import Document
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from cryptography.fernet import Fernet, InvalidToken

import requests
import sqlite3
import os
import base64
import hashlib
import hmac
import logging
import secrets
import psycopg2
import re
import json
import urllib.parse
from io import BytesIO
from xml.sax.saxutils import escape


load_dotenv()

app = Flask(__name__)
logger = logging.getLogger(__name__)
secret_key = os.getenv("SECRET_KEY")

if not secret_key:
    if os.getenv("DATABASE_URL"):
        raise RuntimeError("SECRET_KEY must be configured in production.")

    secret_key = "businessbuilder-local-development-secret"

app.secret_key = secret_key

UPLOAD_FOLDER = "uploads"
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

openai_api_key = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=openai_api_key) if openai_api_key else None

SYSTEM_PROMPT = """
You are BusinessBuilder AI, an AI assistant that helps users create businesses.
Help with business ideas, branding, Shopify planning, Canva ideas, marketing, pricing, and invoices.
Ask useful questions and give step-by-step help.
Give practical guidance for pricing, payment setup, supplier/platform choices, Shopify, Canva, and launch planning.
Do not scrape websites or claim live access to Amazon, Alibaba, AliExpress, Takealot, supplier stock, current prices, delivery times, or availability.
Do not collect card numbers, bank passwords, ID numbers, private banking credentials, API keys, or payment account secrets.
Do not automatically create PayPal, Paystack, Shopify payment, supplier, domain, email-marketing, or subscription accounts.
Do not send mass marketing emails without an official integration and explicit user approval.
Never connect a third-party platform without explicit user permission, and never ask for a platform password.
Create drafts first and require review, approval, and separate final confirmation before supported external actions.
Tell users to verify provider requirements, fees, policies, taxes, customs, supplier quality, shipping, returns, and legal rules directly with the provider.
"""

SHOPIFY_ADMIN_API_VERSION = "2026-04"
CANVA_AUTHORIZATION_URL = "https://www.canva.com/api/oauth/authorize"
CANVA_TOKEN_URL = "https://api.canva.com/rest/v1/oauth/token"
CANVA_PROFILE_URL = "https://api.canva.com/rest/v1/users/me/profile"
CANVA_DESIGNS_URL = "https://api.canva.com/rest/v1/designs"
CANVA_SCOPES = os.getenv(
    "CANVA_SCOPES",
    "design:meta:read design:content:write profile:read"
)

APP_CATALOG = {
    "brevo": {"name": "Brevo", "category": "Email marketing", "best_for": "Beginner-friendly email and SMS campaign planning", "budget_fit": "Very low to medium", "connection_mode": "manual", "draft_actions": ["email_campaign"]},
    "mailchimp": {"name": "Mailchimp", "category": "Email marketing", "best_for": "Beginner-friendly newsletters and email campaigns", "budget_fit": "Low to medium", "connection_mode": "manual", "draft_actions": ["email_campaign"]},
    "klaviyo": {"name": "Klaviyo", "category": "Email marketing", "best_for": "Ecommerce segmentation and automation planning", "budget_fit": "Medium to high", "connection_mode": "manual", "draft_actions": ["email_campaign"]},
    "shopify_email": {"name": "Shopify Email", "category": "Email marketing", "best_for": "Email marketing for Shopify-first businesses", "budget_fit": "Very low to medium", "connection_mode": "shopify", "draft_actions": ["email_campaign"]},
    "canva": {"name": "Canva", "category": "Design", "best_for": "Branding, social graphics, banners, and design briefs", "budget_fit": "Very low to high", "connection_mode": "canva", "draft_actions": ["canva_design_brief"]},
    "shopify": {"name": "Shopify", "category": "Store/ecommerce", "best_for": "Products, collections, pages, and ecommerce store assets", "budget_fit": "Low to high", "connection_mode": "shopify", "draft_actions": ["shopify_product_draft", "website_copy"]},
    "webflow": {"name": "Webflow", "category": "Website builder", "best_for": "Professional visual websites and CMS page planning", "budget_fit": "Medium to high", "connection_mode": "guidance", "draft_actions": ["webflow_page_draft", "website_copy"]},
    "wordpress": {"name": "WordPress / WooCommerce", "category": "Website builder", "best_for": "Flexible websites and ecommerce guidance", "budget_fit": "Low to high", "connection_mode": "guidance", "draft_actions": ["website_copy"]},
    "wix": {"name": "Wix", "category": "Website builder", "best_for": "Beginner-friendly website planning", "budget_fit": "Low to medium", "connection_mode": "guidance", "draft_actions": ["website_copy"]},
    "meta_ads": {"name": "Meta Ads", "category": "Ads", "best_for": "Facebook and Instagram ad-copy and campaign-plan drafts", "budget_fit": "Medium to high", "connection_mode": "guidance", "draft_actions": ["social_ad_draft"]},
    "google_ads": {"name": "Google Ads", "category": "Ads", "best_for": "Search-ad keyword and campaign-plan drafts", "budget_fit": "Medium to high", "connection_mode": "guidance", "draft_actions": ["social_ad_draft"]},
    "tiktok_ads": {"name": "TikTok Ads", "category": "Ads", "best_for": "Short-form creative and campaign-plan drafts", "budget_fit": "Medium to high", "connection_mode": "guidance", "draft_actions": ["social_ad_draft"]},
    "zapier": {"name": "Zapier", "category": "Automation", "best_for": "No-code workflow automation plans", "budget_fit": "Low to high", "connection_mode": "guidance", "draft_actions": ["automation_plan"]},
    "make": {"name": "Make", "category": "Automation", "best_for": "Visual multi-step automation plans", "budget_fit": "Low to high", "connection_mode": "guidance", "draft_actions": ["automation_plan"]}
}

APP_ACTION_TYPES = {
    "email_campaign": "Email Campaign Draft",
    "social_ad_draft": "Ad Copy and Campaign Plan",
    "canva_design_brief": "Canva Design Brief",
    "shopify_product_draft": "Shopify Product Draft",
    "webflow_page_draft": "Webflow Page/CMS Draft",
    "website_copy": "Website Copy Draft",
    "automation_plan": "Automation Plan"
}


# -----------------------------
# VALIDATION / ERROR HELPERS
# -----------------------------

class ExternalServiceError(Exception):
    pass


def get_encryption_key():
    encryption_key = os.getenv("ENCRYPTION_KEY", "").strip()

    if encryption_key:
        try:
            Fernet(encryption_key.encode("utf-8"))
        except (TypeError, ValueError):
            raise ExternalServiceError(
                "Token encryption is misconfigured. Please contact support."
            )

        return encryption_key.encode("utf-8")

    if os.getenv("DATABASE_URL"):
        return None

    # Local development still encrypts tokens at rest without requiring .env
    # changes. Production must use a dedicated ENCRYPTION_KEY environment value.
    local_digest = hashlib.sha256(
        f"local-token-encryption:{app.secret_key}".encode("utf-8")
    ).digest()

    return base64.urlsafe_b64encode(local_digest)


def encrypt_token(value):
    if not value:
        return value

    encryption_key = get_encryption_key()

    if not encryption_key:
        raise ExternalServiceError(
            "Token encryption is not configured. Add ENCRYPTION_KEY before connecting an account."
        )

    # Tokens are encrypted before they leave application memory for storage.
    return Fernet(encryption_key).encrypt(
        value.encode("utf-8")
    ).decode("utf-8")


def decrypt_token(value):
    if not value:
        return value

    encryption_key = get_encryption_key()

    if not encryption_key:
        # Legacy production rows may still contain plaintext until users
        # reconnect. New production writes are blocked by encrypt_token().
        return value

    try:
        return Fernet(encryption_key).decrypt(
            value.encode("utf-8")
        ).decode("utf-8")
    except (InvalidToken, ValueError):
        # Backwards compatibility for connected users whose tokens predate
        # encryption. Reconnecting stores an encrypted replacement.
        return value


def render_error(message, status_code=500, back_url="/dashboard"):
    return render_template(
        "error.html",
        message=message,
        back_url=back_url
    ), status_code


def is_valid_email(email):
    return bool(
        email
        and len(email) <= 254
        and re.fullmatch(
            r"[^@\s]+@[^@\s]+\.[^@\s]+",
            email
        )
    )


def safe_openai_chat_completion(**kwargs):
    if not client:
        raise ExternalServiceError(
            "AI generation is temporarily unavailable. Please try again later."
        )

    try:
        return client.chat.completions.create(**kwargs)
    except Exception:
        raise ExternalServiceError(
            "AI generation is temporarily unavailable. Please try again later."
        )


def send_email(to, subject, body):
    provider = os.getenv("EMAIL_PROVIDER", "").strip().lower()
    api_key = os.getenv("EMAIL_API_KEY", "").strip()
    from_email = os.getenv("FROM_EMAIL", "").strip()

    if not provider or not api_key or not from_email:
        logger.info("Email notification skipped because email is not configured.")
        return False

    if provider != "resend":
        logger.warning("Email notification skipped because the provider is unsupported.")
        return False

    try:
        response = requests.post(
            "https://api.resend.com/emails",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            },
            json={
                "from": from_email,
                "to": [to],
                "subject": subject,
                "text": body
            },
            timeout=10
        )
        response.raise_for_status()
    except requests.RequestException:
        logger.warning("Email notification could not be delivered.")
        return False

    return True


def send_email_notification(to_email, subject, body):
    if not to_email:
        return False

    return send_email(to_email, subject, body)


@app.errorhandler(ExternalServiceError)
def handle_external_service_error(error):
    if request.path == "/chat":
        return jsonify({"reply": str(error)}), 503

    return render_error(str(error), 503)


@app.errorhandler(Exception)
def handle_unexpected_error(error):
    if isinstance(error, HTTPException):
        return error

    if request.path == "/chat":
        return jsonify({
            "reply": "Something went wrong. Please try again."
        }), 500

    return render_error(
        "Something went wrong while processing your request. Please try again.",
        500
    )


# -----------------------------
# DATABASE HELPERS
# -----------------------------

def using_postgres():
    return bool(os.getenv("DATABASE_URL"))


def db():
    database_url = os.getenv("DATABASE_URL")

    if database_url:
        return psycopg2.connect(database_url)

    return sqlite3.connect("business_ai.db")


def sql(query):
    """
    SQLite uses ? placeholders.
    PostgreSQL uses %s placeholders.
    This lets the same app work locally and on Render PostgreSQL.
    """
    if using_postgres():
        return query.replace("?", "%s")

    return query


def init_db():
    conn = db()
    cur = conn.cursor()

    id_type = "SERIAL PRIMARY KEY" if using_postgres() else "INTEGER PRIMARY KEY AUTOINCREMENT"

    def execute_schema(query):
        cur.execute(sql(query))

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS users (
            id {id_type},
            email TEXT UNIQUE NOT NULL,
            password TEXT NOT NULL
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS chats (
            id {id_type},
            user_id INTEGER NOT NULL,
            title TEXT NOT NULL DEFAULT 'New Chat'
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS messages (
            id {id_type},
            user_id INTEGER NOT NULL,
            chat_id INTEGER,
            role TEXT NOT NULL,
            content TEXT NOT NULL
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS projects (
            id {id_type},
            user_id INTEGER NOT NULL,
            title TEXT NOT NULL,
            business_type TEXT NOT NULL,
            description TEXT NOT NULL
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS uploads (
            id {id_type},
            user_id INTEGER NOT NULL,
            filename TEXT NOT NULL,
            filepath TEXT NOT NULL
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS payments (
            id {id_type},
            user_id INTEGER NOT NULL,
            provider TEXT NOT NULL,
            amount INTEGER NOT NULL,
            status TEXT NOT NULL,
            reference TEXT
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS workflow_progress (
            id {id_type},
            user_id INTEGER NOT NULL,
            step_number INTEGER NOT NULL,
            step_name TEXT NOT NULL,
            status TEXT NOT NULL DEFAULT 'completed'
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS workflow_answers (
            id {id_type},
            user_id INTEGER NOT NULL,
            step_number INTEGER NOT NULL,
            step_name TEXT NOT NULL,
            answer TEXT NOT NULL
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS business_plans (
            id {id_type},
            user_id INTEGER NOT NULL,
            title TEXT NOT NULL,
            content TEXT NOT NULL
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS ai_store_builds (
            id {id_type},
            user_id INTEGER NOT NULL,
            title TEXT NOT NULL,
            store_name TEXT NOT NULL,
            content TEXT NOT NULL,
            status TEXT NOT NULL,
            created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS store_agent_tasks (
            id {id_type},
            user_id INTEGER NOT NULL,
            task_type TEXT NOT NULL,
            title TEXT NOT NULL,
            draft_content TEXT NOT NULL,
            status TEXT NOT NULL,
            approved INTEGER NOT NULL DEFAULT 0,
            applied INTEGER NOT NULL DEFAULT 0,
            created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS product_research (
            id {id_type},
            user_id INTEGER NOT NULL,
            title TEXT NOT NULL,
            business_type TEXT NOT NULL,
            target_market TEXT NOT NULL,
            sourcing_preference TEXT NOT NULL,
            budget TEXT NOT NULL,
            country TEXT NOT NULL,
            content TEXT NOT NULL,
            created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS user_packages (
            id {id_type},
            user_id INTEGER NOT NULL,
            package_name TEXT NOT NULL,
            status TEXT NOT NULL,
            created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS user_settings (
            id {id_type},
            user_id INTEGER NOT NULL,
            theme TEXT NOT NULL DEFAULT 'default',
            created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS user_onboarding (
            id {id_type},
            user_id INTEGER NOT NULL,
            business_idea TEXT,
            country TEXT,
            budget TEXT,
            has_shopify TEXT,
            has_canva TEXT,
            product_type TEXT,
            target_customer TEXT,
            business_goal TEXT,
            completed INTEGER NOT NULL DEFAULT 0,
            created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS support_tickets (
            id {id_type},
            user_id INTEGER,
            email TEXT NOT NULL,
            category TEXT NOT NULL,
            subject TEXT NOT NULL,
            message TEXT NOT NULL,
            status TEXT NOT NULL DEFAULT 'open',
            created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS pricing_advice (
            id {id_type},
            user_id INTEGER NOT NULL,
            business_idea TEXT,
            product_type TEXT,
            target_customer TEXT,
            cost_price TEXT,
            desired_profit TEXT,
            competitor_price TEXT,
            country TEXT,
            budget TEXT,
            content TEXT NOT NULL,
            created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS payment_guides (
            id {id_type},
            user_id INTEGER NOT NULL,
            country TEXT,
            business_type TEXT,
            selling_platform TEXT,
            payment_options TEXT,
            budget TEXT,
            content TEXT NOT NULL,
            created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS supplier_recommendations (
            id {id_type},
            user_id INTEGER NOT NULL,
            business_idea TEXT,
            product_type TEXT,
            country TEXT,
            budget TEXT,
            sourcing_preference TEXT,
            risk_level TEXT,
            content TEXT NOT NULL,
            created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS email_campaigns (
            id {id_type},
            user_id INTEGER NOT NULL,
            business_name TEXT,
            business_type TEXT,
            target_customer TEXT,
            campaign_goal TEXT,
            email_platform TEXT,
            offer TEXT,
            tone TEXT,
            content TEXT NOT NULL,
            created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS domain_guides (
            id {id_type},
            user_id INTEGER NOT NULL,
            business_name TEXT,
            preferred_domain TEXT,
            budget TEXT,
            provider_preference TEXT,
            website_platform TEXT,
            country TEXT,
            content TEXT NOT NULL,
            created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS app_recommendations (
            id {id_type},
            user_id INTEGER NOT NULL,
            business_type TEXT,
            country TEXT,
            budget TEXT,
            goal TEXT,
            selling_platform TEXT,
            content TEXT NOT NULL,
            created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS app_connections (
            id {id_type},
            user_id INTEGER NOT NULL,
            platform TEXT NOT NULL,
            category TEXT NOT NULL,
            connection_status TEXT NOT NULL,
            access_token_encrypted TEXT,
            refresh_token_encrypted TEXT,
            api_key_encrypted TEXT,
            account_name TEXT,
            scopes TEXT,
            created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS app_action_drafts (
            id {id_type},
            user_id INTEGER NOT NULL,
            platform TEXT NOT NULL,
            action_type TEXT NOT NULL,
            title TEXT NOT NULL,
            draft_content TEXT NOT NULL,
            status TEXT NOT NULL DEFAULT 'Needs Approval',
            approved INTEGER NOT NULL DEFAULT 0,
            applied INTEGER NOT NULL DEFAULT 0,
            created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS business_projects (
            id {id_type},
            user_id INTEGER NOT NULL,
            name TEXT NOT NULL,
            business_idea TEXT,
            target_customer TEXT,
            country TEXT,
            budget TEXT,
            notes TEXT,
            active INTEGER NOT NULL DEFAULT 0,
            created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP
        )
    """)

    cur.execute(sql("""
        CREATE INDEX IF NOT EXISTS idx_user_settings_user
        ON user_settings (user_id)
    """))

    cur.execute(sql("""
        CREATE INDEX IF NOT EXISTS idx_user_onboarding_user
        ON user_onboarding (user_id)
    """))

    cur.execute(sql("""
        CREATE INDEX IF NOT EXISTS idx_support_tickets_created
        ON support_tickets (created_at)
    """))

    cur.execute(sql("""
        CREATE INDEX IF NOT EXISTS idx_business_projects_user_active
        ON business_projects (user_id, active)
    """))

    cur.execute(sql("""
        CREATE INDEX IF NOT EXISTS idx_pricing_advice_user
        ON pricing_advice (user_id)
    """))

    cur.execute(sql("""
        CREATE INDEX IF NOT EXISTS idx_payment_guides_user
        ON payment_guides (user_id)
    """))

    cur.execute(sql("""
        CREATE INDEX IF NOT EXISTS idx_supplier_recommendations_user
        ON supplier_recommendations (user_id)
    """))

    cur.execute(sql("""
        CREATE INDEX IF NOT EXISTS idx_email_campaigns_user
        ON email_campaigns (user_id)
    """))

    cur.execute(sql("""
        CREATE INDEX IF NOT EXISTS idx_domain_guides_user
        ON domain_guides (user_id)
    """))

    cur.execute(sql("""
        CREATE INDEX IF NOT EXISTS idx_app_recommendations_user
        ON app_recommendations (user_id)
    """))

    cur.execute(sql("""
        CREATE UNIQUE INDEX IF NOT EXISTS idx_app_connections_user_platform
        ON app_connections (user_id, platform)
    """))

    cur.execute(sql("""
        CREATE INDEX IF NOT EXISTS idx_app_action_drafts_user
        ON app_action_drafts (user_id)
    """))

    cur.execute(sql("""
        CREATE INDEX IF NOT EXISTS idx_user_packages_user_status
        ON user_packages (user_id, status)
    """))

    cur.execute(sql("""
        DELETE FROM payments
        WHERE reference IS NOT NULL
        AND id NOT IN (
            SELECT MIN(id)
            FROM payments
            WHERE reference IS NOT NULL
            GROUP BY reference
        )
    """))

    cur.execute(sql("""
        CREATE UNIQUE INDEX IF NOT EXISTS idx_payments_reference_unique
        ON payments (reference)
    """))

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS shopify_plans (
            id {id_type},
            user_id INTEGER NOT NULL,
            title TEXT NOT NULL,
            content TEXT NOT NULL
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS canva_branding_packages (
            id {id_type},
            user_id INTEGER NOT NULL,
            title TEXT NOT NULL,
            content TEXT NOT NULL
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS canva_design_briefs (
            id {id_type},
            user_id INTEGER NOT NULL,
            title TEXT NOT NULL,
            content TEXT NOT NULL
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS canva_designs (
            id {id_type},
            user_id INTEGER NOT NULL,
            title TEXT NOT NULL,
            canva_design_id TEXT NOT NULL,
            edit_url TEXT,
            view_url TEXT,
            status TEXT NOT NULL
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS build_quotes (
            id {id_type},
            user_id INTEGER NOT NULL,
            title TEXT NOT NULL,
            shopify_plan TEXT NOT NULL,
            canva_plan TEXT NOT NULL,
            estimated_shopify_cost TEXT NOT NULL,
            estimated_canva_cost TEXT NOT NULL,
            service_fee TEXT NOT NULL,
            total_estimate TEXT NOT NULL,
            content TEXT NOT NULL
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS shopify_connections (
            id {id_type},
            user_id INTEGER NOT NULL,
            shop_domain TEXT NOT NULL,
            access_token TEXT NOT NULL,
            status TEXT NOT NULL
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS canva_connections (
            id {id_type},
            user_id INTEGER NOT NULL,
            status TEXT NOT NULL,
            connected_email TEXT,
            access_token TEXT,
            refresh_token TEXT
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS canva_oauth_sessions (
            id {id_type},
            user_id INTEGER NOT NULL,
            state TEXT NOT NULL,
            code_verifier TEXT NOT NULL
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS shopify_products (
            id {id_type},
            user_id INTEGER NOT NULL,
            title TEXT NOT NULL,
            description TEXT NOT NULL,
            shopify_product_id TEXT NOT NULL,
            status TEXT NOT NULL
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS shopify_collections (
            id {id_type},
            user_id INTEGER NOT NULL,
            title TEXT NOT NULL,
            description TEXT NOT NULL,
            shopify_collection_id TEXT NOT NULL,
            status TEXT NOT NULL
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS shopify_pages (
            id {id_type},
            user_id INTEGER NOT NULL,
            title TEXT NOT NULL,
            page_type TEXT NOT NULL,
            content TEXT NOT NULL,
            shopify_page_id TEXT NOT NULL,
            status TEXT NOT NULL
        )
    """)

    execute_schema(f"""
        CREATE TABLE IF NOT EXISTS usage_logs (
            id {id_type},
            user_id INTEGER NOT NULL,
            action_type TEXT NOT NULL,
            created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP
        )
    """)

    cur.execute(sql("""
        CREATE INDEX IF NOT EXISTS idx_usage_logs_user_action_created
        ON usage_logs (user_id, action_type, created_at)
    """))

    conn.commit()
    conn.close()


# -----------------------------
# USER / AUTH / CHAT HELPERS
# -----------------------------

def save_message(user_id, chat_id, role, content):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            INSERT INTO messages (user_id, chat_id, role, content)
            VALUES (?, ?, ?, ?)
        """),
        (user_id, chat_id, role, content)
    )

    conn.commit()
    conn.close()


def get_memory(user_id, chat_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT role, content
            FROM messages
            WHERE user_id = ?
            AND chat_id = ?
            ORDER BY id DESC
            LIMIT 10
        """),
        (user_id, chat_id)
    )

    rows = cur.fetchall()
    conn.close()

    rows.reverse()

    return [
        {
            "role": role,
            "content": content
        }
        for role, content in rows
    ]


def get_chat_messages(user_id, chat_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT role, content
            FROM messages
            WHERE user_id = ?
            AND chat_id = ?
            ORDER BY id ASC
        """),
        (user_id, chat_id)
    )

    rows = cur.fetchall()
    conn.close()

    return [
        {
            "role": role,
            "content": content
        }
        for role, content in rows
    ]


def create_chat(user_id, title="New Chat"):
    conn = db()
    cur = conn.cursor()

    if using_postgres():
        cur.execute(
            sql("""
            INSERT INTO chats (user_id, title)
            VALUES (%s, %s)
            RETURNING id
            """),
            (user_id, title)
        )

        chat_id = cur.fetchone()[0]

    else:
        cur.execute(
            sql("""
            INSERT INTO chats (user_id, title)
            VALUES (?, ?)
            """),
            (user_id, title)
        )

        chat_id = cur.lastrowid

    conn.commit()
    conn.close()

    return chat_id


def update_chat_title(chat_id, title):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            UPDATE chats
            SET title = ?
            WHERE id = ?
        """),
        (title, chat_id)
    )

    conn.commit()
    conn.close()


def get_chats(user_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT id, title
            FROM chats
            WHERE user_id = ?
            ORDER BY id DESC
        """),
        (user_id,)
    )

    chats = cur.fetchall()
    conn.close()

    return chats


# -----------------------------
# PROJECT HELPERS
# -----------------------------

def save_project(user_id, title, business_type, description):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            INSERT INTO projects (user_id, title, business_type, description)
            VALUES (?, ?, ?, ?)
        """),
        (user_id, title, business_type, description)
    )

    conn.commit()
    conn.close()


def get_projects(user_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT title, business_type, description
            FROM projects
            WHERE user_id = ?
            ORDER BY id DESC
        """),
        (user_id,)
    )

    projects = cur.fetchall()
    conn.close()

    return projects


# -----------------------------
# UPLOAD / FILE HELPERS
# -----------------------------

def get_latest_uploaded_file(user_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT filepath
            FROM uploads
            WHERE user_id = ?
            ORDER BY id DESC
            LIMIT 1
        """),
        (user_id,)
    )

    file = cur.fetchone()
    conn.close()

    return file[0] if file else None


def extract_file_text(filepath):
    if filepath.endswith(".txt"):
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()

    if filepath.endswith(".pdf"):
        reader = PdfReader(filepath)
        text = ""

        for page in reader.pages:
            text += page.extract_text() or ""

        return text

    if filepath.endswith(".docx"):
        document = Document(filepath)
        text = ""

        for paragraph in document.paragraphs:
            text += paragraph.text + "\n"

        return text

    return ""


def image_to_base64(filepath):
    with open(filepath, "rb") as image_file:
        return base64.b64encode(
            image_file.read()
        ).decode("utf-8")


# -----------------------------
# PAYMENT HELPERS
# -----------------------------

def save_payment(user_id, provider, amount, status, reference):
    if reference and payment_reference_exists(reference):
        return False

    conn = db()
    cur = conn.cursor()

    try:
        cur.execute(
            sql("""
                INSERT INTO payments (user_id, provider, amount, status, reference)
                VALUES (?, ?, ?, ?, ?)
            """),
            (user_id, provider, amount, status, reference)
        )
    except (sqlite3.IntegrityError, psycopg2.IntegrityError):
        conn.rollback()
        conn.close()
        return False

    conn.commit()
    conn.close()

    return True


def payment_reference_exists(reference):
    if not reference:
        return False

    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT id
            FROM payments
            WHERE reference = ?
            LIMIT 1
        """),
        (reference,)
    )

    payment = cur.fetchone()
    conn.close()

    return payment is not None


def get_payment_user_id_by_reference(reference):
    if not reference:
        return None

    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT user_id
            FROM payments
            WHERE reference = ?
            LIMIT 1
        """),
        (reference,)
    )

    payment = cur.fetchone()
    conn.close()

    return payment[0] if payment else None


def verify_paystack_transaction(reference):
    paystack_secret = os.getenv("PAYSTACK_SECRET_KEY")

    if not paystack_secret:
        return None, "Paystack secret key is missing."

    headers = {
        "Authorization": f"Bearer {paystack_secret}"
    }

    try:
        response = requests.get(
            "https://api.paystack.co/transaction/verify/"
            + urllib.parse.quote(reference, safe=""),
            headers=headers,
            timeout=20
        )
        response.raise_for_status()
        result = response.json()
    except (requests.RequestException, ValueError):
        return None, "Could not verify the Paystack transaction."

    transaction = result.get("data") or {}

    if not result.get("status") or transaction.get("status") != "success":
        return None, "Paystack could not confirm a successful payment."

    return transaction, None


PACKAGE_LEVELS = {
    "Starter": 1,
    "Pro": 2,
    "Premium Build": 3
}

PAYSTACK_PLANS = {
    "starter": {
        "package_name": "Starter",
        "amount": 49900,
        "currency": "ZAR"
    },
    "pro": {
        "package_name": "Pro",
        "amount": 99900,
        "currency": "ZAR"
    },
    "premium": {
        "package_name": "Premium Build",
        "amount": 199900,
        "currency": "ZAR"
    }
}

PLAN_DETAILS = {
    "Starter": {
        "best_for": "Beginners who need business planning, product ideas, and launch guidance.",
        "price": "R499"
    },
    "Pro": {
        "best_for": "Entrepreneurs who want Shopify, Canva, product, pricing, and marketing assets.",
        "price": "R999"
    },
    "Premium Build": {
        "best_for": "Users who want the most complete AI-assisted store and business launch package.",
        "price": "R1,999"
    }
}


def get_paystack_plan(plan_name, default=None):
    if not isinstance(plan_name, str):
        return PAYSTACK_PLANS.get(default)

    return PAYSTACK_PLANS.get(plan_name.strip().lower())


def get_paystack_metadata_plan(transaction):
    metadata = transaction.get("metadata")

    if not isinstance(metadata, dict):
        return None

    plan_name = metadata.get("plan")

    if isinstance(plan_name, str):
        return plan_name

    custom_fields = metadata.get("custom_fields")

    if not isinstance(custom_fields, list):
        return None

    for field in custom_fields:
        if (
            isinstance(field, dict)
            and field.get("variable_name") == "plan"
        ):
            return field.get("value")

    return None


def resolve_paystack_plan(transaction, fallback_plan=None):
    metadata_plan = get_paystack_metadata_plan(transaction)
    requested_plan = metadata_plan or fallback_plan or "starter"
    plan = get_paystack_plan(requested_plan)

    if not plan:
        return None

    amount = transaction.get("amount")
    currency = transaction.get("currency")

    if (
        not isinstance(amount, (int, float))
        or isinstance(amount, bool)
        or amount != plan["amount"]
        or (currency is not None and currency != plan["currency"])
    ):
        return None

    return plan

PACKAGE_USAGE_LIMITS = {
    "Starter": {
        "chat_message": {"limit": 50, "period": "daily"},
        "business_plan": {"limit": 3, "period": "total"},
        "shopify_plan": {"limit": 3, "period": "total"},
        "product_research": {"limit": 3, "period": "total"},
        "pricing_advice": {"limit": 3, "period": "total"},
        "supplier_guide": {"limit": 3, "period": "total"},
        "payment_guide": {"limit": 3, "period": "total"},
        "domain_guide": {"limit": 3, "period": "total"},
        "email_campaign": {"limit": 0, "period": "total"},
        "app_action_draft": {"limit": 0, "period": "total"},
        "shopify_product": {"limit": 0, "period": "total"},
        "canva_branding": {"limit": 0, "period": "total"},
        "canva_design_brief": {"limit": 0, "period": "total"},
        "canva_design": {"limit": 0, "period": "total"},
        "launch_package": {"limit": 0, "period": "total"},
        "store_agent_task": {"limit": 3, "period": "total"},
        "pdf_export": {"limit": 20, "period": "total"}
    },
    "Pro": {
        "chat_message": {"limit": 150, "period": "daily"},
        "business_plan": {"limit": 10, "period": "total"},
        "shopify_plan": {"limit": 10, "period": "total"},
        "product_research": {"limit": 10, "period": "total"},
        "pricing_advice": {"limit": 10, "period": "total"},
        "supplier_guide": {"limit": 10, "period": "total"},
        "payment_guide": {"limit": 10, "period": "total"},
        "domain_guide": {"limit": 10, "period": "total"},
        "email_campaign": {"limit": 10, "period": "total"},
        "app_action_draft": {"limit": 0, "period": "total"},
        "store_agent_task": {"limit": 20, "period": "total"},
        "shopify_product": {"limit": 10, "period": "total"},
        "canva_branding": {"limit": 20, "period": "total"},
        "canva_design_brief": {"limit": 20, "period": "total"},
        "canva_design": {"limit": 10, "period": "total"},
        "launch_package": {"limit": 10, "period": "total"},
        "pdf_export": {"limit": 50, "period": "total"}
    },
    "Premium Build": {
        "chat_message": {"limit": 300, "period": "daily"},
        "business_plan": {"limit": None, "period": "total"},
        "shopify_plan": {"limit": None, "period": "total"},
        "product_research": {"limit": None, "period": "total"},
        "pricing_advice": {"limit": None, "period": "total"},
        "supplier_guide": {"limit": None, "period": "total"},
        "payment_guide": {"limit": None, "period": "total"},
        "domain_guide": {"limit": None, "period": "total"},
        "email_campaign": {"limit": 30, "period": "total"},
        "app_action_draft": {"limit": None, "period": "total"},
        "store_agent_task": {"limit": None, "period": "total"},
        "shopify_product": {"limit": 30, "period": "total"},
        "canva_branding": {"limit": None, "period": "total"},
        "canva_design_brief": {"limit": None, "period": "total"},
        "canva_design": {"limit": 30, "period": "total"},
        "launch_package": {"limit": 50, "period": "total"},
        "pdf_export": {"limit": 100, "period": "total"}
    }
}


def set_user_package(user_id, package_name):
    if package_name not in PACKAGE_LEVELS:
        raise ValueError("Unknown package tier.")

    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            UPDATE user_packages
            SET status = ?
            WHERE user_id = ?
            AND status = ?
        """),
        ("inactive", user_id, "active")
    )

    cur.execute(
        sql("""
            INSERT INTO user_packages (user_id, package_name, status)
            VALUES (?, ?, ?)
        """),
        (user_id, package_name, "active")
    )

    conn.commit()
    conn.close()


def get_user_package(user_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT package_name
            FROM user_packages
            WHERE user_id = ?
            AND status = ?
            ORDER BY id DESC
            LIMIT 1
        """),
        (user_id, "active")
    )

    package = cur.fetchone()

    if not package:
        cur.execute(
            sql("""
                SELECT id
                FROM payments
                WHERE user_id = ?
                AND status = ?
                LIMIT 1
            """),
            (user_id, "success")
        )
        package = ("Starter",) if cur.fetchone() else None

    conn.close()

    return package[0] if package else None


def user_has_package(user_id, package_name):
    return get_user_package(user_id) == package_name


def user_package_at_least(user_id, package_name):
    current_package = get_user_package(user_id)

    return (
        current_package in PACKAGE_LEVELS
        and PACKAGE_LEVELS[current_package] >= PACKAGE_LEVELS[package_name]
    )


def user_is_paid(user_id):
    return get_user_package(user_id) is not None


def user_has_paid(user_id):
    """Compatibility alias for existing paid-user checks."""
    return user_is_paid(user_id)


def has_active_user_package(user_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT id
            FROM user_packages
            WHERE user_id = ?
            AND status = ?
            LIMIT 1
        """),
        (user_id, "active")
    )

    package = cur.fetchone()
    conn.close()

    return package is not None


def ensure_starter_package(user_id):
    if not has_active_user_package(user_id):
        set_user_package(user_id, "Starter")


def package_access_redirect(package_name, destination="/dashboard"):
    return redirect(
        f"{destination}?package_required={urllib.parse.quote(package_name)}"
    )


def get_package_required_message(package_name):
    if package_name not in PACKAGE_LEVELS:
        return ""

    package_messages = {
        "Starter": (
            "Starter includes planning, basic product/supplier/pricing/payment/domain guidance, "
            "three AI Store Agent drafts, launch readiness, PDF downloads, and support."
        ),
        "Pro": (
            "Upgrade to Pro to create Shopify draft products, generate full Canva branding "
            "briefs, use email marketing, receive app recommendations, and open the launch package."
        ),
        "Premium Build": (
            "Upgrade to Premium Build for the full AI Store Agent, advanced app connection "
            "drafts and safe applying, higher Shopify/email limits, priority support, "
            "and the premium launch package PDF."
        )
    }

    return (
        f"This feature requires the {package_name} package. "
        f"{package_messages.get(package_name, '')} "
        "Your current plan includes guidance and drafts, but this action requires a higher plan. "
        "Review the available packages to upgrade your account."
    )


def payment_record_exists_for_user(user_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT id
            FROM payments
            WHERE user_id = ?
            AND status = ?
            LIMIT 1
        """),
        (user_id, "success")
    )

    payment = cur.fetchone()
    conn.close()

    return payment is not None


def get_user_email(user_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT email
            FROM users
            WHERE id = ?
            LIMIT 1
        """),
        (user_id,)
    )

    user = cur.fetchone()
    conn.close()

    return user[0] if user else None


def get_user_id_by_email(email):
    if not email:
        return None

    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT id
            FROM users
            WHERE LOWER(email) = ?
            LIMIT 1
        """),
        (email.strip().lower(),)
    )

    user = cur.fetchone()
    conn.close()

    return user[0] if user else None


def get_latest_payment(user_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT provider, amount, status, reference
            FROM payments
            WHERE user_id = ?
            ORDER BY id DESC
            LIMIT 1
        """),
        (user_id,)
    )

    payment = cur.fetchone()
    conn.close()

    return payment


THEME_OPTIONS = {
    "default": "Default Blue",
    "dark": "Dark Mode",
    "light": "Light Mode",
    "teal": "Teal",
    "purple": "Purple",
    "gold": "Gold"
}


def normalize_theme(theme):
    theme = (theme or "default").strip().lower()
    return theme if theme in THEME_OPTIONS else "default"


def get_user_settings(user_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT theme
            FROM user_settings
            WHERE user_id = ?
            ORDER BY id DESC
            LIMIT 1
        """),
        (user_id,)
    )

    settings = cur.fetchone()
    conn.close()

    return {
        "theme": normalize_theme(settings[0] if settings else "default")
    }


def save_user_settings(user_id, theme):
    theme = normalize_theme(theme)
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT id
            FROM user_settings
            WHERE user_id = ?
            ORDER BY id DESC
            LIMIT 1
        """),
        (user_id,)
    )

    settings = cur.fetchone()

    if settings:
        cur.execute(
            sql("""
                UPDATE user_settings
                SET theme = ?,
                    updated_at = CURRENT_TIMESTAMP
                WHERE id = ?
                AND user_id = ?
            """),
            (theme, settings[0], user_id)
        )
    else:
        cur.execute(
            sql("""
                INSERT INTO user_settings (user_id, theme)
                VALUES (?, ?)
            """),
            (user_id, theme)
        )

    conn.commit()
    conn.close()

    return theme


def get_user_onboarding(user_id):
    conn = db()
    cur = conn.cursor()
    cur.execute(
        sql("""
            SELECT id, business_idea, country, budget, has_shopify,
                   has_canva, product_type, target_customer, business_goal,
                   completed
            FROM user_onboarding
            WHERE user_id = ?
            ORDER BY id DESC
            LIMIT 1
        """),
        (user_id,)
    )
    row = cur.fetchone()
    conn.close()
    return row


def save_user_onboarding(user_id, data):
    existing = get_user_onboarding(user_id)
    conn = db()
    cur = conn.cursor()
    values = (
        data.get("business_idea", ""),
        data.get("country", ""),
        data.get("budget", ""),
        data.get("has_shopify", ""),
        data.get("has_canva", ""),
        data.get("product_type", ""),
        data.get("target_customer", ""),
        data.get("business_goal", ""),
        1,
    )

    if existing:
        cur.execute(
            sql("""
                UPDATE user_onboarding
                SET business_idea = ?,
                    country = ?,
                    budget = ?,
                    has_shopify = ?,
                    has_canva = ?,
                    product_type = ?,
                    target_customer = ?,
                    business_goal = ?,
                    completed = ?,
                    updated_at = CURRENT_TIMESTAMP
                WHERE id = ?
                AND user_id = ?
            """),
            values + (existing[0], user_id)
        )
    else:
        cur.execute(
            sql("""
                INSERT INTO user_onboarding (
                    user_id, business_idea, country, budget, has_shopify,
                    has_canva, product_type, target_customer, business_goal,
                    completed
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """),
            (user_id,) + values
        )

    conn.commit()
    conn.close()


def create_business_project_record(
    user_id,
    name,
    business_idea="",
    target_customer="",
    country="",
    budget="",
    notes="",
    make_active=True
):
    conn = db()
    cur = conn.cursor()

    if make_active:
        cur.execute(
            sql("UPDATE business_projects SET active = 0 WHERE user_id = ?"),
            (user_id,)
        )

    if using_postgres():
        cur.execute(
            sql("""
                INSERT INTO business_projects (
                    user_id, name, business_idea, target_customer,
                    country, budget, notes, active
                )
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                RETURNING id
            """),
            (
                user_id, name, business_idea, target_customer,
                country, budget, notes, 1 if make_active else 0
            )
        )
        project_id = cur.fetchone()[0]
    else:
        cur.execute(
            sql("""
                INSERT INTO business_projects (
                    user_id, name, business_idea, target_customer,
                    country, budget, notes, active
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """),
            (
                user_id, name, business_idea, target_customer,
                country, budget, notes, 1 if make_active else 0
            )
        )
        project_id = cur.lastrowid

    conn.commit()
    conn.close()
    return project_id


def get_business_projects(user_id):
    conn = db()
    cur = conn.cursor()
    cur.execute(
        sql("""
            SELECT id, name, business_idea, target_customer, country,
                   budget, notes, active, created_at
            FROM business_projects
            WHERE user_id = ?
            ORDER BY active DESC, id DESC
        """),
        (user_id,)
    )
    projects = cur.fetchall()
    conn.close()
    return projects


def get_active_business_project(user_id):
    conn = db()
    cur = conn.cursor()
    cur.execute(
        sql("""
            SELECT id, name, business_idea, target_customer, country,
                   budget, notes, active, created_at
            FROM business_projects
            WHERE user_id = ?
            AND active = 1
            ORDER BY id DESC
            LIMIT 1
        """),
        (user_id,)
    )
    project = cur.fetchone()
    conn.close()
    return project


def get_business_project(user_id, project_id):
    conn = db()
    cur = conn.cursor()
    cur.execute(
        sql("""
            SELECT id, name, business_idea, target_customer, country,
                   budget, notes, active, created_at
            FROM business_projects
            WHERE user_id = ?
            AND id = ?
            LIMIT 1
        """),
        (user_id, project_id)
    )
    project = cur.fetchone()
    conn.close()
    return project


def set_active_business_project(user_id, project_id):
    conn = db()
    cur = conn.cursor()
    cur.execute(
        sql("UPDATE business_projects SET active = 0 WHERE user_id = ?"),
        (user_id,)
    )
    cur.execute(
        sql("""
            UPDATE business_projects
            SET active = 1,
                updated_at = CURRENT_TIMESTAMP
            WHERE user_id = ?
            AND id = ?
        """),
        (user_id, project_id)
    )
    conn.commit()
    conn.close()


def save_support_ticket(user_id, email, category, subject, message):
    conn = db()
    cur = conn.cursor()

    if using_postgres():
        cur.execute(
            sql("""
                INSERT INTO support_tickets (
                    user_id, email, category, subject, message, status
                )
                VALUES (%s, %s, %s, %s, %s, %s)
                RETURNING id
            """),
            (user_id, email, category, subject, message, "open")
        )
        ticket_id = cur.fetchone()[0]
    else:
        cur.execute(
            sql("""
                INSERT INTO support_tickets (
                    user_id, email, category, subject, message, status
                )
                VALUES (?, ?, ?, ?, ?, ?)
            """),
            (user_id, email, category, subject, message, "open")
        )
        ticket_id = cur.lastrowid

    conn.commit()
    conn.close()
    return ticket_id


def get_recent_support_tickets(limit=10):
    conn = db()
    cur = conn.cursor()
    cur.execute(
        sql("""
            SELECT id, email, category, subject, status, created_at
            FROM support_tickets
            ORDER BY id DESC
            LIMIT ?
        """),
        (limit,)
    )
    tickets = cur.fetchall()
    conn.close()
    return tickets


def save_pricing_advice(user_id, data, content):
    conn = db()
    cur = conn.cursor()

    if using_postgres():
        cur.execute(
            sql("""
                INSERT INTO pricing_advice (
                    user_id, business_idea, product_type, target_customer,
                    cost_price, desired_profit, competitor_price, country,
                    budget, content
                )
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                RETURNING id
            """),
            (
                user_id, data.get("business_idea", ""), data.get("product_type", ""),
                data.get("target_customer", ""), data.get("cost_price", ""),
                data.get("desired_profit", ""), data.get("competitor_price", ""),
                data.get("country", ""), data.get("budget", ""), content
            )
        )
        advice_id = cur.fetchone()[0]
    else:
        cur.execute(
            sql("""
                INSERT INTO pricing_advice (
                    user_id, business_idea, product_type, target_customer,
                    cost_price, desired_profit, competitor_price, country,
                    budget, content
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """),
            (
                user_id, data.get("business_idea", ""), data.get("product_type", ""),
                data.get("target_customer", ""), data.get("cost_price", ""),
                data.get("desired_profit", ""), data.get("competitor_price", ""),
                data.get("country", ""), data.get("budget", ""), content
            )
        )
        advice_id = cur.lastrowid

    conn.commit()
    conn.close()
    return advice_id


def get_pricing_advice_list(user_id):
    conn = db()
    cur = conn.cursor()
    cur.execute(
        sql("""
            SELECT id, business_idea, product_type, target_customer,
                   country, budget, content, created_at
            FROM pricing_advice
            WHERE user_id = ?
            ORDER BY id DESC
        """),
        (user_id,)
    )
    rows = cur.fetchall()
    conn.close()
    return rows


def save_payment_guide(user_id, data, content):
    conn = db()
    cur = conn.cursor()

    if using_postgres():
        cur.execute(
            sql("""
                INSERT INTO payment_guides (
                    user_id, country, business_type, selling_platform,
                    payment_options, budget, content
                )
                VALUES (%s, %s, %s, %s, %s, %s, %s)
                RETURNING id
            """),
            (
                user_id, data.get("country", ""), data.get("business_type", ""),
                data.get("selling_platform", ""), data.get("payment_options", ""),
                data.get("budget", ""), content
            )
        )
        guide_id = cur.fetchone()[0]
    else:
        cur.execute(
            sql("""
                INSERT INTO payment_guides (
                    user_id, country, business_type, selling_platform,
                    payment_options, budget, content
                )
                VALUES (?, ?, ?, ?, ?, ?, ?)
            """),
            (
                user_id, data.get("country", ""), data.get("business_type", ""),
                data.get("selling_platform", ""), data.get("payment_options", ""),
                data.get("budget", ""), content
            )
        )
        guide_id = cur.lastrowid

    conn.commit()
    conn.close()
    return guide_id


def get_payment_guides(user_id):
    conn = db()
    cur = conn.cursor()
    cur.execute(
        sql("""
            SELECT id, country, business_type, selling_platform,
                   payment_options, budget, content, created_at
            FROM payment_guides
            WHERE user_id = ?
            ORDER BY id DESC
        """),
        (user_id,)
    )
    rows = cur.fetchall()
    conn.close()
    return rows


def save_supplier_recommendations(user_id, data, content):
    conn = db()
    cur = conn.cursor()

    if using_postgres():
        cur.execute(
            sql("""
                INSERT INTO supplier_recommendations (
                    user_id, business_idea, product_type, country, budget,
                    sourcing_preference, risk_level, content
                )
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                RETURNING id
            """),
            (
                user_id, data.get("business_idea", ""), data.get("product_type", ""),
                data.get("country", ""), data.get("budget", ""),
                data.get("sourcing_preference", ""), data.get("risk_level", ""),
                content
            )
        )
        recommendation_id = cur.fetchone()[0]
    else:
        cur.execute(
            sql("""
                INSERT INTO supplier_recommendations (
                    user_id, business_idea, product_type, country, budget,
                    sourcing_preference, risk_level, content
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """),
            (
                user_id, data.get("business_idea", ""), data.get("product_type", ""),
                data.get("country", ""), data.get("budget", ""),
                data.get("sourcing_preference", ""), data.get("risk_level", ""),
                content
            )
        )
        recommendation_id = cur.lastrowid

    conn.commit()
    conn.close()
    return recommendation_id


def get_supplier_recommendations(user_id):
    conn = db()
    cur = conn.cursor()
    cur.execute(
        sql("""
            SELECT id, business_idea, product_type, country, budget,
                   sourcing_preference, risk_level, content, created_at
            FROM supplier_recommendations
            WHERE user_id = ?
            ORDER BY id DESC
        """),
        (user_id,)
    )
    rows = cur.fetchall()
    conn.close()
    return rows


def save_email_campaign(user_id, data, content):
    conn = db()
    cur = conn.cursor()
    values = (
        user_id, data.get("business_name", ""), data.get("business_type", ""),
        data.get("target_customer", ""), data.get("campaign_goal", ""),
        data.get("email_platform", ""), data.get("offer", ""),
        data.get("tone", ""), content
    )

    if using_postgres():
        cur.execute(
            sql("""
                INSERT INTO email_campaigns (
                    user_id, business_name, business_type, target_customer,
                    campaign_goal, email_platform, offer, tone, content
                )
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                RETURNING id
            """),
            values
        )
        campaign_id = cur.fetchone()[0]
    else:
        cur.execute(
            sql("""
                INSERT INTO email_campaigns (
                    user_id, business_name, business_type, target_customer,
                    campaign_goal, email_platform, offer, tone, content
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """),
            values
        )
        campaign_id = cur.lastrowid

    conn.commit()
    conn.close()
    return campaign_id


def get_email_campaigns(user_id):
    conn = db()
    cur = conn.cursor()
    cur.execute(
        sql("""
            SELECT id, business_name, business_type, target_customer,
                   campaign_goal, email_platform, offer, tone, content, created_at
            FROM email_campaigns
            WHERE user_id = ?
            ORDER BY id DESC
        """),
        (user_id,)
    )
    rows = cur.fetchall()
    conn.close()
    return rows


def get_email_campaign(user_id, campaign_id):
    conn = db()
    cur = conn.cursor()
    cur.execute(
        sql("""
            SELECT id, business_name, business_type, target_customer,
                   campaign_goal, email_platform, offer, tone, content, created_at
            FROM email_campaigns
            WHERE user_id = ? AND id = ?
            LIMIT 1
        """),
        (user_id, campaign_id)
    )
    row = cur.fetchone()
    conn.close()
    return row


def save_domain_guide(user_id, data, content):
    conn = db()
    cur = conn.cursor()
    values = (
        user_id, data.get("business_name", ""), data.get("preferred_domain", ""),
        data.get("budget", ""), data.get("provider_preference", ""),
        data.get("website_platform", ""), data.get("country", ""), content
    )

    if using_postgres():
        cur.execute(
            sql("""
                INSERT INTO domain_guides (
                    user_id, business_name, preferred_domain, budget,
                    provider_preference, website_platform, country, content
                )
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                RETURNING id
            """),
            values
        )
        guide_id = cur.fetchone()[0]
    else:
        cur.execute(
            sql("""
                INSERT INTO domain_guides (
                    user_id, business_name, preferred_domain, budget,
                    provider_preference, website_platform, country, content
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """),
            values
        )
        guide_id = cur.lastrowid

    conn.commit()
    conn.close()
    return guide_id


def get_domain_guides(user_id):
    conn = db()
    cur = conn.cursor()
    cur.execute(
        sql("""
            SELECT id, business_name, preferred_domain, budget,
                   provider_preference, website_platform, country, content, created_at
            FROM domain_guides
            WHERE user_id = ?
            ORDER BY id DESC
        """),
        (user_id,)
    )
    rows = cur.fetchall()
    conn.close()
    return rows


def get_domain_guide(user_id, guide_id):
    conn = db()
    cur = conn.cursor()
    cur.execute(
        sql("""
            SELECT id, business_name, preferred_domain, budget,
                   provider_preference, website_platform, country, content, created_at
            FROM domain_guides
            WHERE user_id = ? AND id = ?
            LIMIT 1
        """),
        (user_id, guide_id)
    )
    row = cur.fetchone()
    conn.close()
    return row


def save_app_recommendation(user_id, data, content):
    conn = db()
    cur = conn.cursor()
    values = (
        user_id, data.get("business_type", ""), data.get("country", ""),
        data.get("budget", ""), data.get("goal", ""),
        data.get("selling_platform", ""), content
    )
    if using_postgres():
        cur.execute(sql("""
            INSERT INTO app_recommendations (
                user_id, business_type, country, budget, goal,
                selling_platform, content
            ) VALUES (%s, %s, %s, %s, %s, %s, %s)
            RETURNING id
        """), values)
        recommendation_id = cur.fetchone()[0]
    else:
        cur.execute(sql("""
            INSERT INTO app_recommendations (
                user_id, business_type, country, budget, goal,
                selling_platform, content
            ) VALUES (?, ?, ?, ?, ?, ?, ?)
        """), values)
        recommendation_id = cur.lastrowid
    conn.commit()
    conn.close()
    return recommendation_id


def get_app_recommendations(user_id):
    conn = db()
    cur = conn.cursor()
    cur.execute(sql("""
        SELECT id, business_type, country, budget, goal,
               selling_platform, content, created_at
        FROM app_recommendations
        WHERE user_id = ?
        ORDER BY id DESC
    """), (user_id,))
    rows = cur.fetchall()
    conn.close()
    return rows


def save_app_connection(user_id, platform, category, connection_data):
    access_token = connection_data.get("access_token", "").strip()
    refresh_token = connection_data.get("refresh_token", "").strip()
    api_key = connection_data.get("api_key", "").strip()

    encrypted_access_token = encrypt_token(access_token) if access_token else None
    encrypted_refresh_token = encrypt_token(refresh_token) if refresh_token else None
    encrypted_api_key = encrypt_token(api_key) if api_key else None

    conn = db()
    cur = conn.cursor()
    cur.execute(sql("""
        SELECT id FROM app_connections
        WHERE user_id = ? AND platform = ?
        LIMIT 1
    """), (user_id, platform))
    existing = cur.fetchone()
    values = (
        category,
        connection_data.get("connection_status", "configured"),
        encrypted_access_token,
        encrypted_refresh_token,
        encrypted_api_key,
        connection_data.get("account_name", "").strip(),
        connection_data.get("scopes", "").strip()
    )

    if existing:
        cur.execute(sql("""
            UPDATE app_connections
            SET category = ?, connection_status = ?,
                access_token_encrypted = ?, refresh_token_encrypted = ?,
                api_key_encrypted = ?, account_name = ?, scopes = ?,
                updated_at = CURRENT_TIMESTAMP
            WHERE user_id = ? AND platform = ?
        """), values + (user_id, platform))
        connection_id = existing[0]
    elif using_postgres():
        cur.execute(sql("""
            INSERT INTO app_connections (
                user_id, platform, category, connection_status,
                access_token_encrypted, refresh_token_encrypted,
                api_key_encrypted, account_name, scopes
            ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
            RETURNING id
        """), (user_id, platform) + values)
        connection_id = cur.fetchone()[0]
    else:
        cur.execute(sql("""
            INSERT INTO app_connections (
                user_id, platform, category, connection_status,
                access_token_encrypted, refresh_token_encrypted,
                api_key_encrypted, account_name, scopes
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """), (user_id, platform) + values)
        connection_id = cur.lastrowid

    conn.commit()
    conn.close()
    return connection_id


def get_app_connections(user_id):
    conn = db()
    cur = conn.cursor()
    cur.execute(sql("""
        SELECT id, platform, category, connection_status,
               access_token_encrypted, refresh_token_encrypted,
               api_key_encrypted, account_name, scopes, created_at, updated_at
        FROM app_connections
        WHERE user_id = ?
        ORDER BY platform
    """), (user_id,))
    rows = cur.fetchall()
    conn.close()
    return rows


def get_app_connection(user_id, platform):
    conn = db()
    cur = conn.cursor()
    cur.execute(sql("""
        SELECT id, platform, category, connection_status,
               access_token_encrypted, refresh_token_encrypted,
               api_key_encrypted, account_name, scopes, created_at, updated_at
        FROM app_connections
        WHERE user_id = ? AND platform = ?
        LIMIT 1
    """), (user_id, platform))
    row = cur.fetchone()
    conn.close()
    return row


def disconnect_app_connection(user_id, platform):
    conn = db()
    cur = conn.cursor()
    cur.execute(sql("""
        DELETE FROM app_connections
        WHERE user_id = ? AND platform = ?
    """), (user_id, platform))
    conn.commit()
    conn.close()


def create_app_action_draft_record(user_id, platform, action_type, title, content):
    conn = db()
    cur = conn.cursor()
    values = (user_id, platform, action_type, title, content, "Needs Approval", 0, 0)
    if using_postgres():
        cur.execute(sql("""
            INSERT INTO app_action_drafts (
                user_id, platform, action_type, title, draft_content,
                status, approved, applied
            ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
            RETURNING id
        """), values)
        draft_id = cur.fetchone()[0]
    else:
        cur.execute(sql("""
            INSERT INTO app_action_drafts (
                user_id, platform, action_type, title, draft_content,
                status, approved, applied
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """), values)
        draft_id = cur.lastrowid
    conn.commit()
    conn.close()
    return draft_id


def get_app_action_drafts(user_id):
    conn = db()
    cur = conn.cursor()
    cur.execute(sql("""
        SELECT id, platform, action_type, title, draft_content, status,
               approved, applied, created_at, updated_at
        FROM app_action_drafts
        WHERE user_id = ?
        ORDER BY id DESC
    """), (user_id,))
    rows = cur.fetchall()
    conn.close()
    return rows


def get_app_action_draft(user_id, draft_id):
    conn = db()
    cur = conn.cursor()
    cur.execute(sql("""
        SELECT id, platform, action_type, title, draft_content, status,
               approved, applied, created_at, updated_at
        FROM app_action_drafts
        WHERE user_id = ? AND id = ?
        LIMIT 1
    """), (user_id, draft_id))
    row = cur.fetchone()
    conn.close()
    return row


def update_app_action_draft_content(user_id, draft_id, content):
    conn = db()
    cur = conn.cursor()
    cur.execute(sql("""
        UPDATE app_action_drafts
        SET draft_content = ?, status = 'Needs Approval', approved = 0,
            applied = 0, updated_at = CURRENT_TIMESTAMP
        WHERE user_id = ? AND id = ?
    """), (content, user_id, draft_id))
    conn.commit()
    conn.close()


def update_app_action_draft_status(user_id, draft_id, status, approved, applied):
    conn = db()
    cur = conn.cursor()
    cur.execute(sql("""
        UPDATE app_action_drafts
        SET status = ?, approved = ?, applied = ?, updated_at = CURRENT_TIMESTAMP
        WHERE user_id = ? AND id = ?
    """), (status, int(bool(approved)), int(bool(applied)), user_id, draft_id))
    conn.commit()
    conn.close()


def delete_app_action_draft_record(user_id, draft_id):
    conn = db()
    cur = conn.cursor()
    cur.execute(sql("""
        DELETE FROM app_action_drafts
        WHERE user_id = ? AND id = ?
    """), (user_id, draft_id))
    conn.commit()
    conn.close()


def normalize_app_platform(platform):
    return (platform or "").strip().lower().replace("-", "_")


def get_app_connection_status(user_id, platform):
    if platform in {"shopify", "shopify_email"}:
        connection = get_shopify_connection(user_id)
        return "Connected" if connection and connection[3] == "connected" else "Not connected"
    if platform == "canva":
        connection = get_canva_connection(user_id)
        return "Connected" if connection and connection[2] == "connected" else "Not connected"

    catalog_item = APP_CATALOG[platform]
    if catalog_item["connection_mode"] == "guidance":
        return "Guidance only"
    connection = get_app_connection(user_id, platform)
    return "Configured" if connection and connection[3] == "configured" else "Not connected"


def get_app_catalog_cards(user_id):
    cards = []
    for platform, details in APP_CATALOG.items():
        card = dict(details)
        card.update({
            "platform": platform,
            "status": get_app_connection_status(user_id, platform),
            "guide_url": f"/app_connection_guide/{platform}",
            "connect_url": f"/connect_app/{platform}",
            "draft_url": (
                f"/create_app_action_draft/{platform}/{details['draft_actions'][0]}"
                if details["draft_actions"] else ""
            )
        })
        cards.append(card)
    return cards


def get_connected_app_summaries(user_id):
    summaries = []
    shopify = get_shopify_connection(user_id)
    canva = get_canva_connection(user_id)
    if shopify and shopify[3] == "connected":
        summaries.append({"platform": "shopify", "name": "Shopify", "category": "Store/ecommerce", "account_name": shopify[1], "status": "Connected"})
    if canva and canva[2] == "connected":
        summaries.append({"platform": "canva", "name": "Canva", "category": "Design", "account_name": canva[3] or "Canva account", "status": "Connected"})
    for connection in get_app_connections(user_id):
        details = APP_CATALOG.get(connection[1], {})
        summaries.append({
            "platform": connection[1],
            "name": details.get("name", connection[1].replace("_", " ").title()),
            "category": connection[2],
            "account_name": connection[7] or "Account configured",
            "status": "Configured"
        })
    return summaries


def get_app_action_risk(action_type):
    risks = {
        "email_campaign": {"creates": "A campaign draft saved inside BusinessBuilder AI", "cost": "No automatic charge", "public": "No email is sent"},
        "social_ad_draft": {"creates": "Ad copy and a campaign plan saved inside BusinessBuilder AI", "cost": "No ad spend is created", "public": "No ad is published"},
        "canva_design_brief": {"creates": "A Canva-ready design brief", "cost": "No automatic charge", "public": "Nothing is published"},
        "shopify_product_draft": {"creates": "Draft Shopify products if Shopify is connected", "cost": "No automatic charge", "public": "Products remain drafts"},
        "webflow_page_draft": {"creates": "A Webflow page/CMS draft inside BusinessBuilder AI", "cost": "No automatic charge", "public": "No website page is published"},
        "website_copy": {"creates": "Website/store copy saved inside BusinessBuilder AI", "cost": "No automatic charge", "public": "No website change is published"},
        "automation_plan": {"creates": "A manual automation plan", "cost": "No automatic charge", "public": "No automation is enabled"}
    }
    return risks[action_type]


def build_project_context(user_id):
    project = get_active_business_project(user_id)

    if not project:
        return ""

    return (
        "Active business project context:\n"
        f"Project name: {project[1]}\n"
        f"Business idea: {project[2] or 'Not provided'}\n"
        f"Target customer: {project[3] or 'Not provided'}\n"
        f"Country: {project[4] or 'Not provided'}\n"
        f"Budget: {project[5] or 'Not provided'}\n"
        f"Notes: {project[6] or 'Not provided'}"
    )


def get_launch_readiness(user_id):
    completed_steps = get_completed_steps(user_id)
    business_plans = get_business_plans(user_id)
    product_research_list = get_product_research_list(user_id)
    store_agent_tasks = get_store_agent_tasks(user_id)
    latest_tasks = get_latest_store_agent_tasks_by_type(user_id)
    shopify_connection = get_shopify_connection(user_id)
    canva_connection = get_canva_connection(user_id)
    shopify_products = get_shopify_products(user_id)
    canva_branding = get_canva_branding_packages(user_id)
    canva_briefs = get_canva_design_briefs(user_id)
    ai_store_builds = get_ai_store_builds(user_id)
    domain_guides = get_domain_guides(user_id)
    email_campaigns = get_email_campaigns(user_id)

    checks = [
        {
            "title": "Business workflow completed",
            "complete": len(completed_steps) >= len(WORKFLOW_STEPS),
            "description": "Complete all guided setup steps.",
            "url": "/business_workflow",
            "action": "Open Workflow"
        },
        {
            "title": "Product research completed",
            "complete": bool(product_research_list),
            "description": "Generate product ideas, sourcing notes, and competitor research.",
            "url": "/product_finder",
            "action": "Find Products"
        },
        {
            "title": "Business plan generated",
            "complete": bool(business_plans),
            "description": "Create a saved business plan from workflow answers.",
            "url": "/generate_business_plan",
            "action": "Generate Plan"
        },
        {
            "title": "AI Store Agent task created",
            "complete": bool(store_agent_tasks),
            "description": "Create at least one store agent draft.",
            "url": "/ai_store_agent",
            "action": "Open Agent"
        },
        {
            "title": "Shopify connected",
            "complete": bool(shopify_connection and shopify_connection[3] == "connected"),
            "description": "Connect Shopify for approved draft store actions.",
            "url": "/shopify_settings",
            "action": "Connect Shopify"
        },
        {
            "title": "Canva connected",
            "complete": bool(canva_connection and canva_connection[2] == "connected"),
            "description": "Connect Canva for editable design drafts.",
            "url": "/canva_settings",
            "action": "Connect Canva"
        },
        {
            "title": "Shopify draft products created",
            "complete": bool(shopify_products),
            "description": "Create draft Shopify products for review.",
            "url": "/build_approval?action=shopify_products",
            "action": "Create Products"
        },
        {
            "title": "Canva branding or brief created",
            "complete": bool(canva_branding or canva_briefs or latest_tasks.get("canva_branding")),
            "description": "Generate branding direction or Canva design briefs.",
            "url": "/ai_store_agent",
            "action": "Create Branding"
        },
        {
            "title": "Homepage/store draft approved",
            "complete": bool(
                latest_tasks.get("homepage_design")
                and latest_tasks["homepage_design"][6]
            ),
            "description": "Approve a homepage or store draft before applying safe changes.",
            "url": "/ai_store_agent",
            "action": "Review Store Drafts"
        },
        {
            "title": "Shipping plan created",
            "complete": bool(latest_tasks.get("shipping_zones")),
            "description": "Create shipping zone guidance.",
            "url": "/generate_store_agent_task/shipping_zones",
            "action": "Create Shipping Plan"
        },
        {
            "title": "Payment setup checklist created",
            "complete": bool(latest_tasks.get("payments_setup")),
            "description": "Create a payment setup checklist.",
            "url": "/generate_store_agent_task/payments_setup",
            "action": "Create Payment Checklist"
        },
        {
            "title": "Domain setup guidance created",
            "complete": bool(domain_guides or latest_tasks.get("domain_setup")),
            "description": "Create optional domain and DNS guidance without buying anything automatically.",
            "url": "/domain_helper",
            "action": "Open Domain Helper",
            "optional": True
        },
        {
            "title": "Email marketing campaign created",
            "complete": bool(email_campaigns),
            "description": "Create an optional campaign draft to support the launch.",
            "url": "/email_marketing",
            "action": "Create Email Campaign",
            "optional": True
        },
        {
            "title": "Launch package generated/downloaded",
            "complete": bool(ai_store_builds and latest_tasks.get("launch_checklist")),
            "description": "Generate your launch checklist and package.",
            "url": "/launch_package",
            "action": "Open Launch Package"
        }
    ]
    completed = sum(1 for check in checks if check["complete"])
    core_checks = [check for check in checks if not check.get("optional")]
    optional_checks = [check for check in checks if check.get("optional")]
    core_completed = sum(1 for check in core_checks if check["complete"])
    optional_completed = sum(1 for check in optional_checks if check["complete"])
    core_score = int(round((core_completed / len(core_checks)) * 90)) if core_checks else 0
    optional_score = int(round((optional_completed / len(optional_checks)) * 10)) if optional_checks else 0
    score = min(100, core_score + optional_score)
    next_action = next(
        (check for check in checks if not check["complete"]),
        {
            "title": "Review your final launch package",
            "description": "Your core checklist is complete. Review and launch manually when ready.",
            "url": "/launch_package",
            "action": "View Launch Package"
        }
    )

    return {
        "score": score,
        "completed": completed,
        "total": len(checks),
        "checks": checks,
        "next_action": next_action
    }


@app.context_processor
def inject_global_template_context():
    user_id = session.get("user_id")

    if not user_id:
        return {
            "current_user_email": "",
            "current_user_theme": "default",
            "theme_options": THEME_OPTIONS
        }

    return {
        "current_user_email": get_user_email(user_id) or "",
        "current_user_theme": get_user_settings(user_id)["theme"],
        "theme_options": THEME_OPTIONS
    }


USAGE_LABELS = {
    "chat_message": "chat messages",
    "business_plan": "business plan generation",
    "shopify_plan": "Shopify setup plan generation",
    "product_research": "AI Product Finder research reports",
    "pricing_advice": "pricing advice reports",
    "supplier_guide": "supplier guides",
    "payment_guide": "payment setup guides",
    "domain_guide": "domain and DNS guides",
    "email_campaign": "email campaign drafts",
    "app_action_draft": "advanced app action drafts",
    "store_agent_task": "AI Store Agent task drafts",
    "shopify_product": "Shopify product creation",
    "canva_branding": "Canva branding package generation",
    "canva_design_brief": "Canva design brief generation",
    "canva_design": "Canva design draft creation",
    "launch_package": "launch package views",
    "pdf_export": "PDF exports"
}


def log_usage(user_id, action_type):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            INSERT INTO usage_logs (user_id, action_type)
            VALUES (?, ?)
        """),
        (user_id, action_type)
    )

    conn.commit()
    conn.close()


def get_usage_count(user_id, action_type, period="daily"):
    conn = db()
    cur = conn.cursor()

    if period == "daily":
        cur.execute(
            sql("""
                SELECT COUNT(*)
                FROM usage_logs
                WHERE user_id = ?
                AND action_type = ?
                AND created_at >= CURRENT_DATE
            """),
            (user_id, action_type)
        )
    elif period == "total":
        cur.execute(
            sql("""
                SELECT COUNT(*)
                FROM usage_logs
                WHERE user_id = ?
                AND action_type = ?
            """),
            (user_id, action_type)
        )
    else:
        conn.close()
        raise ValueError("Usage period must be daily or total.")

    usage_count = cur.fetchone()[0]
    conn.close()

    return usage_count


def usage_limit_reached(user_id, action_type):
    package_name = get_user_package(user_id) or "Starter"
    package_limits = PACKAGE_USAGE_LIMITS.get(
        package_name,
        PACKAGE_USAGE_LIMITS["Starter"]
    )
    limit_config = package_limits.get(action_type)

    if not limit_config:
        return True

    if limit_config["limit"] is None:
        return False

    return get_usage_count(
        user_id,
        action_type,
        limit_config["period"]
    ) >= limit_config["limit"]


def usage_limit_redirect(action_type, destination="/dashboard"):
    return redirect(
        f"{destination}?usage_limit={urllib.parse.quote(action_type)}"
    )


def get_usage_limit_message(action_type, user_id=None):
    if not action_type:
        return ""

    package_name = get_user_package(user_id) if user_id else "Starter"
    package_name = package_name or "Starter"
    package_limits = PACKAGE_USAGE_LIMITS.get(
        package_name,
        PACKAGE_USAGE_LIMITS["Starter"]
    )
    limit_config = package_limits.get(action_type)

    if not limit_config:
        return (
            f"{USAGE_LABELS.get(action_type, 'This action')} is not included "
            f"in the {package_name} package. Review the available packages to upgrade."
        )

    if limit_config["limit"] is None:
        return ""

    period_label = " today" if limit_config["period"] == "daily" else ""

    return (
        f"You have reached the {limit_config['limit']}"
        f"{period_label} limit for {USAGE_LABELS[action_type]}. "
        "Your saved work is still available from the dashboard."
    )


def get_usage_summary(user_id):
    package_name = get_user_package(user_id) or "Starter"
    package_limits = PACKAGE_USAGE_LIMITS.get(
        package_name,
        PACKAGE_USAGE_LIMITS["Starter"]
    )
    counts = {
        action_type: get_usage_count(
            user_id,
            action_type,
            package_limits.get(action_type, {"period": "total"})["period"]
        )
        for action_type in USAGE_LABELS
    }

    return {
        "package_name": package_name,
        "chat_messages": counts["chat_message"],
        "chat_message_limit": package_limits["chat_message"]["limit"],
        "business_plans": counts["business_plan"],
        "business_plan_limit": package_limits["business_plan"]["limit"],
        "product_research_reports": counts["product_research"],
        "product_research_limit": (
            package_limits.get("product_research", {}).get("limit")
        ),
        "pricing_advice": counts["pricing_advice"],
        "pricing_advice_limit": package_limits.get("pricing_advice", {}).get("limit"),
        "supplier_guides": counts["supplier_guide"],
        "supplier_guide_limit": package_limits.get("supplier_guide", {}).get("limit"),
        "payment_guides": counts["payment_guide"],
        "payment_guide_limit": package_limits.get("payment_guide", {}).get("limit"),
        "domain_guides": counts["domain_guide"],
        "domain_guide_limit": package_limits.get("domain_guide", {}).get("limit"),
        "email_campaigns": counts["email_campaign"],
        "email_campaign_limit": package_limits.get("email_campaign", {}).get("limit"),
        "app_action_drafts": counts["app_action_draft"],
        "app_action_draft_limit": package_limits.get("app_action_draft", {}).get("limit"),
        "store_agent_tasks": counts["store_agent_task"],
        "store_agent_task_limit": (
            package_limits.get("store_agent_task", {}).get("limit")
        ),
        "shopify_actions": (
            counts["shopify_plan"]
            + counts["shopify_product"]
        ),
        "shopify_plans": counts["shopify_plan"],
        "shopify_plan_limit": package_limits["shopify_plan"]["limit"],
        "shopify_products": counts["shopify_product"],
        "shopify_product_limit": (
            package_limits.get("shopify_product", {}).get("limit")
        ),
        "canva_actions": (
            counts["canva_branding"]
            + counts["canva_design_brief"]
            + counts["canva_design"]
        ),
        "canva_branding": counts["canva_branding"],
        "canva_branding_limit": (
            package_limits.get("canva_branding", {}).get("limit")
        ),
        "canva_design_briefs": counts["canva_design_brief"],
        "canva_design_brief_limit": (
            package_limits.get("canva_design_brief", {}).get("limit")
        ),
        "canva_designs": counts["canva_design"],
        "canva_design_limit": package_limits.get("canva_design", {}).get("limit"),
        "launch_packages": counts["launch_package"],
        "launch_package_limit": (
            package_limits.get("launch_package", {}).get("limit")
        ),
        "pdf_exports": counts["pdf_export"],
        "pdf_export_limit": package_limits["pdf_export"]["limit"]
    }


def is_admin_user():
    admin_email = os.getenv("ADMIN_EMAIL", "").strip().lower()
    user_id = session.get("user_id")

    if not admin_email or not user_id:
        return False

    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT email
            FROM users
            WHERE id = ?
            LIMIT 1
        """),
        (user_id,)
    )

    user = cur.fetchone()
    conn.close()

    return bool(
        user
        and user[0]
        and user[0].strip().lower() == admin_email
    )


def get_admin_dashboard_data():
    conn = db()
    cur = conn.cursor()

    def count(query, params=()):
        cur.execute(sql(query), params)
        return cur.fetchone()[0]

    metrics = {
        "total_users": count("SELECT COUNT(*) FROM users"),
        "paid_users": count("""
            SELECT COUNT(*)
            FROM (
                SELECT user_id
                FROM payments
                WHERE status = ?
                UNION
                SELECT user_id
                FROM user_packages
                WHERE status = ?
            ) AS paid_accounts
        """, ("success", "active")),
        "total_payments": count("SELECT COUNT(*) FROM payments"),
        "business_plans": count("SELECT COUNT(*) FROM business_plans"),
        "shopify_connections": count("SELECT COUNT(*) FROM shopify_connections"),
        "canva_connections": count("SELECT COUNT(*) FROM canva_connections"),
        "shopify_products": count("SELECT COUNT(*) FROM shopify_products"),
        "canva_designs": count("SELECT COUNT(*) FROM canva_designs"),
        "product_research_reports": count("SELECT COUNT(*) FROM product_research"),
        "store_builds": count("SELECT COUNT(*) FROM ai_store_builds"),
        "store_agent_tasks": count("SELECT COUNT(*) FROM store_agent_tasks"),
        "support_tickets": count("SELECT COUNT(*) FROM support_tickets"),
        "open_support_tickets": count("SELECT COUNT(*) FROM support_tickets WHERE status = ?", ("open",)),
        "total_usage_logs": count("SELECT COUNT(*) FROM usage_logs")
    }

    cur.execute(sql("""
        SELECT package_name, COUNT(DISTINCT user_id)
        FROM user_packages
        WHERE status = ?
        GROUP BY package_name
    """), ("active",))
    package_counts = {
        package_name: package_count
        for package_name, package_count in cur.fetchall()
    }

    legacy_starter_users = count("""
        SELECT COUNT(DISTINCT payments.user_id)
        FROM payments
        WHERE payments.status = ?
        AND NOT EXISTS (
            SELECT 1
            FROM user_packages
            WHERE user_packages.user_id = payments.user_id
            AND user_packages.status = ?
        )
    """, ("success", "active"))

    metrics["starter_users"] = package_counts.get("Starter", 0) + legacy_starter_users
    metrics["pro_users"] = package_counts.get("Pro", 0)
    metrics["premium_build_users"] = package_counts.get("Premium Build", 0)
    metrics["free_users"] = max(metrics["total_users"] - metrics["paid_users"], 0)

    cur.execute(sql("""
        SELECT id, email
        FROM users
        ORDER BY id DESC
        LIMIT 10
    """))
    recent_users = cur.fetchall()

    cur.execute(sql("""
        SELECT
            payments.id,
            users.email,
            payments.provider,
            payments.amount,
            payments.status,
            payments.reference
        FROM payments
        LEFT JOIN users ON users.id = payments.user_id
        ORDER BY payments.id DESC
        LIMIT 10
    """))
    recent_payments = cur.fetchall()

    cur.execute(sql("""
        SELECT action_type, COUNT(*) AS usage_count
        FROM usage_logs
        GROUP BY action_type
        ORDER BY usage_count DESC, action_type ASC
        LIMIT 10
    """))
    most_used_actions = cur.fetchall()

    cur.execute(sql("""
        SELECT id, email, category, subject, status, created_at, message
        FROM support_tickets
        ORDER BY id DESC
        LIMIT 10
    """))
    recent_support_tickets = cur.fetchall()

    conn.close()

    return {
        "metrics": metrics,
        "recent_users": recent_users,
        "recent_payments": recent_payments,
        "most_used_actions": most_used_actions,
        "recent_support_tickets": recent_support_tickets
    }


# -----------------------------
# WORKFLOW HELPERS
# -----------------------------

WORKFLOW_STEPS = {
    1: {
        "name": "Business Idea",
        "question": "Describe your business idea. What problem does it solve, who is it for, and what makes it valuable?"
    },
    2: {
        "name": "Brand Name",
        "question": "Describe the type of brand name, slogan, personality, and identity you want for your business."
    },
    3: {
        "name": "Target Market",
        "question": "Describe your ideal customer. Include their age, location, interests, problems, buying behavior, and where you can reach them."
    },
    4: {
        "name": "Products / Services",
        "question": "Describe the products or services your business will offer, including packages, features, benefits, and value."
    },
    5: {
        "name": "Pricing",
        "question": "Describe your pricing ideas, package options, target profit margins, discounts, and launch offers."
    },
    6: {
        "name": "Marketing Plan",
        "question": "Describe your marketing goals, social media ideas, launch strategy, customer acquisition channels, and content ideas."
    },
    7: {
        "name": "Shopify Setup Plan",
        "question": "Describe how you want your Shopify store to work, including pages, products, collections, policies, and checkout flow."
    },
    8: {
        "name": "Canva Branding Plan",
        "question": "Describe your visual branding needs, including logo direction, colors, fonts, social templates, banners, and product graphics."
    },
    9: {
        "name": "Launch Checklist",
        "question": "Describe what still needs to be completed before launch, including branding, products, website, payments, marketing, and customer support."
    }
}


def mark_workflow_step_complete(user_id, step_number, step_name):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT id
            FROM workflow_progress
            WHERE user_id = ?
            AND step_number = ?
        """),
        (user_id, step_number)
    )

    existing = cur.fetchone()

    if not existing:
        cur.execute(
            sql("""
                INSERT INTO workflow_progress
                (user_id, step_number, step_name, status)
                VALUES (?, ?, ?, ?)
            """),
            (user_id, step_number, step_name, "completed")
        )

    conn.commit()
    conn.close()


def get_completed_steps(user_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT step_number
            FROM workflow_progress
            WHERE user_id = ?
            AND status = ?
        """),
        (user_id, "completed")
    )

    rows = cur.fetchall()
    conn.close()

    return [row[0] for row in rows]


def save_workflow_answer(user_id, step_number, step_name, answer):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT id
            FROM workflow_answers
            WHERE user_id = ?
            AND step_number = ?
        """),
        (user_id, step_number)
    )

    existing = cur.fetchone()

    if existing:
        cur.execute(
            sql("""
                UPDATE workflow_answers
                SET answer = ?
                WHERE user_id = ?
                AND step_number = ?
            """),
            (answer, user_id, step_number)
        )
    else:
        cur.execute(
            sql("""
                INSERT INTO workflow_answers
                (user_id, step_number, step_name, answer)
                VALUES (?, ?, ?, ?)
            """),
            (user_id, step_number, step_name, answer)
        )

    conn.commit()
    conn.close()


def get_workflow_answer(user_id, step_number):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT answer
            FROM workflow_answers
            WHERE user_id = ?
            AND step_number = ?
            LIMIT 1
        """),
        (user_id, step_number)
    )

    row = cur.fetchone()
    conn.close()

    if row:
        return row[0]

    return ""


def get_all_workflow_answers(user_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT step_number, step_name, answer
            FROM workflow_answers
            WHERE user_id = ?
            ORDER BY step_number ASC
        """),
        (user_id,)
    )

    rows = cur.fetchall()
    conn.close()

    return rows


def get_nonempty_workflow_answers(user_id):
    return [
        answer
        for answer in get_all_workflow_answers(user_id)
        if answer[2] and answer[2].strip()
    ]


BUILD_APPROVAL_ACTIONS = {
    "shopify_products": "/create_shopify_products",
    "shopify_store": "/build_shopify_store_draft",
    "canva_designs": "/create_canva_designs",
    "full_build": "/build_shopify_store_draft",
    "ai_store": "/generate_full_store",
    "product_research_products": "/create_product_research_shopify_products"
}


def grant_build_approval(action_type):
    if action_type == "full_build":
        approved_actions = ["shopify_store", "canva_designs"]
        session["continue_full_build"] = True
    else:
        approved_actions = [action_type]

    session["build_approvals"] = approved_actions


def consume_build_approval(action_type):
    approved_actions = session.get("build_approvals", [])

    if action_type not in approved_actions:
        return False

    approved_actions.remove(action_type)

    if approved_actions:
        session["build_approvals"] = approved_actions
    else:
        session.pop("build_approvals", None)

    return True


def revoke_build_approval(action_type):
    approved_actions = session.get("build_approvals", [])

    if action_type in approved_actions:
        approved_actions.remove(action_type)

    if approved_actions:
        session["build_approvals"] = approved_actions
    else:
        session.pop("build_approvals", None)


def build_approval_redirect(action_type):
    return redirect(
        f"/build_approval?action={urllib.parse.quote(action_type)}"
    )


def save_business_plan(user_id, title, content):
    conn = db()
    cur = conn.cursor()

    if using_postgres():
        cur.execute(
            sql("""
            INSERT INTO business_plans (user_id, title, content)
            VALUES (%s, %s, %s)
            RETURNING id
            """),
            (user_id, title, content)
        )
        plan_id = cur.fetchone()[0]
    else:
        cur.execute(
            sql("""
            INSERT INTO business_plans (user_id, title, content)
            VALUES (?, ?, ?)
            """),
            (user_id, title, content)
        )
        plan_id = cur.lastrowid

    conn.commit()
    conn.close()

    return plan_id


def get_business_plans(user_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT id, title, content
            FROM business_plans
            WHERE user_id = ?
            ORDER BY id DESC
        """),
        (user_id,)
    )

    plans = cur.fetchall()
    conn.close()

    return plans


def get_business_plan(user_id, plan_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT id, title, content
            FROM business_plans
            WHERE id = ?
            AND user_id = ?
            LIMIT 1
        """),
        (plan_id, user_id)
    )

    plan = cur.fetchone()
    conn.close()

    return plan


def save_ai_store_build(user_id, title, store_name, content, status):
    conn = db()
    cur = conn.cursor()

    if using_postgres():
        cur.execute(
            sql("""
            INSERT INTO ai_store_builds (
                user_id,
                title,
                store_name,
                content,
                status
            )
            VALUES (%s, %s, %s, %s, %s)
            RETURNING id
            """),
            (user_id, title, store_name, content, status)
        )
        build_id = cur.fetchone()[0]
    else:
        cur.execute(
            sql("""
            INSERT INTO ai_store_builds (
                user_id,
                title,
                store_name,
                content,
                status
            )
            VALUES (?, ?, ?, ?, ?)
            """),
            (user_id, title, store_name, content, status)
        )
        build_id = cur.lastrowid

    conn.commit()
    conn.close()

    return build_id


def get_ai_store_builds(user_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT id, title, store_name, content, status, created_at
            FROM ai_store_builds
            WHERE user_id = ?
            ORDER BY id DESC
        """),
        (user_id,)
    )

    builds = cur.fetchall()
    conn.close()

    return builds


def get_ai_store_build(user_id, build_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT id, title, store_name, content, status, created_at
            FROM ai_store_builds
            WHERE id = ?
            AND user_id = ?
            LIMIT 1
        """),
        (build_id, user_id)
    )

    build = cur.fetchone()
    conn.close()

    return build


def save_product_research(
    user_id,
    title,
    business_type,
    target_market,
    sourcing_preference,
    budget,
    country,
    content
):
    conn = db()
    cur = conn.cursor()

    if using_postgres():
        cur.execute(
            sql("""
                INSERT INTO product_research (
                    user_id,
                    title,
                    business_type,
                    target_market,
                    sourcing_preference,
                    budget,
                    country,
                    content
                )
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                RETURNING id
            """),
            (
                user_id,
                title,
                business_type,
                target_market,
                sourcing_preference,
                budget,
                country,
                content
            )
        )
        research_id = cur.fetchone()[0]
    else:
        cur.execute(
            sql("""
                INSERT INTO product_research (
                    user_id,
                    title,
                    business_type,
                    target_market,
                    sourcing_preference,
                    budget,
                    country,
                    content
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """),
            (
                user_id,
                title,
                business_type,
                target_market,
                sourcing_preference,
                budget,
                country,
                content
            )
        )
        research_id = cur.lastrowid

    conn.commit()
    conn.close()

    return research_id


def get_product_research_list(user_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT id, title, business_type, target_market,
                   sourcing_preference, budget, country, content, created_at
            FROM product_research
            WHERE user_id = ?
            ORDER BY id DESC
        """),
        (user_id,)
    )

    research = cur.fetchall()
    conn.close()

    return research


def get_product_research(user_id, research_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT id, title, business_type, target_market,
                   sourcing_preference, budget, country, content, created_at
            FROM product_research
            WHERE id = ?
            AND user_id = ?
            LIMIT 1
        """),
        (research_id, user_id)
    )

    research = cur.fetchone()
    conn.close()

    return research


def parse_product_research_content(content):
    try:
        return json.loads(content)
    except (TypeError, ValueError, json.JSONDecodeError):
        return None


STORE_AGENT_TASK_DEFINITIONS = {
    "theme_recommendation": {
        "section": "Store Theme",
        "title": "Store Theme Recommendation",
        "mode": "Advice Mode",
        "explanation": "Get theme suggestions and setup guidance without installing or changing a live theme.",
        "apply_supported": False
    },
    "homepage_design": {
        "section": "Homepage Design",
        "title": "Homepage Design Draft",
        "mode": "Draft Mode",
        "explanation": "Generate homepage sections, CTA copy, trust badges, FAQ, footer copy, and a Canva banner brief.",
        "apply_supported": True
    },
    "theme_sections": {
        "section": "Homepage Design",
        "title": "Theme Section Draft",
        "mode": "Draft Mode",
        "explanation": "Generate section-by-section Shopify theme guidance without editing live theme files.",
        "apply_supported": True
    },
    "products": {
        "section": "Shopify Assets",
        "title": "Shopify Draft Products",
        "mode": "Action Mode",
        "explanation": "Generate product drafts and apply them as draft Shopify products after approval.",
        "apply_supported": True
    },
    "product_sourcing": {
        "section": "Product Sourcing",
        "title": "Product Sourcing Research",
        "mode": "Advice Mode",
        "explanation": "Find product ideas, supplier search paths, competitor examples, and the first products to add to Shopify.",
        "apply_supported": False
    },
    "supplier_finder": {
        "section": "Supplier Finder",
        "title": "Supplier and Platform Recommendation",
        "mode": "Advice Mode",
        "explanation": "Compare Amazon-style research, Alibaba, AliExpress, local suppliers, print-on-demand, digital platforms, wholesalers, and other options based on budget.",
        "apply_supported": False
    },
    "pricing_advisor": {
        "section": "Pricing Advisor",
        "title": "Pricing Strategy Advice",
        "mode": "Advice Mode",
        "explanation": "Generate suggested price ranges, margin guidance, break-even notes, competitor pricing advice, and first-price tests.",
        "apply_supported": False
    },
    "pages": {
        "section": "Shopify Pages",
        "title": "Shopify Page Drafts",
        "mode": "Action Mode",
        "explanation": "Generate About, Contact, FAQ, Shipping, and Refund page drafts for unpublished Shopify pages.",
        "apply_supported": True
    },
    "shipping_zones": {
        "section": "Shipping Zones",
        "title": "Shipping Zone Setup Plan",
        "mode": "Advice Mode",
        "explanation": "Create shipping recommendations and manual Shopify setup steps. No shipping settings are changed automatically.",
        "apply_supported": True
    },
    "payments_setup": {
        "section": "Payment Setup Guide",
        "title": "Payments Setup Checklist",
        "mode": "Advice Mode",
        "explanation": "Create guidance for PayPal, Paystack, Shopify payments, EFT, card payments, cash on delivery, and other options. Banking information must be entered directly with the provider.",
        "apply_supported": True
    },
    "domain_setup": {
        "section": "Domain Setup",
        "title": "Domain Setup Guide",
        "mode": "Advice Mode",
        "explanation": "Suggest domains, providers, costs, and connection steps. Domains are never bought automatically.",
        "apply_supported": True
    },
    "canva_branding": {
        "section": "Canva Branding",
        "title": "Canva Branding Brief",
        "mode": "Draft Mode",
        "explanation": "Generate logo, brand kit, banner, product mockup, and visual identity briefs.",
        "apply_supported": True
    },
    "canva_marketing": {
        "section": "Canva Marketing Assets",
        "title": "Canva Marketing Asset Brief",
        "mode": "Draft Mode",
        "explanation": "Generate social post, story, ad, launch graphic, and product promotion briefs.",
        "apply_supported": True
    },
    "email_marketing": {
        "section": "Email Marketing",
        "title": "Email Marketing Campaign Draft",
        "mode": "Draft Mode",
        "explanation": "Create email strategy, subject lines, body copy, calls to action, platform guidance, and a sending checklist. Emails are never sent automatically.",
        "apply_supported": False
    },
    "launch_checklist": {
        "section": "Launch Checklist",
        "title": "Store Launch Checklist",
        "mode": "Advice Mode",
        "explanation": "Create a review checklist that guides the user through launch without publishing automatically.",
        "apply_supported": True
    }
}


STORE_AGENT_SECTION_ORDER = [
    "theme_recommendation",
    "homepage_design",
    "products",
    "product_sourcing",
    "supplier_finder",
    "pricing_advisor",
    "pages",
    "shipping_zones",
    "payments_setup",
    "domain_setup",
    "canva_branding",
    "canva_marketing",
    "email_marketing",
    "launch_checklist"
]


def normalize_store_agent_task_type(task_type):
    return task_type


def create_store_agent_task(user_id, task_type, title, draft_content, status):
    task_type = normalize_store_agent_task_type(task_type)
    conn = db()
    cur = conn.cursor()

    if using_postgres():
        cur.execute(
            sql("""
                INSERT INTO store_agent_tasks (
                    user_id,
                    task_type,
                    title,
                    draft_content,
                    status
                )
                VALUES (%s, %s, %s, %s, %s)
                RETURNING id
            """),
            (user_id, task_type, title, draft_content, status)
        )
        task_id = cur.fetchone()[0]
    else:
        cur.execute(
            sql("""
                INSERT INTO store_agent_tasks (
                    user_id,
                    task_type,
                    title,
                    draft_content,
                    status
                )
                VALUES (?, ?, ?, ?, ?)
            """),
            (user_id, task_type, title, draft_content, status)
        )
        task_id = cur.lastrowid

    conn.commit()
    conn.close()

    return task_id


def get_store_agent_tasks(user_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT id, user_id, task_type, title, draft_content, status,
                   approved, applied, created_at, updated_at
            FROM store_agent_tasks
            WHERE user_id = ?
            ORDER BY id DESC
        """),
        (user_id,)
    )

    tasks = cur.fetchall()
    conn.close()

    return tasks


def get_store_agent_task(user_id, task_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT id, user_id, task_type, title, draft_content, status,
                   approved, applied, created_at, updated_at
            FROM store_agent_tasks
            WHERE id = ?
            AND user_id = ?
            LIMIT 1
        """),
        (task_id, user_id)
    )

    task = cur.fetchone()
    conn.close()

    return task


def update_store_agent_task_status(task_id, user_id, status, approved, applied):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            UPDATE store_agent_tasks
            SET status = ?,
                approved = ?,
                applied = ?,
                updated_at = CURRENT_TIMESTAMP
            WHERE id = ?
            AND user_id = ?
        """),
        (status, int(bool(approved)), int(bool(applied)), task_id, user_id)
    )

    conn.commit()
    conn.close()


def update_store_agent_task_content(task_id, user_id, draft_content):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            UPDATE store_agent_tasks
            SET draft_content = ?,
                status = ?,
                approved = ?,
                applied = ?,
                updated_at = CURRENT_TIMESTAMP
            WHERE id = ?
            AND user_id = ?
        """),
        (draft_content, "Needs Approval", 0, 0, task_id, user_id)
    )

    conn.commit()
    conn.close()


def delete_store_agent_task_record(task_id, user_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            DELETE FROM store_agent_tasks
            WHERE id = ?
            AND user_id = ?
        """),
        (task_id, user_id)
    )

    conn.commit()
    conn.close()


def get_latest_store_agent_tasks_by_type(user_id):
    latest = {}

    for task in get_store_agent_tasks(user_id):
        task_type = task[2]

        if task_type not in latest:
            latest[task_type] = task

    return latest


def save_shopify_plan(user_id, title, content):
    conn = db()
    cur = conn.cursor()

    if using_postgres():
        cur.execute(
            sql("""
            INSERT INTO shopify_plans (user_id, title, content)
            VALUES (%s, %s, %s)
            RETURNING id
            """),
            (user_id, title, content)
        )
        plan_id = cur.fetchone()[0]
    else:
        cur.execute(
            sql("""
            INSERT INTO shopify_plans (user_id, title, content)
            VALUES (?, ?, ?)
            """),
            (user_id, title, content)
        )
        plan_id = cur.lastrowid

    conn.commit()
    conn.close()

    return plan_id


def get_shopify_plans(user_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT id, title, content
            FROM shopify_plans
            WHERE user_id = ?
            ORDER BY id DESC
        """),
        (user_id,)
    )

    plans = cur.fetchall()
    conn.close()

    return plans


def get_shopify_plan(user_id, plan_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT id, title, content
            FROM shopify_plans
            WHERE id = ?
            AND user_id = ?
            LIMIT 1
        """),
        (plan_id, user_id)
    )

    plan = cur.fetchone()
    conn.close()

    return plan


def save_canva_branding_package(user_id, title, content):
    conn = db()
    cur = conn.cursor()

    if using_postgres():
        cur.execute(
            sql("""
            INSERT INTO canva_branding_packages (user_id, title, content)
            VALUES (%s, %s, %s)
            RETURNING id
            """),
            (user_id, title, content)
        )
        package_id = cur.fetchone()[0]
    else:
        cur.execute(
            sql("""
            INSERT INTO canva_branding_packages (user_id, title, content)
            VALUES (?, ?, ?)
            """),
            (user_id, title, content)
        )
        package_id = cur.lastrowid

    conn.commit()
    conn.close()

    return package_id


def get_canva_branding_packages(user_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT id, title, content
            FROM canva_branding_packages
            WHERE user_id = ?
            ORDER BY id DESC
        """),
        (user_id,)
    )

    packages = cur.fetchall()
    conn.close()

    return packages


def get_canva_branding_package(user_id, package_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT id, title, content
            FROM canva_branding_packages
            WHERE id = ?
            AND user_id = ?
            LIMIT 1
        """),
        (package_id, user_id)
    )

    package = cur.fetchone()
    conn.close()

    return package


def save_canva_design_brief(user_id, title, content):
    conn = db()
    cur = conn.cursor()

    if using_postgres():
        cur.execute(
            sql("""
            INSERT INTO canva_design_briefs (user_id, title, content)
            VALUES (%s, %s, %s)
            RETURNING id
            """),
            (user_id, title, content)
        )
        brief_id = cur.fetchone()[0]
    else:
        cur.execute(
            sql("""
            INSERT INTO canva_design_briefs (user_id, title, content)
            VALUES (?, ?, ?)
            """),
            (user_id, title, content)
        )
        brief_id = cur.lastrowid

    conn.commit()
    conn.close()

    return brief_id


def get_canva_design_briefs(user_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT id, title, content
            FROM canva_design_briefs
            WHERE user_id = ?
            ORDER BY id DESC
        """),
        (user_id,)
    )

    briefs = cur.fetchall()
    conn.close()

    return briefs


def get_canva_design_brief(user_id, brief_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT id, title, content
            FROM canva_design_briefs
            WHERE id = ?
            AND user_id = ?
            LIMIT 1
        """),
        (brief_id, user_id)
    )

    brief = cur.fetchone()
    conn.close()

    return brief


def get_latest_canva_design_brief(user_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT id, title, content
            FROM canva_design_briefs
            WHERE user_id = ?
            ORDER BY id DESC
            LIMIT 1
        """),
        (user_id,)
    )

    brief = cur.fetchone()
    conn.close()

    return brief


def save_canva_design(
    user_id,
    title,
    canva_design_id,
    edit_url,
    view_url,
    status
):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            INSERT INTO canva_designs (
                user_id,
                title,
                canva_design_id,
                edit_url,
                view_url,
                status
            )
            VALUES (?, ?, ?, ?, ?, ?)
        """),
        (
            user_id,
            title,
            canva_design_id,
            edit_url,
            view_url,
            status
        )
    )

    conn.commit()
    conn.close()


def get_canva_designs(user_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT title, canva_design_id, edit_url, view_url, status
            FROM canva_designs
            WHERE user_id = ?
            ORDER BY id DESC
        """),
        (user_id,)
    )

    designs = cur.fetchall()
    conn.close()

    return designs


def save_build_quote(
    user_id,
    title,
    shopify_plan,
    canva_plan,
    estimated_shopify_cost,
    estimated_canva_cost,
    service_fee,
    total_estimate,
    content
):
    conn = db()
    cur = conn.cursor()

    if using_postgres():
        cur.execute(
            sql("""
            INSERT INTO build_quotes (
                user_id,
                title,
                shopify_plan,
                canva_plan,
                estimated_shopify_cost,
                estimated_canva_cost,
                service_fee,
                total_estimate,
                content
            )
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
            RETURNING id
            """),
            (
                user_id,
                title,
                shopify_plan,
                canva_plan,
                estimated_shopify_cost,
                estimated_canva_cost,
                service_fee,
                total_estimate,
                content
            )
        )
        quote_id = cur.fetchone()[0]
    else:
        cur.execute(
            sql("""
            INSERT INTO build_quotes (
                user_id,
                title,
                shopify_plan,
                canva_plan,
                estimated_shopify_cost,
                estimated_canva_cost,
                service_fee,
                total_estimate,
                content
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """),
            (
                user_id,
                title,
                shopify_plan,
                canva_plan,
                estimated_shopify_cost,
                estimated_canva_cost,
                service_fee,
                total_estimate,
                content
            )
        )
        quote_id = cur.lastrowid

    conn.commit()
    conn.close()

    return quote_id


def get_build_quotes(user_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT id, title, content
            FROM build_quotes
            WHERE user_id = ?
            ORDER BY id DESC
        """),
        (user_id,)
    )

    quotes = cur.fetchall()
    conn.close()

    return quotes


def get_build_quote(user_id, quote_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT
                id,
                title,
                shopify_plan,
                canva_plan,
                estimated_shopify_cost,
                estimated_canva_cost,
                service_fee,
                total_estimate,
                content
            FROM build_quotes
            WHERE id = ?
            AND user_id = ?
            LIMIT 1
        """),
        (quote_id, user_id)
    )

    quote = cur.fetchone()
    conn.close()

    return quote


def save_shopify_connection(user_id, shop_domain, access_token, status):
    encrypted_access_token = encrypt_token(access_token)
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT id
            FROM shopify_connections
            WHERE user_id = ?
            LIMIT 1
        """),
        (user_id,)
    )

    connection = cur.fetchone()

    if connection:
        cur.execute(
            sql("""
                UPDATE shopify_connections
                SET shop_domain = ?,
                    access_token = ?,
                    status = ?
                WHERE user_id = ?
            """),
            (shop_domain, encrypted_access_token, status, user_id)
        )
    else:
        cur.execute(
            sql("""
                INSERT INTO shopify_connections (
                    user_id,
                    shop_domain,
                    access_token,
                    status
                )
                VALUES (?, ?, ?, ?)
            """),
            (user_id, shop_domain, encrypted_access_token, status)
        )

    conn.commit()
    conn.close()


def get_shopify_connection(user_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT id, shop_domain, access_token, status
            FROM shopify_connections
            WHERE user_id = ?
            LIMIT 1
        """),
        (user_id,)
    )

    connection = cur.fetchone()
    conn.close()

    return connection


def get_canva_connection(user_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT
                id,
                user_id,
                status,
                connected_email,
                access_token,
                refresh_token
            FROM canva_connections
            WHERE user_id = ?
            LIMIT 1
        """),
        (user_id,)
    )

    connection = cur.fetchone()
    conn.close()

    return connection


def save_canva_connection(
    user_id,
    status,
    connected_email,
    access_token,
    refresh_token
):
    encrypted_access_token = encrypt_token(access_token)
    encrypted_refresh_token = encrypt_token(refresh_token)
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT id
            FROM canva_connections
            WHERE user_id = ?
            LIMIT 1
        """),
        (user_id,)
    )

    connection = cur.fetchone()

    if connection:
        cur.execute(
            sql("""
                UPDATE canva_connections
                SET status = ?,
                    connected_email = ?,
                    access_token = ?,
                    refresh_token = ?
                WHERE user_id = ?
            """),
            (
                status,
                connected_email,
                encrypted_access_token,
                encrypted_refresh_token,
                user_id
            )
        )
    else:
        cur.execute(
            sql("""
                INSERT INTO canva_connections (
                    user_id,
                    status,
                    connected_email,
                    access_token,
                    refresh_token
                )
                VALUES (?, ?, ?, ?, ?)
            """),
            (
                user_id,
                status,
                connected_email,
                encrypted_access_token,
                encrypted_refresh_token
            )
        )

    conn.commit()
    conn.close()


def create_canva_design(user_id, title, width=1080, height=1080):
    connection = get_canva_connection(user_id)

    if not connection or connection[2] != "connected":
        return None, None, None, None, "Connect Canva before creating a design draft."

    access_token = decrypt_token(connection[4])

    if not access_token:
        return None, None, None, None, "Your Canva connection is missing an access token. Reconnect Canva and try again."

    try:
        response = requests.post(
            CANVA_DESIGNS_URL,
            headers={
                "Authorization": f"Bearer {access_token}",
                "Content-Type": "application/json"
            },
            json={
                "type": "type_and_asset",
                "design_type": {
                    "type": "custom",
                    "width": width,
                    "height": height
                },
                "title": title
            },
            timeout=10
        )
        result = response.json()
    except requests.RequestException:
        return None, None, None, None, "Could not reach Canva. Please try again."
    except ValueError:
        return None, None, None, None, "Canva returned an unexpected response. Please try again."

    design = result.get("design") or result.get("data", {}).get("design") or {}
    canva_design_id = design.get("id") or design.get("design_id")

    if response.status_code != 200 or not canva_design_id:
        return None, None, None, None, "Canva could not create the design draft. Reconnect Canva or try again."

    urls = design.get("urls", {})

    return (
        canva_design_id,
        urls.get("edit_url") or design.get("edit_url", ""),
        urls.get("view_url") or design.get("view_url", ""),
        "draft",
        None
    )


def save_canva_oauth_session(user_id, state, code_verifier):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            DELETE FROM canva_oauth_sessions
            WHERE user_id = ?
        """),
        (user_id,)
    )

    cur.execute(
        sql("""
            INSERT INTO canva_oauth_sessions (
                user_id,
                state,
                code_verifier
            )
            VALUES (?, ?, ?)
        """),
        (user_id, state, code_verifier)
    )

    conn.commit()
    conn.close()


def pop_canva_oauth_verifier(user_id, state):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT code_verifier
            FROM canva_oauth_sessions
            WHERE user_id = ?
            AND state = ?
            LIMIT 1
        """),
        (user_id, state)
    )

    row = cur.fetchone()

    cur.execute(
        sql("""
            DELETE FROM canva_oauth_sessions
            WHERE user_id = ?
        """),
        (user_id,)
    )

    conn.commit()
    conn.close()

    return row[0] if row else None


def normalize_shopify_domain(shop_domain):
    domain = shop_domain.strip().lower()
    domain = re.sub(r"^https?://", "", domain)
    domain = domain.split("/", 1)[0]

    if re.fullmatch(r"[a-z0-9][a-z0-9-]*\.myshopify\.com", domain):
        return domain

    return None


def test_shopify_connection(shop_domain, access_token):
    endpoint = (
        f"https://{shop_domain}/admin/api/"
        f"{SHOPIFY_ADMIN_API_VERSION}/graphql.json"
    )
    query = """
    {
        shop {
            name
            myshopifyDomain
        }
    }
    """

    try:
        response = requests.post(
            endpoint,
            headers={
                "Content-Type": "application/json",
                "X-Shopify-Access-Token": access_token
            },
            json={"query": query},
            timeout=10
        )
        result = response.json()
    except requests.RequestException:
        return False, "Could not reach Shopify. Check the store domain and try again."
    except ValueError:
        return False, "Shopify returned an unexpected response. Check your connection details."

    if response.status_code != 200 or result.get("errors"):
        return False, "Shopify connection failed. Check the store domain and Admin API access token."

    shop = result.get("data", {}).get("shop")

    if not shop:
        return False, "Shopify connection failed. The shop details could not be retrieved."

    return True, f"Connected to {shop['name']} ({shop['myshopifyDomain']})."


def save_shopify_product(
    user_id,
    title,
    description,
    shopify_product_id,
    status
):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            INSERT INTO shopify_products (
                user_id,
                title,
                description,
                shopify_product_id,
                status
            )
            VALUES (?, ?, ?, ?, ?)
        """),
        (user_id, title, description, shopify_product_id, status)
    )

    conn.commit()
    conn.close()


def get_shopify_products(user_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT title, description, shopify_product_id, status
            FROM shopify_products
            WHERE user_id = ?
            ORDER BY id DESC
        """),
        (user_id,)
    )

    products = cur.fetchall()
    conn.close()

    return products


def save_shopify_collection(
    user_id,
    title,
    description,
    shopify_collection_id,
    status
):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            INSERT INTO shopify_collections (
                user_id,
                title,
                description,
                shopify_collection_id,
                status
            )
            VALUES (?, ?, ?, ?, ?)
        """),
        (user_id, title, description, shopify_collection_id, status)
    )

    conn.commit()
    conn.close()


def get_shopify_collections(user_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT title, description, shopify_collection_id, status
            FROM shopify_collections
            WHERE user_id = ?
            ORDER BY id DESC
        """),
        (user_id,)
    )

    collections = cur.fetchall()
    conn.close()

    return collections


def save_shopify_page(
    user_id,
    title,
    page_type,
    content,
    shopify_page_id,
    status
):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            INSERT INTO shopify_pages (
                user_id,
                title,
                page_type,
                content,
                shopify_page_id,
                status
            )
            VALUES (?, ?, ?, ?, ?, ?)
        """),
        (user_id, title, page_type, content, shopify_page_id, status)
    )

    conn.commit()
    conn.close()


def get_shopify_pages(user_id):
    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT title, page_type, content, shopify_page_id, status
            FROM shopify_pages
            WHERE user_id = ?
            ORDER BY id DESC
        """),
        (user_id,)
    )

    pages = cur.fetchall()
    conn.close()

    return pages


def send_shopify_graphql(user_id, query, variables):
    connection = get_shopify_connection(user_id)

    if not connection or connection[3] != "connected":
        return None, "Connect and test your Shopify store before building store drafts."

    shop_domain = connection[1]
    # Decrypt only at the API boundary; database rows remain encrypted.
    access_token = decrypt_token(connection[2])
    endpoint = (
        f"https://{shop_domain}/admin/api/"
        f"{SHOPIFY_ADMIN_API_VERSION}/graphql.json"
    )

    try:
        response = requests.post(
            endpoint,
            headers={
                "Content-Type": "application/json",
                "X-Shopify-Access-Token": access_token
            },
            json={
                "query": query,
                "variables": variables
            },
            timeout=10
        )
        result = response.json()
    except requests.RequestException:
        return None, "Could not reach Shopify. Check your connection and try again."
    except ValueError:
        return None, "Shopify returned an unexpected response."

    if response.status_code != 200 or result.get("errors"):
        return None, "Shopify could not create the requested draft item."

    return result, None


def create_shopify_collection(user_id, title, description):
    mutation = """
    mutation CreateCollection($input: CollectionInput!) {
        collectionCreate(input: $input) {
            collection {
                id
            }
            userErrors {
                field
                message
            }
        }
    }
    """
    result, error = send_shopify_graphql(
        user_id,
        mutation,
        {
            "input": {
                "title": title,
                "descriptionHtml": (
                    f"<p>{escape(description).replace(chr(10), '<br>')}</p>"
                )
            }
        }
    )

    if error:
        return None, None, error

    collection_create = result.get("data", {}).get("collectionCreate") or {}

    if collection_create.get("userErrors"):
        return None, None, "Shopify could not create the unpublished collection."

    collection = collection_create.get("collection")

    if not collection:
        return None, None, "Shopify did not return the created collection."

    return collection["id"], "unpublished", None


def create_shopify_page(user_id, title, content):
    mutation = """
    mutation CreatePage($page: PageCreateInput!) {
        pageCreate(page: $page) {
            page {
                id
            }
            userErrors {
                field
                message
            }
        }
    }
    """
    result, error = send_shopify_graphql(
        user_id,
        mutation,
        {
            "page": {
                "title": title,
                "body": f"<p>{escape(content).replace(chr(10), '<br>')}</p>",
                "isPublished": False
            }
        }
    )

    if error:
        return None, None, error

    page_create = result.get("data", {}).get("pageCreate") or {}

    if page_create.get("userErrors"):
        return None, None, "Shopify could not create the unpublished page."

    page = page_create.get("page")

    if not page:
        return None, None, "Shopify did not return the created page."

    return page["id"], "draft", None


def create_shopify_product(
    user_id,
    title,
    description,
    suggested_price=None,
    category=None
):
    connection = get_shopify_connection(user_id)

    if not connection or connection[3] != "connected":
        return None, None, "Connect and test your Shopify store before creating products."

    shop_domain = connection[1]
    # Decrypt only at the API boundary; database rows remain encrypted.
    access_token = decrypt_token(connection[2])
    endpoint = (
        f"https://{shop_domain}/admin/api/"
        f"{SHOPIFY_ADMIN_API_VERSION}/graphql.json"
    )
    mutation = """
    mutation CreateProduct($product: ProductCreateInput!) {
        productCreate(product: $product) {
            product {
                id
                status
            }
            userErrors {
                field
                message
            }
        }
    }
    """
    description_lines = [description]

    if suggested_price:
        description_lines.append(f"Suggested price: {suggested_price}")

    if category:
        description_lines.append(f"Collection / category: {category}")

    full_description = "\n\n".join(description_lines)
    product_input = {
        "title": title,
        "descriptionHtml": f"<p>{escape(full_description).replace(chr(10), '<br>')}</p>",
        "vendor": "BusinessBuilder AI",
        "status": "DRAFT"
    }

    if category:
        product_input["productType"] = category

    def send_product_create(input_data):
        try:
            response = requests.post(
                endpoint,
                headers={
                    "Content-Type": "application/json",
                    "X-Shopify-Access-Token": access_token
                },
                json={
                    "query": mutation,
                    "variables": {"product": input_data}
                },
                timeout=10
            )
            return response.status_code, response.json()
        except requests.RequestException:
            return None, None
        except ValueError:
            return None, None

    status_code, result = send_product_create(product_input)

    if not result:
        return None, None, "Could not reach Shopify. Try again after checking your connection."

    graphql_errors = result.get("errors", [])
    product_create = result.get("data", {}).get("productCreate") or {}
    user_errors = product_create.get("userErrors", [])
    all_errors = graphql_errors + user_errors

    if all_errors and any("status" in str(error).lower() for error in all_errors):
        product_input.pop("status", None)
        status_code, result = send_product_create(product_input)

        if not result:
            return None, None, "Could not reach Shopify. Try again after checking your connection."

        graphql_errors = result.get("errors", [])
        product_create = result.get("data", {}).get("productCreate") or {}
        user_errors = product_create.get("userErrors", [])

    if status_code != 200 or graphql_errors or user_errors:
        return None, None, "Shopify could not create the draft product. Check your Admin API permissions."

    product = product_create.get("product")

    if not product:
        return None, None, "Shopify did not return the created draft product."

    return product["id"], product.get("status", "DRAFT"), None


def create_business_plan_pdf(title, content):
    pdf_buffer = BytesIO()
    document = SimpleDocTemplate(
        pdf_buffer,
        pagesize=A4,
        rightMargin=20 * mm,
        leftMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "BusinessPlanTitle",
        parent=styles["Title"],
        alignment=TA_CENTER,
        textColor=colors.HexColor("#172554"),
        fontSize=22,
        leading=28,
        spaceAfter=8
    )
    brand_style = ParagraphStyle(
        "BusinessPlanBrand",
        parent=styles["Normal"],
        alignment=TA_CENTER,
        textColor=colors.HexColor("#2563eb"),
        fontSize=10,
        leading=14,
        spaceAfter=14
    )
    heading_style = ParagraphStyle(
        "BusinessPlanHeading",
        parent=styles["Heading2"],
        textColor=colors.HexColor("#1e3a8a"),
        fontSize=14,
        leading=18,
        spaceBefore=10,
        spaceAfter=5
    )
    body_style = ParagraphStyle(
        "BusinessPlanBody",
        parent=styles["BodyText"],
        textColor=colors.HexColor("#334155"),
        fontSize=10.5,
        leading=15,
        spaceAfter=6
    )

    story = [
        Paragraph("BusinessBuilder AI", brand_style),
        Paragraph(escape(title), title_style),
        Spacer(1, 5 * mm)
    ]

    for line in content.splitlines():
        text = line.strip()

        if not text:
            story.append(Spacer(1, 2.5 * mm))
            continue

        clean_text = escape(text)
        is_heading = (
            text.startswith("#")
            or (
                len(text) <= 90
                and (
                    text.endswith(":")
                    or text[0].isdigit() and "." in text[:4]
                )
            )
        )

        if text.startswith("#"):
            clean_text = escape(text.lstrip("#").strip())

        story.append(
            Paragraph(
                clean_text,
                heading_style if is_heading else body_style
            )
        )

    document.build(story)
    pdf_buffer.seek(0)

    return pdf_buffer


def create_launch_package_pdf(launch_data):
    pdf_buffer = BytesIO()
    document = SimpleDocTemplate(
        pdf_buffer,
        pagesize=A4,
        rightMargin=20 * mm,
        leftMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "LaunchPackageTitle",
        parent=styles["Title"],
        alignment=TA_CENTER,
        textColor=colors.HexColor("#172554"),
        fontSize=22,
        leading=28,
        spaceAfter=8
    )
    brand_style = ParagraphStyle(
        "LaunchPackageBrand",
        parent=styles["Normal"],
        alignment=TA_CENTER,
        textColor=colors.HexColor("#2563eb"),
        fontSize=10,
        leading=14,
        spaceAfter=6
    )
    subtitle_style = ParagraphStyle(
        "LaunchPackageSubtitle",
        parent=styles["BodyText"],
        alignment=TA_CENTER,
        textColor=colors.HexColor("#64748b"),
        fontSize=10,
        leading=14,
        spaceAfter=12
    )
    heading_style = ParagraphStyle(
        "LaunchPackageHeading",
        parent=styles["Heading2"],
        textColor=colors.HexColor("#1e3a8a"),
        fontSize=14,
        leading=18,
        spaceBefore=12,
        spaceAfter=5
    )
    subheading_style = ParagraphStyle(
        "LaunchPackageSubheading",
        parent=styles["Heading3"],
        textColor=colors.HexColor("#334155"),
        fontSize=11.5,
        leading=15,
        spaceBefore=7,
        spaceAfter=4
    )
    body_style = ParagraphStyle(
        "LaunchPackageBody",
        parent=styles["BodyText"],
        textColor=colors.HexColor("#334155"),
        fontSize=10,
        leading=14,
        spaceAfter=5
    )
    bullet_style = ParagraphStyle(
        "LaunchPackageBullet",
        parent=body_style,
        leftIndent=10,
        firstLineIndent=-8
    )

    story = [
        Paragraph("BusinessBuilder AI", brand_style),
        Paragraph("Final Launch Package", title_style),
        Paragraph(
            "Your business strategy, draft store assets, branding outputs, and launch priorities.",
            subtitle_style
        ),
        Spacer(1, 3 * mm),
        Paragraph(
            (
                f"Shopify products: {len(launch_data['shopify_products'])} &nbsp; | &nbsp; "
                f"Shopify collections: {len(launch_data['shopify_collections'])} &nbsp; | &nbsp; "
                f"Shopify pages: {len(launch_data['shopify_pages'])} &nbsp; | &nbsp; "
                f"Canva drafts: {len(launch_data['canva_designs'])}"
            ),
            subtitle_style
        )
    ]

    def add_heading(title, style=heading_style):
        story.append(Paragraph(escape(title), style))

    def add_text(content):
        for line in str(content or "").splitlines():
            text = line.strip()

            if not text:
                story.append(Spacer(1, 2 * mm))
                continue

            clean_text = escape(text)
            is_heading = (
                text.startswith("#")
                or (
                    len(text) <= 90
                    and (
                        text.endswith(":")
                        or text[0].isdigit() and "." in text[:4]
                    )
                )
            )

            if text.startswith("#"):
                clean_text = escape(text.lstrip("#").strip())

            story.append(
                Paragraph(
                    clean_text,
                    subheading_style if is_heading else body_style
                )
            )

    def add_bullet(text):
        story.append(Paragraph(f"&bull; {escape(str(text))}", bullet_style))

    def workflow_answer(step_number, fallback):
        return launch_data["answers_by_step"].get(step_number, {}).get(
            "answer",
            fallback
        )

    add_heading("Business Summary")
    add_text(workflow_answer(1, "Complete workflow step 1 to add your business summary."))

    add_heading("Brand Summary")
    add_text(workflow_answer(2, "Complete workflow step 2 to add your brand summary."))

    add_heading("Target Market")
    add_text(workflow_answer(3, "Complete workflow step 3 to define your target market."))

    add_heading("Products / Services")
    add_text(workflow_answer(4, "Complete workflow step 4 to define your products or services."))

    add_heading("Pricing")
    add_text(workflow_answer(5, "Complete workflow step 5 to add your pricing strategy."))

    add_heading("Marketing Plan")
    add_text(workflow_answer(6, "Complete workflow step 6 to add your marketing plan."))

    add_heading("Shopify Assets Created")
    add_text("Review and publish these draft or unpublished Shopify assets manually.")

    add_heading("Products", subheading_style)
    if launch_data["shopify_products"]:
        for product in launch_data["shopify_products"]:
            add_bullet(f"{product[0]} - {product[3]}")
    else:
        add_text("No draft Shopify products created yet.")

    add_heading("Collections", subheading_style)
    if launch_data["shopify_collections"]:
        for collection in launch_data["shopify_collections"]:
            add_bullet(f"{collection[0]} - {collection[3]}")
    else:
        add_text("No unpublished Shopify collections created yet.")

    add_heading("Pages", subheading_style)
    if launch_data["shopify_pages"]:
        for page in launch_data["shopify_pages"]:
            add_bullet(f"{page[0]} - {page[4]}")
    else:
        add_text("No unpublished Shopify pages created yet.")

    if launch_data["latest_shopify_plan"]:
        add_heading("Latest Shopify Setup Plan", subheading_style)
        add_text(launch_data["latest_shopify_plan"][2])

    add_heading("Canva Assets Created")
    if launch_data["canva_designs"]:
        for design in launch_data["canva_designs"]:
            add_bullet(f"{design[0]} - {design[4]}")
    else:
        add_text("No Canva design drafts created yet.")

    if launch_data["latest_canva_branding_package"]:
        add_heading("Latest Canva Branding Package", subheading_style)
        add_text(launch_data["latest_canva_branding_package"][2])

    if launch_data["latest_canva_design_brief"]:
        add_heading("Latest Canva Design Brief", subheading_style)
        add_text(launch_data["latest_canva_design_brief"][2])

    if launch_data["latest_business_plan"]:
        add_heading("Latest Generated Business Plan")
        add_text(launch_data["latest_business_plan"][2])

    if launch_data["latest_build_quote"]:
        add_heading("Latest Store + Branding Quote")
        add_text(launch_data["latest_build_quote"][2])

    add_heading("Launch Checklist")
    add_text(workflow_answer(9, "Complete workflow step 9 to add your launch checklist."))

    add_heading("Next 30-Day Action Plan")
    add_heading("Days 1-7", subheading_style)
    add_text("Review your business plan, validate pricing, and refine your product or service offer.")
    add_heading("Days 8-14", subheading_style)
    add_text("Review Shopify draft products, collections, and pages. Add final images and publish only approved assets.")
    add_heading("Days 15-21", subheading_style)
    add_text("Edit your Canva drafts, prepare social content, and schedule your launch marketing.")
    add_heading("Days 22-30", subheading_style)
    add_text("Run checkout tests, complete the launch checklist, publish approved assets, and monitor your first customer feedback.")

    document.build(story)
    pdf_buffer.seek(0)

    return pdf_buffer


# -----------------------------
# PAGE ROUTES
# -----------------------------

@app.route("/")
def home():
    if "user_id" not in session:
        return redirect("/landing")

    chats = get_chats(session["user_id"])

    return render_template(
        "index.html",
        chats=chats
    )


@app.route("/landing")
def landing():
    if "user_id" in session:
        return redirect("/dashboard")

    return render_template("landing.html")


@app.route("/signup", methods=["GET", "POST"])
def signup():
    if request.method == "GET":
        return render_template("signup.html")

    email = request.form.get("email", "").strip().lower()
    raw_password = request.form.get("password", "")

    if not is_valid_email(email):
        return render_template(
            "signup.html",
            message="Enter a valid email address."
        ), 400

    if not raw_password:
        return render_template(
            "signup.html",
            message="Enter a password."
        ), 400

    password = generate_password_hash(raw_password)
    conn = None

    try:
        conn = db()
        cur = conn.cursor()

        cur.execute(
            sql("""
                SELECT id
                FROM users
                WHERE LOWER(email) = ?
                LIMIT 1
            """),
            (email,)
        )

        if cur.fetchone():
            conn.close()

            return render_template(
                "signup.html",
                message="An account with that email already exists."
            ), 400

        cur.execute(
            sql("""
                INSERT INTO users (email, password)
                VALUES (?, ?)
            """),
            (email, password)
        )

        conn.commit()
        cur.execute(
            sql("""
                SELECT id
                FROM users
                WHERE LOWER(email) = ?
                LIMIT 1
            """),
            (email,)
        )
        user = cur.fetchone()
        conn.close()

        if not user:
            raise RuntimeError("Created user could not be loaded.")

        session["user_id"] = user[0]

        send_email(
            email,
            "Welcome to BusinessBuilder AI",
            (
                "Your BusinessBuilder AI workspace is ready. "
                "Start by choosing your package, then complete the "
                "guided workflow to build your plans, store drafts, and brand assets."
            )
        )

        return redirect("/pricing?onboarding=choose_plan")

    except (sqlite3.IntegrityError, psycopg2.IntegrityError):
        if conn:
            conn.rollback()
            conn.close()

        return render_template(
            "signup.html",
            message="An account with that email already exists."
        ), 400
    except Exception:
        if conn:
            conn.close()

        raise


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "GET":
        return render_template("login.html")

    email = request.form.get("email", "").strip().lower()
    password = request.form.get("password", "")

    if not is_valid_email(email) or not password:
        return render_template(
            "login.html",
            message="Enter a valid email address and password."
        ), 400

    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT id, password
            FROM users
            WHERE LOWER(email) = ?
        """),
        (email,)
    )

    user = cur.fetchone()
    conn.close()

    if user and check_password_hash(user[1], password):
        session["user_id"] = user[0]
        return redirect("/")

    return render_template(
        "login.html",
        message="Invalid email or password."
    ), 401


@app.route("/logout")
def logout():
    session.clear()
    return redirect("/login")


@app.route("/dashboard")
def dashboard():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    projects = get_projects(user_id)
    current_package = get_user_package(user_id)
    paid = user_is_paid(user_id)
    latest_payment = get_latest_payment(user_id)
    business_plans = get_business_plans(user_id)
    shopify_plans = get_shopify_plans(user_id)
    canva_branding_packages = get_canva_branding_packages(user_id)
    canva_design_briefs = get_canva_design_briefs(user_id)
    canva_designs = get_canva_designs(user_id)
    build_quotes = get_build_quotes(user_id)
    ai_store_builds = get_ai_store_builds(user_id)
    store_agent_tasks = get_store_agent_tasks(user_id)
    product_research_list = get_product_research_list(user_id)
    supplier_recommendations = get_supplier_recommendations(user_id)
    pricing_advice_list = get_pricing_advice_list(user_id)
    payment_guides = get_payment_guides(user_id)
    shopify_connection = get_shopify_connection(user_id)
    canva_connection = get_canva_connection(user_id)
    shopify_products = get_shopify_products(user_id)
    shopify_collections = get_shopify_collections(user_id)
    shopify_pages = get_shopify_pages(user_id)
    workflow_answers = get_nonempty_workflow_answers(user_id)
    onboarding = get_user_onboarding(user_id)
    active_project = get_active_business_project(user_id)
    launch_readiness = get_launch_readiness(user_id)

    if not paid:
        next_action = {
            "title": "Choose a package",
            "description": "Activate Starter, Pro, or Premium Build to unlock the guided business workspace.",
            "url": "/pricing",
            "label": "View Pricing"
        }
    elif not onboarding or not onboarding[9]:
        next_action = {
            "title": "Start your business setup",
            "description": "Answer a few questions so BusinessBuilder AI can personalize product, pricing, payment, supplier, and store recommendations.",
            "url": "/onboarding",
            "label": "Start Onboarding"
        }
    elif not workflow_answers:
        next_action = {
            "title": "Complete your business workflow",
            "description": "Save your business details once, then reuse them for plans, research, store assets, and launch packages.",
            "url": "/business_workflow",
            "label": "Open Workflow"
        }
    elif onboarding and onboarding[8] == "find products" and not product_research_list:
        next_action = {
            "title": "Find products for your idea",
            "description": "Your onboarding goal points to product discovery. Start with product research before building store drafts.",
            "url": "/product_finder",
            "label": "Find Products"
        }
    elif not product_research_list:
        next_action = {
            "title": "Find products before building the store",
            "description": "Use AI Product Finder to identify product ideas, sourcing paths, competitors, pricing, and first draft products.",
            "url": "/product_finder",
            "label": "Find Products"
        }
    elif not supplier_recommendations:
        next_action = {
            "title": "Compare supplier and platform choices",
            "description": "Use Supplier Finder to compare Amazon-style research, Alibaba, AliExpress, local suppliers, print-on-demand, digital platforms, and wholesale options.",
            "url": "/supplier_finder",
            "label": "Compare Suppliers"
        }
    elif not pricing_advice_list:
        next_action = {
            "title": "Generate pricing advice",
            "description": "Estimate price ranges, margins, break-even points, competitor positioning, and the first price to test.",
            "url": "/pricing_advisor",
            "label": "Get Pricing Advice"
        }
    elif not payment_guides:
        next_action = {
            "title": "Choose safe payment methods",
            "description": "Plan PayPal, Paystack, EFT, card payments, Shopify payments, COD, mobile money, and international payment options without sharing sensitive details.",
            "url": "/payment_guide",
            "label": "Open Payment Guide"
        }
    elif user_package_at_least(user_id, "Pro") and not ai_store_builds:
        next_action = {
            "title": "Generate your store package",
            "description": "Use AI Store Builder to create the full Shopify store package, then approve safe draft asset creation.",
            "url": "/store_builder",
            "label": "Generate My Store"
        }
    else:
        next_action = {
            "title": "Open the Build Center",
            "description": "Review progress, approvals, draft assets, and launch package steps from the command center.",
            "url": "/build_center",
            "label": "Open Build Center"
        }

    shopify_connection_summary = None

    if shopify_connection:
        shopify_connection_summary = {
            "shop_domain": shopify_connection[1],
            "status": shopify_connection[3]
        }

    canva_connection_summary = None

    if canva_connection:
        canva_connection_summary = {
            "status": canva_connection[2]
        }

    return render_template(
        "dashboard.html",
        projects=projects,
        paid=paid,
        latest_payment=latest_payment,
        business_plans=business_plans,
        shopify_plans=shopify_plans,
        canva_branding_packages=canva_branding_packages,
        canva_design_briefs=canva_design_briefs,
        canva_designs=canva_designs,
        build_quotes=build_quotes,
        ai_store_builds=ai_store_builds,
        store_agent_tasks=store_agent_tasks,
        product_research_list=product_research_list,
        shopify_connection=shopify_connection_summary,
        canva_connection=canva_connection_summary,
        shopify_products=shopify_products,
        shopify_collections=shopify_collections,
        shopify_pages=shopify_pages,
        usage_summary=get_usage_summary(user_id),
        next_action=next_action,
        onboarding=onboarding,
        active_project=active_project,
        launch_readiness=launch_readiness,
        usage_limit_message=get_usage_limit_message(
            request.args.get("usage_limit"),
            user_id
        ),
        package_required_message=get_package_required_message(
            request.args.get("package_required")
        ),
        current_package=current_package,
        current_plan_details=PLAN_DETAILS.get(current_package, {}),
        pro_package=user_package_at_least(user_id, "Pro"),
        premium_build=user_package_at_least(user_id, "Premium Build"),
        is_admin=is_admin_user()
    )


@app.route("/admin")
def admin():
    if "user_id" not in session:
        return redirect("/login")

    if not is_admin_user():
        return redirect("/dashboard")

    return render_template(
        "admin.html",
        **get_admin_dashboard_data()
    )


@app.route("/build_center")
def build_center():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]
    current_package = get_user_package(user_id)
    paid = user_is_paid(user_id)

    if not paid:
        return redirect("/dashboard")

    completed_steps = get_completed_steps(user_id)
    completed_count = len(completed_steps)
    total_steps = len(WORKFLOW_STEPS)
    workflow_answers = get_all_workflow_answers(user_id)
    business_plans = get_business_plans(user_id)
    shopify_plans = get_shopify_plans(user_id)
    shopify_products = get_shopify_products(user_id)
    shopify_collections = get_shopify_collections(user_id)
    shopify_pages = get_shopify_pages(user_id)
    ai_store_builds = get_ai_store_builds(user_id)
    canva_branding_packages = get_canva_branding_packages(user_id)
    canva_design_briefs = get_canva_design_briefs(user_id)
    canva_designs = get_canva_designs(user_id)
    build_quotes = get_build_quotes(user_id)
    latest_payment = get_latest_payment(user_id)
    shopify_connection = get_shopify_connection(user_id)
    canva_connection = get_canva_connection(user_id)
    store_agent_tasks = get_store_agent_tasks(user_id)
    product_research_list = get_product_research_list(user_id)
    supplier_recommendations = get_supplier_recommendations(user_id)
    pricing_advice_list = get_pricing_advice_list(user_id)
    payment_guides = get_payment_guides(user_id)
    domain_guides = get_domain_guides(user_id)
    email_campaigns = get_email_campaigns(user_id)
    app_recommendations = get_app_recommendations(user_id)
    connected_apps = get_connected_app_summaries(user_id)
    app_action_drafts = get_app_action_drafts(user_id)
    marketing_app_drafts = [draft for draft in app_action_drafts if draft[2] in {"email_campaign", "social_ad_draft"}]
    design_app_drafts = [draft for draft in app_action_drafts if draft[2] == "canva_design_brief"]
    website_store_app_drafts = [draft for draft in app_action_drafts if draft[2] in {"shopify_product_draft", "webflow_page_draft", "website_copy", "automation_plan"}]

    def status_for(completed, started=False):
        if completed:
            return "Completed"

        if started:
            return "In Progress"

        return "Not Started"

    workflow_started = bool(completed_steps or workflow_answers)
    shopify_connected = bool(
        shopify_connection
        and shopify_connection[3] == "connected"
    )
    canva_connected = bool(
        canva_connection
        and canva_connection[2] == "connected"
    )
    latest_canva_design_url = ""

    if canva_designs:
        latest_canva_design_url = canva_designs[0][2] or canva_designs[0][3]

    launch_readiness = get_launch_readiness(user_id)

    build_items = [
        {
            "title": f"{current_package} Package",
            "status": status_for(paid),
            "description": "Your verified package controls your available tools and usage limits.",
            "url": "/dashboard",
            "action": "View Payment Status",
            "count": "Verified" if latest_payment else "Active"
        },
        {
            "title": "Business Workflow",
            "status": status_for(completed_count == total_steps, workflow_started),
            "description": f"{completed_count} of {total_steps} guided setup steps completed.",
            "url": "/business_workflow",
            "action": "Open Workflow",
            "count": f"{completed_count} / {total_steps} steps"
        },
        {
            "title": "Launch Readiness",
            "status": status_for(launch_readiness["score"] >= 80, launch_readiness["score"] > 0),
            "description": "Track the checklist that shows how close this business is to launch.",
            "url": "/launch_readiness",
            "action": "View Launch Checklist",
            "count": f"{launch_readiness['score']}% ready"
        },
        {
            "title": "AI Store Agent",
            "status": status_for(
                any(task[5] == "Applied" for task in store_agent_tasks),
                bool(store_agent_tasks)
            ),
            "description": "Review, edit, approve, and safely apply Shopify and Canva store tasks from one command center.",
            "url": "/ai_store_agent",
            "action": "Open AI Store Agent",
            "count": f"{len(store_agent_tasks)} tasks"
        },
        {
            "title": "AI Product Finder",
            "status": status_for(product_research_list, workflow_started),
            "description": "Discover product ideas, supplier search paths, competitor examples, and first Shopify products.",
            "url": (
                f"/product_research/{product_research_list[0][0]}"
                if product_research_list
                else "/product_finder"
            ),
            "action": "View Product Research" if product_research_list else "Find Products",
            "secondary_url": "/product_finder",
            "secondary_action": "New Product Search",
            "count": f"{len(product_research_list)} saved"
        },
        {
            "title": "Supplier Finder",
            "status": status_for(supplier_recommendations, bool(product_research_list or workflow_started)),
            "description": "Compare Amazon-style research, Alibaba, AliExpress, local suppliers, print-on-demand, digital platforms, and wholesalers.",
            "url": "/supplier_finder",
            "action": "Find Suppliers",
            "count": f"{len(supplier_recommendations)} saved"
        },
        {
            "title": "Pricing Advisor",
            "status": status_for(pricing_advice_list, bool(product_research_list or workflow_started)),
            "description": "Generate price ranges, margin notes, break-even guidance, competitor pricing advice, and first-price tests.",
            "url": "/pricing_advisor",
            "action": "Get Pricing Advice",
            "count": f"{len(pricing_advice_list)} saved"
        },
        {
            "title": "Payment Guide",
            "status": status_for(payment_guides, bool(workflow_started or shopify_plans)),
            "description": "Choose PayPal, Paystack, EFT, card, Shopify payment, COD, mobile money, and international payment options safely.",
            "url": "/payment_guide",
            "action": "Open Payment Guide",
            "count": f"{len(payment_guides)} saved"
        },
        {
            "title": "Domain/DNS Helper",
            "status": status_for(domain_guides, bool(workflow_started)),
            "description": "Compare domain options and create safe purchase, DNS, Shopify, SSL, and connection guidance.",
            "url": (
                f"/domain_guide/{domain_guides[0][0]}"
                if domain_guides
                else "/domain_helper"
            ),
            "action": "View Latest Domain Guide" if domain_guides else "Open Domain Helper",
            "secondary_url": "/domain_helper",
            "secondary_action": "New Domain Guide",
            "count": f"{len(domain_guides)} saved"
        },
        {
            "title": "AI Store Builder",
            "status": status_for(ai_store_builds, workflow_started),
            "description": "Generate a complete Shopify store package and create safe draft store assets.",
            "url": (
                f"/store_build/{ai_store_builds[0][0]}"
                if ai_store_builds
                else "/store_builder"
            ),
            "action": "View Latest Store Build" if ai_store_builds else "Generate My Store",
            "secondary_url": "/store_builder",
            "secondary_action": "Open Store Builder",
            "count": f"{len(ai_store_builds)} saved"
        },
        {
            "title": "Business Plan",
            "status": status_for(business_plans, workflow_started),
            "description": "Generate and review your complete saved business strategy.",
            "url": (
                f"/business_plan/{business_plans[0][0]}"
                if business_plans
                else "/generate_business_plan"
            ),
            "action": "View Plan" if business_plans else "Generate Plan",
            "count": f"{len(business_plans)} saved"
        },
        {
            "title": "Shopify Setup Plan",
            "status": status_for(shopify_plans, 7 in completed_steps),
            "description": "Plan your store pages, products, collections, policies, and launch.",
            "url": (
                f"/shopify_plan/{shopify_plans[0][0]}"
                if shopify_plans
                else "/generate_shopify_plan"
            ),
            "action": "View Shopify Plan" if shopify_plans else "Generate Shopify Plan",
            "count": f"{len(shopify_plans)} saved"
        },
        {
            "title": "Shopify Connection",
            "status": status_for(shopify_connected),
            "description": (
                "Your Shopify Admin API connection is ready for draft product creation."
                if shopify_connected
                else "Connect and test Shopify before creating draft store products."
            ),
            "url": "/shopify_settings",
            "action": "Manage Shopify" if shopify_connected else "Connect Shopify"
        },
        {
            "title": "Shopify Products",
            "status": status_for(shopify_products, bool(shopify_plans)),
            "description": f"{len(shopify_products)} draft Shopify products created.",
            "url": "/business_workflow",
            "action": "Open Product Tools",
            "count": f"{len(shopify_products)} created"
        },
        {
            "title": "Shopify Store Draft",
            "status": status_for(
                shopify_collections and shopify_pages,
                bool(shopify_products or shopify_collections or shopify_pages)
            ),
            "description": "Create draft products, unpublished collections, and unpublished store pages in one guided build.",
            "url": "/build_approval?action=shopify_store",
            "action": "Build Shopify Store Draft",
            "secondary_url": (
                "/shopify_build_summary"
                if shopify_collections or shopify_pages
                else ""
            ),
            "secondary_action": "View Summary",
            "count": (
                f"{len(shopify_collections)} collections / "
                f"{len(shopify_pages)} pages"
            )
        },
        {
            "title": "Canva Branding Package",
            "status": status_for(canva_branding_packages, 8 in completed_steps),
            "description": "Create your logo, color, typography, and visual asset direction.",
            "url": (
                f"/canva_branding/{canva_branding_packages[0][0]}"
                if canva_branding_packages
                else "/generate_canva_branding"
            ),
            "action": "View Branding Package" if canva_branding_packages else "Generate Branding Package",
            "count": f"{len(canva_branding_packages)} saved"
        },
        {
            "title": "Canva Design Briefs",
            "status": status_for(canva_design_briefs, bool(canva_branding_packages)),
            "description": "Turn your brand direction into practical Canva creation instructions.",
            "url": (
                f"/canva_design_brief/{canva_design_briefs[0][0]}"
                if canva_design_briefs
                else "/generate_canva_design_brief"
            ),
            "action": "View Design Brief" if canva_design_briefs else "Generate Design Brief",
            "count": f"{len(canva_design_briefs)} saved"
        },
        {
            "title": "Canva Connection",
            "status": status_for(canva_connected),
            "description": (
                "Your Canva OAuth connection is ready for editable design drafts."
                if canva_connected
                else "Connect Canva before creating editable design drafts."
            ),
            "url": "/canva_settings",
            "action": "Manage Canva" if canva_connected else "Connect Canva"
        },
        {
            "title": "Canva Design Drafts",
            "status": status_for(canva_designs, bool(canva_design_briefs)),
            "description": f"{len(canva_designs)} editable Canva design drafts created.",
            "url": "/build_approval?action=canva_designs",
            "action": "Create Canva Design Drafts",
            "secondary_url": latest_canva_design_url,
            "secondary_action": "Open Latest Canva Draft",
            "secondary_external": bool(latest_canva_design_url),
            "count": f"{len(canva_designs)} created"
        },
        {
            "title": "Email Marketing Assistant",
            "status": status_for(email_campaigns, bool(ai_store_builds or workflow_started)),
            "description": "Create reviewable email strategy, subject lines, campaign copy, platform advice, and setup checklists.",
            "url": (
                f"/email_campaign/{email_campaigns[0][0]}"
                if email_campaigns
                else "/email_marketing"
            ),
            "action": "View Latest Campaign" if email_campaigns else "Create Email Campaign",
            "secondary_url": "/email_marketing",
            "secondary_action": "New Email Campaign",
            "count": f"{len(email_campaigns)} saved"
        },
        {
            "title": "Store + Branding Quote",
            "status": status_for(build_quotes, workflow_started),
            "description": "Review your saved informational store and branding estimate.",
            "url": (
                f"/build_quote/{build_quotes[0][0]}"
                if build_quotes
                else "/generate_build_quote"
            ),
            "action": "View Quote" if build_quotes else "Generate Quote",
            "count": f"{len(build_quotes)} saved"
        },
        {
            "title": "Launch Checklist",
            "status": status_for(9 in completed_steps, workflow_started),
            "description": "Complete your final pre-launch checks for branding, store, marketing, and support.",
            "url": "/workflow_step/9",
            "action": "Open Launch Checklist"
        }
    ]

    build_items.extend([
        {
            "title": "App Recommendations",
            "status": status_for(app_recommendations, workflow_started),
            "description": "Compare apps by budget, country, business type, goal, platform, and skill level.",
            "url": "/recommend_apps",
            "action": "Recommend Apps",
            "count": f"{len(app_recommendations)} saved"
        },
        {
            "title": "App Connections",
            "status": status_for(connected_apps, bool(app_recommendations)),
            "description": "Manage permission-first Shopify, Canva, email-platform, and guidance-only integrations.",
            "url": "/app_connection_agent",
            "action": "Manage Connections",
            "count": f"{len(connected_apps)} configured"
        },
        {
            "title": "Marketing Drafts",
            "status": status_for(marketing_app_drafts, bool(connected_apps or app_recommendations)),
            "description": "Create reviewable email and advertising plans without sending, publishing, or spending.",
            "url": "/app_connection_agent",
            "action": "Create Marketing Draft",
            "count": f"{len(marketing_app_drafts)} drafts"
        },
        {
            "title": "Design Drafts",
            "status": status_for(design_app_drafts, bool(canva_branding_packages or connected_apps)),
            "description": "Create Canva-ready briefs and visual asset instructions for approval.",
            "url": "/app_connection_agent",
            "action": "Create Design Draft",
            "count": f"{len(design_app_drafts)} drafts"
        },
        {
            "title": "Website/Store Drafts",
            "status": status_for(website_store_app_drafts, bool(ai_store_builds or app_recommendations)),
            "description": "Create website copy, Webflow page plans, and safe Shopify product drafts.",
            "url": "/app_connection_agent",
            "action": "Create Website/Store Draft",
            "count": f"{len(website_store_app_drafts)} drafts"
        }
    ])

    pro_features = {
        "AI Store Builder",
        "Shopify Connection",
        "Shopify Products",
        "Shopify Store Draft",
        "Canva Design Briefs",
        "Canva Connection",
        "Canva Design Drafts",
        "Email Marketing Assistant",
        "App Recommendations"
    }
    premium_features = {
        "App Connections",
        "Marketing Drafts",
        "Design Drafts",
        "Website/Store Drafts"
    }

    for item in build_items:
        if (
            item["title"] in pro_features
            and not user_package_at_least(user_id, "Pro")
        ):
            item["status"] = "Not Started"
            item["url"] = "/pricing"
            item["action"] = "Upgrade to Pro"
            item["secondary_url"] = ""
            item["required_plan"] = "Pro"
        elif (
            item["title"] in premium_features
            and not user_package_at_least(user_id, "Premium Build")
        ):
            item["status"] = "Locked"
            item["url"] = "/pricing"
            item["action"] = "Upgrade to Premium Build"
            item["secondary_url"] = ""
            item["required_plan"] = "Premium Build"

    next_recommended_action = next(
        item
        for item in build_items
        if item["status"] != "Completed"
    ) if any(
        item["status"] != "Completed"
        for item in build_items
    ) else {
        "title": "Review Your Build Center",
        "description": "Your core business build is complete. Review your saved outputs and prepare for launch.",
        "url": "/build_center",
        "action": "Review Build Center"
    }

    roadmap_steps = [
        {
            "number": "01",
            "title": "Onboarding / Business Details",
            "status": status_for(completed_count == total_steps, workflow_started),
            "description": "Complete the guided answers that power plans, store copy, branding, and launch tasks.",
            "url": "/onboarding" if not workflow_started else "/business_workflow",
            "action": "Start Onboarding" if not workflow_started else "Open Workflow"
        },
        {
            "number": "02",
            "title": "Product Finder",
            "status": status_for(product_research_list, workflow_started),
            "description": "Research product ideas, sourcing paths, competitor examples, and price angles.",
            "url": "/product_finder",
            "action": "Find Products"
        },
        {
            "number": "03",
            "title": "Supplier Finder",
            "status": status_for(supplier_recommendations, bool(product_research_list)),
            "description": "Compare supplier and platform options before buying inventory or choosing a fulfillment path.",
            "url": "/supplier_finder",
            "action": "Compare Suppliers"
        },
        {
            "number": "04",
            "title": "Pricing Advisor",
            "status": status_for(pricing_advice_list, bool(product_research_list or supplier_recommendations)),
            "description": "Estimate price ranges, margins, break-even points, and a first price to test.",
            "url": "/pricing_advisor",
            "action": "Get Pricing Advice"
        },
        {
            "number": "05",
            "title": "Payment Guide",
            "status": status_for(payment_guides, bool(pricing_advice_list or workflow_started)),
            "description": "Plan PayPal, Paystack, Shopify, EFT, card, COD, mobile money, and international payment setup.",
            "url": "/payment_guide",
            "action": "Choose Payments"
        },
        {
            "number": "06",
            "title": "Domain Helper",
            "status": status_for(domain_guides, bool(payment_guides or workflow_started)),
            "description": "Choose a domain and prepare safe DNS, SSL, Shopify, and custom-domain connection steps.",
            "url": "/domain_helper",
            "action": "Create Domain Guide"
        },
        {
            "number": "07",
            "title": "Store Draft",
            "status": status_for(ai_store_builds or store_agent_tasks, product_research_list),
            "description": "Generate reviewable store drafts and AI Store Agent tasks before applying anything.",
            "url": "/ai_store_agent",
            "action": "Open Agent"
        },
        {
            "number": "08",
            "title": "Shopify Assets",
            "status": status_for(shopify_products or shopify_collections or shopify_pages, shopify_connected),
            "description": "Create supported draft products, collections, pages, and setup guidance after approval.",
            "url": "/shopify_settings" if not shopify_connected else "/shopify_build_summary",
            "action": "Manage Shopify" if not shopify_connected else "View Assets"
        },
        {
            "number": "09",
            "title": "Canva Branding",
            "status": status_for(canva_branding_packages or canva_design_briefs or canva_designs, canva_connected),
            "description": "Create brand direction, logo briefs, social ideas, and marketing asset briefs.",
            "url": "/canva_settings" if not canva_connected else "/generate_canva_branding",
            "action": "Manage Canva" if not canva_connected else "Create Branding"
        },
        {
            "number": "10",
            "title": "Email Marketing",
            "status": status_for(email_campaigns, bool(ai_store_builds or canva_branding_packages)),
            "description": "Create a reviewable launch, welcome, newsletter, promotion, or ecommerce email campaign draft.",
            "url": "/email_marketing",
            "action": "Create Campaign"
        },
        {
            "number": "11",
            "title": "App Recommendations",
            "status": status_for(app_recommendations, workflow_started),
            "description": "Choose suitable platforms based on budget, market, business type, goals, and skill level.",
            "url": "/recommend_apps",
            "action": "Recommend Apps"
        },
        {
            "number": "12",
            "title": "App Connections",
            "status": status_for(connected_apps, bool(app_recommendations)),
            "description": "Connect supported platforms with permission or use guidance-only setup paths.",
            "url": "/app_connection_agent",
            "action": "Manage Connections"
        },
        {
            "number": "13",
            "title": "Marketing Drafts",
            "status": status_for(marketing_app_drafts, bool(app_recommendations or connected_apps)),
            "description": "Prepare email and advertising drafts without sending, publishing, or spending.",
            "url": "/app_connection_agent",
            "action": "Create Marketing Draft"
        },
        {
            "number": "14",
            "title": "Design Drafts",
            "status": status_for(design_app_drafts, bool(canva_branding_packages or connected_apps)),
            "description": "Prepare Canva-ready design briefs for review and approval.",
            "url": "/app_connection_agent",
            "action": "Create Design Draft"
        },
        {
            "number": "15",
            "title": "Website/Store Drafts",
            "status": status_for(website_store_app_drafts, bool(ai_store_builds or app_recommendations)),
            "description": "Prepare website copy, page plans, automation plans, and draft store assets.",
            "url": "/app_connection_agent",
            "action": "Create Website/Store Draft"
        },
        {
            "number": "16",
            "title": "Launch Readiness",
            "status": status_for(launch_readiness["score"] >= 80, launch_readiness["score"] > 0),
            "description": "Check completed and missing launch items before you go live.",
            "url": "/launch_readiness",
            "action": "Check Score"
        },
        {
            "number": "17",
            "title": "Launch Package",
            "status": status_for(user_package_at_least(user_id, "Pro") and launch_readiness["score"] >= 80, launch_readiness["score"] > 0),
            "description": "Bundle strategy, research, store assets, branding, and next steps into one package.",
            "url": "/launch_package" if user_package_at_least(user_id, "Pro") else "/pricing",
            "action": "Open Package" if user_package_at_least(user_id, "Pro") else "Upgrade to Pro"
        }
    ]

    roadmap_requirements = {
        "Store Draft": "Pro", "Shopify Assets": "Pro", "Canva Branding": "Pro",
        "Email Marketing": "Pro", "App Recommendations": "Pro", "Launch Package": "Pro",
        "App Connections": "Premium Build", "Marketing Drafts": "Premium Build",
        "Design Drafts": "Premium Build", "Website/Store Drafts": "Premium Build"
    }
    for step in roadmap_steps:
        required_plan = roadmap_requirements.get(step["title"], "Starter")
        step["required_plan"] = required_plan
        step["unlocked"] = user_package_at_least(user_id, required_plan)
        if not step["unlocked"]:
            step["status"] = "Locked"
            step["url"] = "/pricing"
            step["action"] = f"Upgrade to {required_plan}"

    return render_template(
        "build_center.html",
        build_items=build_items,
        roadmap_steps=roadmap_steps,
        completed_count=completed_count,
        total_steps=total_steps,
        next_recommended_action=next_recommended_action,
        current_package=current_package,
        current_plan_details=PLAN_DETAILS.get(current_package, {}),
        shopify_connected=shopify_connected,
        canva_connected=canva_connected,
        pro_package=user_package_at_least(user_id, "Pro"),
        premium_build=user_package_at_least(user_id, "Premium Build")
    )


@app.route("/store_builder")
def store_builder():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    workflow_answers = get_nonempty_workflow_answers(user_id)
    completed_steps = get_completed_steps(user_id)
    shopify_connection = get_shopify_connection(user_id)
    canva_connection = get_canva_connection(user_id)

    shopify_connected = bool(
        shopify_connection
        and shopify_connection[3] == "connected"
    )
    canva_connected = bool(
        canva_connection
        and canva_connection[2] == "connected"
    )
    pro_package = user_package_at_least(user_id, "Pro")
    generate_url = ""

    if pro_package:
        generate_url = (
            "/build_approval?action=ai_store"
            if shopify_connected
            else "/generate_full_store"
        )

    return render_template(
        "store_builder.html",
        workflow_answers=workflow_answers,
        workflow_completed=len(completed_steps) == len(WORKFLOW_STEPS),
        has_workflow_answers=bool(workflow_answers),
        shopify_connected=shopify_connected,
        canva_connected=canva_connected,
        paid_plan_active=True,
        pro_package=pro_package,
        store_builds=get_ai_store_builds(user_id),
        generate_url=generate_url
    )


def get_store_agent_sections(user_id):
    latest_tasks = get_latest_store_agent_tasks_by_type(user_id)
    domain_guides = get_domain_guides(user_id)
    email_campaigns = get_email_campaigns(user_id)
    sections = []

    for task_type in STORE_AGENT_SECTION_ORDER:
        definition = STORE_AGENT_TASK_DEFINITIONS[task_type]
        task = latest_tasks.get(task_type)
        status = "Not Started"

        if task:
            status = task[5]
        elif task_type == "domain_setup" and domain_guides:
            status = "Completed"
        elif task_type == "email_marketing" and email_campaigns:
            status = "Completed"

        sections.append({
            "task_type": task_type,
            "section": definition["section"],
            "title": definition["title"],
            "mode": definition["mode"],
            "explanation": definition["explanation"],
            "status": status,
            "task": task,
            "apply_supported": definition["apply_supported"],
            "url": {
                "product_sourcing": "/product_finder",
                "supplier_finder": "/supplier_finder",
                "pricing_advisor": "/pricing_advisor",
                "payments_setup": "/payment_guide",
                "domain_setup": "/domain_helper",
                "email_marketing": "/email_marketing"
            }.get(task_type, "")
        })

    return sections


def store_agent_json_prompt(task_type):
    if task_type == "products":
        return """
Return JSON only with one top-level key named "products".
The products value must be an array of 5 to 10 objects.
Each product must include title, description, suggested_price, category, seo_title, and meta_description.
Products must be suitable for draft Shopify creation and must not be published.
"""

    if task_type == "pages":
        return """
Return JSON only with one top-level key named "pages".
The pages value must be an array of objects.
Each page object must include title, page_type, and content.
Use these page_type values where appropriate: about, contact, faq, shipping_policy, refund_policy.
Pages must be customer-facing drafts and must not claim the store is live.
"""

    return ""


def build_store_agent_prompt(task_type, workflow_text, extra_context=""):
    definition = STORE_AGENT_TASK_DEFINITIONS[task_type]
    section = definition["section"]
    safety = """
Safety rules:
- Do not claim the store is live or published.
- Do not buy domains, charge the user, collect banking details, or configure sensitive payment information.
- Do not install themes, publish products, publish pages, or edit live theme files.
- Clearly separate advice, draft content, and actions that need user approval.
"""
    task_instructions = {
        "theme_recommendation": """
Ask and answer these theme planning points:
1. Whether the user already has a Shopify theme installed.
2. Whether free or paid theme suggestions fit their current budget.
3. Recommended styles such as luxury, simple, bold, playful, tech, fashion, beauty, or another fit.
Recommend Shopify themes and explain why each one fits. Include manual installation guidance only.
""",
        "homepage_design": """
Generate a full homepage design draft with:
1. Hero section
2. CTA buttons
3. Featured collections
4. Product sections
5. About section
6. Trust badges
7. FAQ section
8. Footer copy
9. Canva banner brief
Do not edit theme files.
""",
        "theme_sections": """
Generate a Shopify theme section plan with section names, order, copy, image/design notes, and manual setup instructions.
Do not edit live theme files or install a theme.
""",
        "products": """
Generate Shopify product drafts with descriptions, categories, suggested prices, SEO titles, and meta descriptions.
""",
        "product_sourcing": """
Generate product sourcing research, supplier search paths, competitor examples to study, and initial product recommendations.
Do not scrape Amazon or claim live supplier access.
""",
        "supplier_finder": """
Compare supplier and platform options for the user's business, budget, country, and product type.
Cover Amazon-style research, Alibaba, AliExpress, local suppliers, print-on-demand, digital product platforms, wholesalers, handmade/local manufacturing, and direct manufacturers where relevant.
Do not scrape websites, claim live prices, claim live stock, or claim delivery availability.
Tell the user what to verify before buying inventory or choosing a supplier.
""",
        "pricing_advisor": """
Generate practical pricing advice with suggested selling price range, low/standard/premium options, profit margin notes, break-even explanation, competitor pricing advice, discount ideas, too-low and too-high warnings, first price to test, and customer-facing price explanation.
Use estimates only and tell the user to verify costs, fees, taxes, packaging, shipping, and competitor positioning.
""",
        "pages": """
Generate Shopify page drafts for About, Contact, FAQ, Shipping Policy, and Refund Policy.
""",
        "shipping_zones": """
Generate recommended shipping zones, local delivery options, free shipping thresholds, courier notes, and Shopify manual setup steps.
If API support is not available, make clear the user must configure this inside Shopify.
""",
        "payments_setup": """
Generate a payment setup checklist covering PayPal, Paystack, Shopify Payments, EFT/bank transfer, card payments, cash on delivery, mobile money, and international payment considerations where relevant.
Explain that requirements vary by country and provider.
Explain that banking details, identity checks, tax details, card numbers, passwords, and private credentials must be entered directly with the payment provider or Shopify, never inside BusinessBuilder AI.
Do not create payment accounts automatically.
""",
        "domain_setup": """
Suggest domain names, providers, cheaper offers, likely costs, buying decision questions, and Shopify connection instructions.
Make clear the user must approve billing and payment directly with the domain provider or Shopify.
""",
        "canva_branding": """
Generate a Canva branding package with logo brief, banner brief, product mockup brief, social post brief, and brand kit suggestions.
""",
        "canva_marketing": """
Generate Canva marketing asset briefs for launch posts, product posts, story templates, promo banners, ads, and short-form video covers.
""",
        "email_marketing": """
Generate an email campaign strategy, subject and preview options, full email body, CTA copy, audience segment, sending-time guidance, follow-up idea, platform recommendation, and setup checklist.
Do not send emails or create external email-platform accounts.
""",
        "launch_checklist": """
Generate a launch checklist that guides the user through reviewing store content, legal pages, payments, shipping, domains, Canva assets, test orders, and final manual publishing.
"""
    }

    return f"""
You are the BusinessBuilder AI Store Agent working in {definition["mode"]}.
Generate the draft for this command-center section: {section}.

{safety}

Task instructions:
{task_instructions[task_type]}

{store_agent_json_prompt(task_type)}

Extra user preferences:
{extra_context or "No extra preferences supplied."}

User workflow answers:
{workflow_text}
"""


def parse_task_json_content(content):
    try:
        return json.loads(content)
    except (TypeError, ValueError, json.JSONDecodeError):
        pass

    match = re.search(r"\{.*\}", content or "", re.DOTALL)

    if not match:
        return None

    try:
        return json.loads(match.group(0))
    except (TypeError, ValueError, json.JSONDecodeError):
        return None


@app.route("/ai_store_agent")
def ai_store_agent():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    app_drafts = get_app_action_drafts(user_id)
    connected_apps = get_connected_app_summaries(user_id)

    return render_template(
        "ai_store_agent.html",
        sections=get_store_agent_sections(user_id),
        task_definitions=STORE_AGENT_TASK_DEFINITIONS,
        workflow_answers=get_nonempty_workflow_answers(user_id),
        shopify_connected=bool(
            get_shopify_connection(user_id)
            and get_shopify_connection(user_id)[3] == "connected"
        ),
        canva_connected=bool(
            get_canva_connection(user_id)
            and get_canva_connection(user_id)[2] == "connected"
        ),
        connected_app_count=len(connected_apps),
        pending_app_draft_count=sum(1 for draft in app_drafts if not draft[7]),
        current_package=get_user_package(user_id) or "Starter"
    )


@app.route("/generate_store_agent_task/<task_type>", methods=["GET", "POST"])
def generate_store_agent_task(task_type):
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]
    task_type = normalize_store_agent_task_type(task_type)

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    if task_type not in STORE_AGENT_TASK_DEFINITIONS:
        return redirect("/ai_store_agent")

    if usage_limit_reached(user_id, "store_agent_task"):
        return usage_limit_redirect("store_agent_task", "/ai_store_agent")

    answers = get_nonempty_workflow_answers(user_id)

    if not answers:
        return redirect("/ai_store_agent?agent_error=no_answers")

    extra_context = ""

    if task_type == "theme_recommendation":
        extra_context = (
            f"Existing Shopify theme installed: {request.values.get('has_theme', 'Not answered')}\n"
            f"Theme budget preference: {request.values.get('theme_budget', 'Not answered')}\n"
            f"Preferred style: {request.values.get('theme_style', 'Not answered')}"
        )
    elif request.values.get("notes"):
        extra_context = request.values.get("notes", "").strip()

    workflow_text = build_workflow_text(answers)
    prompt = build_store_agent_prompt(task_type, workflow_text, extra_context)

    completion_kwargs = {
        "model": "gpt-4.1-mini",
        "messages": [
            {
                "role": "system",
                "content": SYSTEM_PROMPT
            },
            {
                "role": "user",
                "content": prompt
            }
        ]
    }

    if task_type in {"products", "pages"}:
        completion_kwargs["response_format"] = {"type": "json_object"}

    try:
        response = safe_openai_chat_completion(
            **completion_kwargs
        )
    except Exception:
        return redirect("/ai_store_agent?agent_error=ai_response")

    draft_content = response.choices[0].message.content.strip()
    definition = STORE_AGENT_TASK_DEFINITIONS[task_type]
    task_id = create_store_agent_task(
        user_id,
        task_type,
        definition["title"],
        draft_content,
        "Needs Approval"
    )
    log_usage(user_id, "store_agent_task")

    return redirect(f"/review_store_agent_task/{task_id}")


@app.route("/review_store_agent_task/<int:task_id>")
def review_store_agent_task(task_id):
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    task = get_store_agent_task(user_id, task_id)

    if not task:
        return redirect("/ai_store_agent")

    return render_template(
        "review_store_agent_task.html",
        task=task,
        definition=STORE_AGENT_TASK_DEFINITIONS.get(task[2], {}),
        pro_package=user_package_at_least(user_id, "Pro"),
        ai_edit_mode=request.args.get("edit") == "ai",
        manual_edit_mode=request.args.get("edit") == "manual"
    )


@app.route("/approve_store_agent_task/<int:task_id>", methods=["POST"])
def approve_store_agent_task(task_id):
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    task = get_store_agent_task(user_id, task_id)

    if not task:
        return redirect("/ai_store_agent")

    update_store_agent_task_status(task_id, user_id, "Approved", True, task[7])

    return redirect(f"/review_store_agent_task/{task_id}?agent_notice=approved")


@app.route("/edit_store_agent_task_ai/<int:task_id>", methods=["GET", "POST"])
def edit_store_agent_task_ai(task_id):
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    task = get_store_agent_task(user_id, task_id)

    if not task:
        return redirect("/ai_store_agent")

    if request.method == "GET":
        return redirect(f"/review_store_agent_task/{task_id}?edit=ai")

    feedback = request.form.get("feedback", "").strip()

    if not feedback:
        return redirect(f"/review_store_agent_task/{task_id}?edit=ai&agent_error=feedback")

    prompt = f"""
Revise this BusinessBuilder AI Store Agent draft using the user's feedback.
Keep the same task type and keep all safety rules: do not publish, charge, buy domains,
collect banking information, install themes, or edit live theme files.

User feedback:
{feedback}

Current draft:
{task[4]}
"""

    completion_kwargs = {
        "model": "gpt-4.1-mini",
        "messages": [
            {
                "role": "system",
                "content": SYSTEM_PROMPT
            },
            {
                "role": "user",
                "content": prompt
            }
        ]
    }

    if task[2] in {"products", "pages"}:
        completion_kwargs["response_format"] = {"type": "json_object"}

    try:
        response = safe_openai_chat_completion(
            **completion_kwargs
        )
    except Exception:
        return redirect(f"/review_store_agent_task/{task_id}?agent_error=ai_response")

    update_store_agent_task_content(
        task_id,
        user_id,
        response.choices[0].message.content.strip()
    )

    return redirect(f"/review_store_agent_task/{task_id}?agent_notice=updated")


@app.route("/edit_store_agent_task_manual/<int:task_id>", methods=["GET", "POST"])
def edit_store_agent_task_manual(task_id):
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    task = get_store_agent_task(user_id, task_id)

    if not task:
        return redirect("/ai_store_agent")

    if request.method == "GET":
        return redirect(f"/review_store_agent_task/{task_id}?edit=manual")

    draft_content = request.form.get("draft_content", "").strip()
    save_status = request.form.get("save_status", "needs_approval")

    if not draft_content:
        return redirect(f"/review_store_agent_task/{task_id}?edit=manual&agent_error=empty")

    update_store_agent_task_content(task_id, user_id, draft_content)

    if save_status == "approved":
        update_store_agent_task_status(task_id, user_id, "Approved", True, False)

    return redirect(f"/review_store_agent_task/{task_id}?agent_notice=updated")


@app.route("/delete_store_agent_task/<int:task_id>", methods=["POST"])
def delete_store_agent_task(task_id):
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    delete_store_agent_task_record(task_id, user_id)

    return redirect("/ai_store_agent?agent_notice=deleted")


def apply_store_agent_products(user_id, task):
    parsed = parse_task_json_content(task[4]) or {}
    products = parsed.get("products", [])
    created_count = 0
    failed_count = 0

    if not isinstance(products, list):
        return False, "The product draft is not valid JSON. Edit it before applying."

    if not get_shopify_connection(user_id) or get_shopify_connection(user_id)[3] != "connected":
        return False, "Connect Shopify before applying product drafts."

    for product in products:
        if usage_limit_reached(user_id, "shopify_product"):
            break

        try:
            title = str(product["title"]).strip()
            description = str(product["description"]).strip()
            suggested_price = str(product.get("suggested_price", "")).strip()
            category = str(product.get("category", "")).strip()
            seo_title = str(product.get("seo_title", "")).strip()
            meta_description = str(product.get("meta_description", "")).strip()
        except (KeyError, TypeError):
            failed_count += 1
            continue

        shopify_product_id, status, error = create_shopify_product(
            user_id,
            title,
            description,
            suggested_price,
            category
        )

        if error:
            failed_count += 1
            continue

        save_shopify_product(
            user_id,
            title,
            (
                f"{description}\n\nSuggested price: {suggested_price}\n"
                f"Collection / category: {category}\nSEO title: {seo_title}\n"
                f"Meta description: {meta_description}"
            ),
            shopify_product_id,
            status
        )
        log_usage(user_id, "shopify_product")
        created_count += 1

    if not created_count:
        return False, "No draft Shopify products were created."

    return True, f"{created_count} draft products created. {failed_count} products failed."


def apply_store_agent_pages(user_id, task):
    parsed = parse_task_json_content(task[4]) or {}
    pages = parsed.get("pages", [])
    created_count = 0
    failed_count = 0

    if not isinstance(pages, list):
        return False, "The page draft is not valid JSON. Edit it before applying."

    if not get_shopify_connection(user_id) or get_shopify_connection(user_id)[3] != "connected":
        return False, "Connect Shopify before applying page drafts."

    for page in pages:
        try:
            title = str(page["title"]).strip()
            page_type = str(page.get("page_type", "custom")).strip()
            content = str(page["content"]).strip()
        except (KeyError, TypeError):
            failed_count += 1
            continue

        shopify_page_id, status, error = create_shopify_page(
            user_id,
            title,
            content
        )

        if error:
            failed_count += 1
            continue

        save_shopify_page(
            user_id,
            title,
            page_type,
            content,
            shopify_page_id,
            status
        )
        created_count += 1

    if not created_count:
        return False, "No unpublished Shopify pages were created."

    return True, f"{created_count} unpublished pages created. {failed_count} pages failed."


def apply_store_agent_canva(user_id, task):
    if usage_limit_reached(user_id, "canva_design_brief"):
        return False, get_usage_limit_message("canva_design_brief", user_id)

    brief_id = save_canva_design_brief(
        user_id,
        task[3],
        task[4]
    )
    created_design = False

    if get_canva_connection(user_id) and get_canva_connection(user_id)[2] == "connected":
        if not usage_limit_reached(user_id, "canva_design"):
            canva_design_id, edit_url, view_url, status, error = create_canva_design(
                user_id,
                f"{task[3]} Draft"[:255]
            )

            if not error:
                save_canva_design(
                    user_id,
                    f"{task[3]} Draft"[:255],
                    canva_design_id,
                    edit_url,
                    view_url,
                    status
                )
                log_usage(user_id, "canva_design")
                created_design = True

    log_usage(user_id, "canva_design_brief")

    if created_design:
        return True, f"Canva design brief saved and one Canva draft created. Brief ID: {brief_id}."

    return True, f"Canva design brief saved. Brief ID: {brief_id}."


@app.route("/apply_store_agent_task/<int:task_id>", methods=["POST"])
def apply_store_agent_task(task_id):
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    task = get_store_agent_task(user_id, task_id)

    if not task:
        return redirect("/ai_store_agent")

    if not user_package_at_least(user_id, "Pro"):
        return package_access_redirect("Pro", f"/review_store_agent_task/{task_id}")

    if not task[6]:
        return redirect(f"/review_store_agent_task/{task_id}?agent_error=not_approved")

    task_type = task[2]

    if task_type == "products":
        applied, message = apply_store_agent_products(user_id, task)
    elif task_type == "pages":
        applied, message = apply_store_agent_pages(user_id, task)
    elif task_type in {"canva_branding", "canva_marketing"}:
        applied, message = apply_store_agent_canva(user_id, task)
    elif task_type in {"homepage_design", "theme_sections"}:
        save_shopify_plan(user_id, "Approved Homepage Design Plan", task[4])
        applied, message = True, "Homepage design plan saved. Theme files were not edited."
    elif task_type == "theme_recommendation":
        save_shopify_plan(user_id, "Approved Theme Recommendation", task[4])
        applied, message = True, "Theme recommendation saved. No theme was installed."
    elif task_type in {"shipping_zones", "payments_setup", "domain_setup", "launch_checklist"}:
        save_shopify_plan(user_id, task[3], task[4])
        applied = True
        message = "Approved guidance saved. Manual review is still required inside Shopify or the provider account."
    else:
        applied, message = False, "This task type is not supported."

    if applied:
        update_store_agent_task_status(task_id, user_id, "Applied", True, True)
        return redirect(
            f"/review_store_agent_task/{task_id}?agent_notice=applied"
            f"&apply_message={urllib.parse.quote(message)}"
        )

    return redirect(
        f"/review_store_agent_task/{task_id}?agent_error=apply_failed"
        f"&apply_message={urllib.parse.quote(message)}"
    )


def product_research_to_text(research_data):
    if not isinstance(research_data, dict):
        return ""

    lines = []
    labels = [
        ("recommended_product_ideas", "Recommended product ideas"),
        ("supplier_source_options", "Supplier/source options"),
        ("amazon_research_suggestions", "Amazon product research suggestions"),
        ("competitor_examples", "Similar business/competitor examples to research"),
        ("product_differentiation_ideas", "Product differentiation ideas"),
        ("recommended_first_5_shopify_products", "Recommended first 5 products to add to Shopify"),
        ("risky_product_warnings", "Warnings about risky products"),
        ("next_steps", "Next steps")
    ]

    for key, label in labels:
        value = research_data.get(key)

        if not value:
            continue

        lines.append(f"{label}:")

        if isinstance(value, list):
            for item in value:
                if isinstance(item, dict):
                    title = item.get("name") or item.get("title") or "Item"
                    lines.append(f"- {title}")

                    for detail_key, detail_value in item.items():
                        if detail_key in {"name", "title"}:
                            continue
                        lines.append(f"  {detail_key.replace('_', ' ').title()}: {detail_value}")
                else:
                    lines.append(f"- {item}")
        elif isinstance(value, dict):
            for detail_key, detail_value in value.items():
                lines.append(f"- {detail_key.replace('_', ' ').title()}: {detail_value}")
        else:
            lines.append(str(value))

        lines.append("")

    return "\n".join(lines).strip()


@app.route("/product_finder")
def product_finder():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    return render_template(
        "product_finder.html",
        product_research_list=get_product_research_list(user_id)
    )


@app.route("/generate_product_research", methods=["GET", "POST"])
def generate_product_research():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    if request.method == "GET":
        return redirect("/product_finder")

    if usage_limit_reached(user_id, "product_research"):
        return usage_limit_redirect("product_research", "/product_finder")

    business_type = request.form.get("business_type", "").strip()
    target_market = request.form.get("target_market", "").strip()
    country = request.form.get("country", "").strip()
    budget = request.form.get("budget", "").strip()
    product_type = request.form.get("product_type", "").strip()
    sourcing_preference = request.form.get("sourcing_preference", "").strip()
    risk_level = request.form.get("risk_level", "").strip()

    if not business_type or not target_market:
        return redirect("/product_finder?research_error=missing_fields")

    prompt = f"""
Create AI product sourcing and competitor discovery research for a new Shopify business.
Return JSON only.

Important Amazon rule:
- Do not scrape Amazon.
- Do not pretend to access Amazon live data.
- Generate Amazon search suggestions and research links only.
- Make clear that users must verify prices, suppliers, stock, policies, reviews, and restrictions themselves.

Use exactly these top-level keys:
- recommended_product_ideas
- supplier_source_options
- amazon_research_suggestions
- competitor_examples
- product_differentiation_ideas
- recommended_first_5_shopify_products
- risky_product_warnings
- next_steps

recommended_product_ideas must be an array of 8 to 12 objects with:
- name
- why_it_fits
- suggested_selling_price_range
- estimated_difficulty_level
- shipping_difficulty
- branding_angle
- profit_margin_notes

recommended_first_5_shopify_products must be an array of exactly 5 objects with:
- title
- description
- suggested_price
- category

supplier_source_options must include options for the user's sourcing preference and at least two alternatives.
competitor_examples must be research examples or search queries to investigate, not claims of live competitor data.

Business/store idea: {business_type}
Target customer: {target_market}
Country/market: {country}
Budget: {budget}
Product type: {product_type}
Sourcing preference: {sourcing_preference}
Risk level: {risk_level}
"""

    try:
        response = safe_openai_chat_completion(
            model="gpt-4.1-mini",
            response_format={"type": "json_object"},
            messages=[
                {
                    "role": "system",
                    "content": SYSTEM_PROMPT
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ]
        )
        research_data = json.loads(response.choices[0].message.content)
    except Exception:
        return redirect("/product_finder?research_error=ai_response")

    title = f"Product Research - {business_type[:80]}"
    research_id = save_product_research(
        user_id,
        title,
        business_type,
        target_market,
        sourcing_preference,
        budget,
        country,
        json.dumps(research_data, indent=2)
    )
    log_usage(user_id, "product_research")

    return redirect(f"/product_research/{research_id}")


@app.route("/product_research/<int:research_id>")
def product_research_detail(research_id):
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    research = get_product_research(user_id, research_id)

    if not research:
        return redirect("/product_finder")

    research_data = parse_product_research_content(research[7])

    return render_template(
        "product_research.html",
        research=research,
        research_data=research_data,
        research_text=product_research_to_text(research_data),
        shopify_connected=bool(
            get_shopify_connection(user_id)
            and get_shopify_connection(user_id)[3] == "connected"
        )
    )


@app.route("/create_product_research_shopify_products")
def create_product_research_shopify_products():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    if not consume_build_approval("product_research_products"):
        return build_approval_redirect("product_research_products")

    research_id = session.pop("product_research_id", None)

    if not research_id:
        return redirect("/product_finder?research_error=no_research")

    research = get_product_research(user_id, int(research_id))

    if not research:
        return redirect("/product_finder?research_error=no_research")

    connection = get_shopify_connection(user_id)

    if not connection or connection[3] != "connected":
        return redirect(f"/product_research/{research_id}?research_error=shopify_connection")

    research_data = parse_product_research_content(research[7]) or {}
    products = research_data.get("recommended_first_5_shopify_products", [])

    if not isinstance(products, list) or not products:
        return redirect(f"/product_research/{research_id}?research_error=no_products")

    created_count = 0
    failed_count = 0

    for product in products[:5]:
        if usage_limit_reached(user_id, "shopify_product"):
            break

        try:
            title = str(product["title"]).strip()
            description = str(product["description"]).strip()
            suggested_price = str(product.get("suggested_price", "")).strip()
            category = str(product.get("category", "")).strip()
        except (KeyError, TypeError):
            failed_count += 1
            continue

        if not title or not description:
            failed_count += 1
            continue

        shopify_product_id, status, error = create_shopify_product(
            user_id,
            title,
            description,
            suggested_price,
            category
        )

        if error:
            failed_count += 1
            continue

        save_shopify_product(
            user_id,
            title,
            (
                f"{description}\n\n"
                f"Suggested price: {suggested_price}\n"
                f"Collection / category: {category}\n"
                "Source: AI Product Finder research"
            ),
            shopify_product_id,
            status
        )
        log_usage(user_id, "shopify_product")
        created_count += 1

    if not created_count:
        return redirect(f"/product_research/{research_id}?research_error=create_failed")

    return redirect(
        f"/product_research/{research_id}?research_notice=products_created"
        f"&created_count={created_count}&failed_count={failed_count}"
    )


def get_launch_package_data(user_id):
    workflow_answers = get_all_workflow_answers(user_id)
    answers_by_step = {
        step_number: {
            "name": step_name,
            "answer": answer
        }
        for step_number, step_name, answer in workflow_answers
    }
    business_plans = get_business_plans(user_id)
    shopify_plans = get_shopify_plans(user_id)
    shopify_products = get_shopify_products(user_id)
    shopify_collections = get_shopify_collections(user_id)
    shopify_pages = get_shopify_pages(user_id)
    canva_branding_packages = get_canva_branding_packages(user_id)
    canva_design_briefs = get_canva_design_briefs(user_id)
    canva_designs = get_canva_designs(user_id)
    build_quotes = get_build_quotes(user_id)

    return {
        "answers_by_step": answers_by_step,
        "latest_business_plan": business_plans[0] if business_plans else None,
        "latest_shopify_plan": shopify_plans[0] if shopify_plans else None,
        "shopify_products": shopify_products,
        "shopify_collections": shopify_collections,
        "shopify_pages": shopify_pages,
        "latest_canva_branding_package": (
            canva_branding_packages[0]
            if canva_branding_packages
            else None
        ),
        "latest_canva_design_brief": (
            canva_design_briefs[0]
            if canva_design_briefs
            else None
        ),
        "canva_designs": canva_designs,
        "latest_build_quote": build_quotes[0] if build_quotes else None
    }


def send_launch_package_email_once(user_id):
    if session.get("launch_package_email_sent"):
        return

    email = get_user_email(user_id)

    if email:
        send_email(
            email,
            "Your BusinessBuilder AI launch package is ready",
            (
                "Your final launch package is ready to review. "
                "Open BusinessBuilder AI to check your strategy, draft assets, "
                "launch checklist, and PDF export."
            )
        )

    session["launch_package_email_sent"] = True


@app.route("/launch_package")
def launch_package():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_package_at_least(user_id, "Pro"):
        return package_access_redirect("Pro")

    if usage_limit_reached(user_id, "launch_package"):
        return usage_limit_redirect("launch_package")

    send_launch_package_email_once(user_id)

    launch_data = get_launch_package_data(user_id)
    launch_data["premium_build"] = user_package_at_least(user_id, "Premium Build")
    response = render_template("launch_package.html", **launch_data)

    log_usage(user_id, "launch_package")

    return response


@app.route("/download_launch_package")
def download_launch_package():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_package_at_least(user_id, "Premium Build"):
        return package_access_redirect("Premium Build")

    if usage_limit_reached(user_id, "pdf_export"):
        return usage_limit_redirect("pdf_export")

    send_launch_package_email_once(user_id)

    response = send_file(
        create_launch_package_pdf(get_launch_package_data(user_id)),
        mimetype="application/pdf",
        as_attachment=True,
        download_name="businessbuilder-ai-launch-package.pdf"
    )

    log_usage(user_id, "pdf_export")

    return response


@app.route("/business_plan/<int:plan_id>")
def business_plan(plan_id):
    if "user_id" not in session:
        return redirect("/login")

    plan = get_business_plan(
        session["user_id"],
        plan_id
    )

    if not plan:
        return redirect("/dashboard")

    return render_template(
        "business_plan.html",
        plan=plan
    )


@app.route("/download_business_plan/<int:plan_id>")
def download_business_plan(plan_id):
    if "user_id" not in session:
        return redirect("/login")

    plan = get_business_plan(
        session["user_id"],
        plan_id
    )

    if not plan:
        return redirect("/dashboard")

    if not user_is_paid(session["user_id"]):
        return package_access_redirect("Starter")

    if usage_limit_reached(session["user_id"], "pdf_export"):
        return usage_limit_redirect("pdf_export")

    filename = re.sub(r"[_-]+", "-", secure_filename(plan[1])).strip("-")
    filename = filename or "business-plan"

    response = send_file(
        create_business_plan_pdf(plan[1], plan[2]),
        mimetype="application/pdf",
        as_attachment=True,
        download_name=f"{filename}.pdf"
    )

    log_usage(session["user_id"], "pdf_export")

    return response


@app.route("/shopify_plan/<int:plan_id>")
def shopify_plan(plan_id):
    if "user_id" not in session:
        return redirect("/login")

    plan = get_shopify_plan(
        session["user_id"],
        plan_id
    )

    if not plan:
        return redirect("/dashboard")

    return render_template(
        "shopify_plan.html",
        plan=plan
    )


@app.route("/canva_branding/<int:package_id>")
def canva_branding(package_id):
    if "user_id" not in session:
        return redirect("/login")

    package = get_canva_branding_package(
        session["user_id"],
        package_id
    )

    if not package:
        return redirect("/dashboard")

    return render_template(
        "canva_branding.html",
        package=package
    )


@app.route("/canva_design_brief/<int:brief_id>")
def canva_design_brief(brief_id):
    if "user_id" not in session:
        return redirect("/login")

    brief = get_canva_design_brief(
        session["user_id"],
        brief_id
    )

    if not brief:
        return redirect("/dashboard")

    return render_template(
        "canva_design_brief.html",
        brief=brief
    )


@app.route("/build_quote/<int:quote_id>")
def build_quote(quote_id):
    if "user_id" not in session:
        return redirect("/login")

    quote = get_build_quote(
        session["user_id"],
        quote_id
    )

    if not quote:
        return redirect("/dashboard")

    return render_template(
        "build_quote.html",
        quote=quote
    )


@app.route("/shopify_settings", methods=["GET", "POST"])
def shopify_settings():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    if not user_package_at_least(user_id, "Pro"):
        return package_access_redirect("Pro")

    connection = get_shopify_connection(user_id)
    connection_summary = None
    message = None

    if connection:
        connection_summary = {
            "shop_domain": connection[1],
            "status": connection[3]
        }

    if request.method == "POST":
        shop_domain = normalize_shopify_domain(
            request.form.get("shop_domain", "")
        )
        access_token = request.form.get("access_token", "").strip()

        if not shop_domain:
            message = "Enter a valid Shopify domain such as your-store.myshopify.com."
        elif not access_token:
            message = "Enter an Admin API access token to save and test the connection."
        else:
            connected, message = test_shopify_connection(
                shop_domain,
                access_token
            )
            status = "connected" if connected else "not_connected"

            save_shopify_connection(
                user_id,
                shop_domain,
                access_token,
                status
            )

            connection_summary = {
                "shop_domain": shop_domain,
                "status": status
            }

    return render_template(
        "shopify_settings.html",
        connection=connection_summary,
        message=message
    )


@app.route("/shopify_build_summary")
def shopify_build_summary():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    if not user_package_at_least(user_id, "Pro"):
        return package_access_redirect("Pro")

    return render_template(
        "shopify_build_summary.html",
        shopify_products=get_shopify_products(user_id),
        shopify_collections=get_shopify_collections(user_id),
        shopify_pages=get_shopify_pages(user_id)
    )


@app.route("/canva_settings")
def canva_settings():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    if not user_package_at_least(user_id, "Pro"):
        return package_access_redirect("Pro")

    connection = get_canva_connection(user_id)
    message = {
        "connected": "Canva connected successfully.",
        "configuration": "Canva OAuth is not configured yet. Add the required environment variables.",
        "state": "Canva connection could not be verified. Please start the connection again.",
        "denied": "Canva authorization was not completed.",
        "token": "Canva could not complete the token exchange. Please try again."
    }.get(request.args.get("canva_oauth"))

    connection_summary = None

    if connection:
        connection_summary = {
            "status": connection[2]
        }

    return render_template(
        "canva_settings.html",
        connection=connection_summary,
        message=message
    )


@app.route("/settings")
def settings():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]
    latest_payment = get_latest_payment(user_id)

    return render_template(
        "settings.html",
        email=get_user_email(user_id),
        current_package=get_user_package(user_id) or "No active package",
        current_plan_details=PLAN_DETAILS.get(get_user_package(user_id), {}),
        latest_payment=latest_payment,
        usage_summary=get_usage_summary(user_id),
        settings=get_user_settings(user_id),
        theme_options=THEME_OPTIONS,
        shopify_connection=get_shopify_connection(user_id),
        canva_connection=get_canva_connection(user_id),
        connected_apps=get_connected_app_summaries(user_id),
        pro_package=user_package_at_least(user_id, "Pro"),
        premium_build=user_package_at_least(user_id, "Premium Build")
    )


@app.route("/update_settings", methods=["POST"])
def update_settings():
    if "user_id" not in session:
        if request.is_json:
            return jsonify({"error": "Login required."}), 401

        return redirect("/login")

    theme = normalize_theme(
        (request.get_json(silent=True) or {}).get("theme")
        if request.is_json
        else request.form.get("theme")
    )
    save_user_settings(session["user_id"], theme)

    if request.is_json:
        return jsonify({
            "status": "saved",
            "theme": theme
        })

    return redirect("/settings?settings_notice=saved")


@app.route("/connect_canva")
def connect_canva():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    if not user_package_at_least(user_id, "Pro"):
        return package_access_redirect("Pro")

    client_id = os.getenv("CANVA_CLIENT_ID")
    redirect_uri = os.getenv("CANVA_REDIRECT_URI")

    if not client_id or not redirect_uri:
        return redirect("/canva_settings?canva_oauth=configuration")

    code_verifier = secrets.token_urlsafe(64)
    code_challenge = base64.urlsafe_b64encode(
        hashlib.sha256(code_verifier.encode("ascii")).digest()
    ).decode("ascii").rstrip("=")
    state = secrets.token_urlsafe(32)

    session["canva_oauth_state"] = state
    save_canva_oauth_session(
        user_id,
        state,
        code_verifier
    )

    authorization_url = CANVA_AUTHORIZATION_URL + "?" + urllib.parse.urlencode({
        "code_challenge": code_challenge,
        "code_challenge_method": "S256",
        "scope": CANVA_SCOPES,
        "response_type": "code",
        "client_id": client_id,
        "state": state,
        "redirect_uri": redirect_uri
    })

    return redirect(authorization_url)


@app.route("/canva_callback")
def canva_callback():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    if not user_package_at_least(user_id, "Pro"):
        return package_access_redirect("Pro")

    expected_state = session.pop("canva_oauth_state", None)
    received_state = request.args.get("state", "")

    if (
        not expected_state
        or not received_state
        or not secrets.compare_digest(expected_state, received_state)
    ):
        return redirect("/canva_settings?canva_oauth=state")

    code_verifier = pop_canva_oauth_verifier(
        user_id,
        received_state
    )

    if not code_verifier:
        return redirect("/canva_settings?canva_oauth=state")

    if request.args.get("error"):
        return redirect("/canva_settings?canva_oauth=denied")

    code = request.args.get("code")
    client_id = os.getenv("CANVA_CLIENT_ID")
    client_secret = os.getenv("CANVA_CLIENT_SECRET")
    redirect_uri = os.getenv("CANVA_REDIRECT_URI")

    if not code or not client_id or not client_secret or not redirect_uri:
        return redirect("/canva_settings?canva_oauth=configuration")

    try:
        response = requests.post(
            CANVA_TOKEN_URL,
            auth=(client_id, client_secret),
            headers={
                "Content-Type": "application/x-www-form-urlencoded"
            },
            data={
                "grant_type": "authorization_code",
                "code_verifier": code_verifier,
                "code": code,
                "redirect_uri": redirect_uri
            },
            timeout=10
        )
        token_data = response.json()
    except requests.RequestException:
        return redirect("/canva_settings?canva_oauth=token")
    except ValueError:
        return redirect("/canva_settings?canva_oauth=token")

    access_token = token_data.get("access_token")

    if response.status_code != 200 or not access_token:
        return redirect("/canva_settings?canva_oauth=token")

    refresh_token = token_data.get("refresh_token", "")
    connected_email = "Connected Canva User"

    try:
        profile_response = requests.get(
            CANVA_PROFILE_URL,
            headers={
                "Authorization": f"Bearer {access_token}"
            },
            timeout=10
        )
        profile_data = profile_response.json()

        if profile_response.status_code == 200:
            connected_email = (
                profile_data.get("profile", {}).get("display_name")
                or connected_email
            )
    except (requests.RequestException, ValueError):
        pass

    save_canva_connection(
        user_id,
        "connected",
        connected_email,
        access_token,
        refresh_token
    )

    return redirect("/canva_settings?canva_oauth=connected")


@app.route("/health_check")
def health_check():
    if "user_id" not in session:
        return jsonify({"error": "Authentication required."}), 401

    database_connected = False
    conn = None

    try:
        conn = db()
        cur = conn.cursor()
        cur.execute(sql("SELECT 1"))
        database_connected = cur.fetchone() is not None
    except (sqlite3.Error, psycopg2.Error):
        database_connected = False
    finally:
        if conn:
            conn.close()

    return jsonify({
        "database_connected": database_connected,
        "openai_configured": bool(os.getenv("OPENAI_API_KEY")),
        "paystack_configured": bool(os.getenv("PAYSTACK_SECRET_KEY")),
        "database_type": "postgres" if using_postgres() else "sqlite"
    })


@app.route("/app_connection_agent")
def app_connection_agent():
    if "user_id" not in session:
        return redirect("/login")
    user_id = session["user_id"]
    if not user_package_at_least(user_id, "Pro"):
        return package_access_redirect("Pro")
    drafts = get_app_action_drafts(user_id)
    return render_template(
        "app_connection_agent.html",
        app_cards=get_app_catalog_cards(user_id),
        connected_apps=get_connected_app_summaries(user_id),
        recommendations=get_app_recommendations(user_id),
        pending_drafts=[draft for draft in drafts if not draft[7]],
        all_drafts=drafts,
        premium_actions=user_package_at_least(user_id, "Premium Build"),
        current_package=get_user_package(user_id) or "Starter"
    )


@app.route("/recommend_apps")
def recommend_apps():
    if "user_id" not in session:
        return redirect("/login")
    user_id = session["user_id"]
    if not user_package_at_least(user_id, "Pro"):
        return package_access_redirect("Pro")
    return render_template(
        "app_recommendations.html",
        onboarding=get_user_onboarding(user_id),
        recommendations=get_app_recommendations(user_id)
    )


@app.route("/generate_app_recommendations", methods=["GET", "POST"])
def generate_app_recommendations():
    if "user_id" not in session:
        return redirect("/login")
    if not user_package_at_least(session["user_id"], "Pro"):
        return package_access_redirect("Pro")
    if request.method == "GET":
        return redirect("/recommend_apps")

    user_id = session["user_id"]
    data = {
        "business_type": request.form.get("business_type", "").strip(),
        "country": request.form.get("country", "").strip(),
        "budget": request.form.get("budget", "").strip(),
        "goal": request.form.get("goal", "").strip(),
        "selling_platform": request.form.get("selling_platform", "").strip(),
        "skill_level": request.form.get("skill_level", "").strip()
    }
    if not data["business_type"] or not data["country"] or not data["budget"]:
        return redirect("/recommend_apps?app_error=missing")

    prompt = f"""
Recommend approved business apps for this BusinessBuilder AI user.

Safety rules:
- Do not claim live prices. Tell the user to verify provider pricing, limits, availability, and country support.
- Do not connect accounts, create accounts, spend ad money, publish ads, send mass emails, or publish website/store changes.
- Explain what requires explicit user approval and what should stay manual.
- Recommend official OAuth or API-key flows only. Never ask for passwords.

Budget logic to apply:
- Very low: Canva free/low-cost, Brevo, Shopify Email for Shopify users, WhatsApp Business guidance, and organic social content.
- Low: Brevo, Mailchimp, Canva, Shopify, and Webflow guidance only when a website is the goal.
- Medium: Klaviyo for ecommerce, small controlled Meta Ads planning, Mailchimp/Brevo, and Canva Pro only if chosen.
- High: Klaviyo, Meta Ads, Google Ads, Webflow, and advanced Zapier/Make automation planning.

Include:
1. Best app/platform
2. Cheapest option
3. Best beginner option
4. Best professional option
5. What each recommended platform can do
6. What BusinessBuilder AI can draft through it
7. What requires user approval
8. What may cost money
9. What should stay manual for safety
10. Recommended next step

Business type: {data['business_type']}
Country/market: {data['country']}
Monthly budget: {data['budget']}
Main goal: {data['goal']}
Selling platform: {data['selling_platform']}
Skill level: {data['skill_level']}
{build_project_context(user_id)}
"""
    try:
        response = safe_openai_chat_completion(
            model="gpt-4.1-mini",
            messages=[{"role": "system", "content": SYSTEM_PROMPT}, {"role": "user", "content": prompt}]
        )
    except Exception:
        return redirect("/recommend_apps?app_error=ai_response")
    save_app_recommendation(user_id, data, response.choices[0].message.content.strip())
    return redirect("/recommend_apps?app_notice=created")


@app.route("/app_connection_guide/<platform>")
def app_connection_guide(platform):
    if "user_id" not in session:
        return redirect("/login")
    if not user_package_at_least(session["user_id"], "Pro"):
        return package_access_redirect("Pro")
    platform = normalize_app_platform(platform)
    if platform not in APP_CATALOG:
        return redirect("/app_connection_agent?app_error=unsupported")
    return render_template(
        "app_connection_guide.html",
        platform=platform,
        app_info=APP_CATALOG[platform],
        status=get_app_connection_status(session["user_id"], platform),
        connection=bool(get_app_connection(session["user_id"], platform)) if APP_CATALOG[platform]["connection_mode"] == "manual" else False,
        premium_actions=user_package_at_least(session["user_id"], "Premium Build")
    )


@app.route("/connect_app/<platform>", methods=["GET", "POST"])
def connect_app(platform):
    if "user_id" not in session:
        return redirect("/login")
    platform = normalize_app_platform(platform)
    if platform not in APP_CATALOG:
        return redirect("/app_connection_agent?app_error=unsupported")
    mode = APP_CATALOG[platform]["connection_mode"]
    if mode == "canva":
        return redirect("/canva_settings")
    if mode == "shopify":
        return redirect("/shopify_settings")
    if mode == "guidance":
        return redirect(f"/app_connection_guide/{platform}")
    if not user_package_at_least(session["user_id"], "Premium Build"):
        return package_access_redirect("Premium Build", f"/app_connection_guide/{platform}")
    if request.method == "GET":
        return redirect(f"/app_connection_guide/{platform}")

    api_key = request.form.get("api_key", "").strip()
    account_name = request.form.get("account_name", "").strip()
    permission_confirmed = request.form.get("permission_confirmed") == "yes"
    if not api_key or not permission_confirmed:
        return redirect(f"/app_connection_guide/{platform}?connection_error=missing")
    try:
        save_app_connection(session["user_id"], platform, APP_CATALOG[platform]["category"], {
            "connection_status": "configured",
            "api_key": api_key,
            "account_name": account_name,
            "scopes": "User-supplied API key; permissions controlled by provider"
        })
    except ExternalServiceError:
        return redirect(f"/app_connection_guide/{platform}?connection_error=encryption")
    return redirect(f"/app_connection_guide/{platform}?connection_notice=saved")


@app.route("/disconnect_app/<platform>", methods=["GET", "POST"])
def disconnect_app(platform):
    if "user_id" not in session:
        return redirect("/login")
    platform = normalize_app_platform(platform)
    if platform not in APP_CATALOG:
        return redirect("/app_connection_agent?app_error=unsupported")
    if APP_CATALOG[platform]["connection_mode"] != "manual":
        return redirect(f"/app_connection_guide/{platform}")
    if request.method == "GET":
        return redirect(f"/app_connection_guide/{platform}?confirm_disconnect=1")
    disconnect_app_connection(session["user_id"], platform)
    return redirect(f"/app_connection_guide/{platform}?connection_notice=disconnected")


def build_app_action_prompt(platform, action_type, user_id):
    app_name = APP_CATALOG[platform]["name"]
    instructions = {
        "email_campaign": "Create strategy, subject lines, preview text, full email body, CTA, audience segment, sending-time test, follow-up, and platform setup checklist.",
        "social_ad_draft": "Create audience ideas, campaign objective, ad-copy options, headlines, CTA options, creative brief, tracking checklist, controlled-budget planning notes, and risks. Do not publish or spend.",
        "canva_design_brief": "Create a Canva-ready design brief with dimensions, layout, copy, colors, imagery, accessibility, and export guidance.",
        "shopify_product_draft": "Return JSON only with a top-level products array containing 3 objects. Each needs title, description, suggested_price, category, seo_title, and meta_description. Products must remain drafts.",
        "webflow_page_draft": "Create a Webflow page/CMS draft with page goal, sections, copy, CMS fields, SEO title, meta description, and manual build checklist.",
        "website_copy": "Create reviewable website/store copy with hero, benefits, CTA, trust section, FAQ, SEO title, and meta description.",
        "automation_plan": "Create a trigger/action automation plan, field mapping, failure handling, privacy risks, test plan, and manual enablement checklist. Do not enable it."
    }[action_type]
    return f"""
Create a {APP_ACTION_TYPES[action_type]} for {app_name}.
{instructions}

Safety rules:
- Create a draft only. The user must review, edit, approve, and separately confirm any supported apply action.
- Never request passwords or expose tokens.
- Never spend ad money, publish ads, send mass emails, or publish website/store changes.
- Clearly state provider costs are not live and must be verified.
- Separate draft content from manual setup and approval steps.

{build_project_context(user_id)}
"""


@app.route("/create_app_action_draft/<platform>/<action_type>", methods=["GET", "POST"])
def create_app_action_draft(platform, action_type):
    if "user_id" not in session:
        return redirect("/login")
    if not user_package_at_least(session["user_id"], "Premium Build"):
        return package_access_redirect("Premium Build", "/app_connection_agent")
    if usage_limit_reached(session["user_id"], "app_action_draft"):
        return usage_limit_redirect("app_action_draft", "/app_connection_agent")
    platform = normalize_app_platform(platform)
    action_type = (action_type or "").strip().lower()
    if platform not in APP_CATALOG or action_type not in APP_ACTION_TYPES or action_type not in APP_CATALOG[platform]["draft_actions"]:
        return redirect("/app_connection_agent?draft_error=unsupported")
    kwargs = {
        "model": "gpt-4.1-mini",
        "messages": [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": build_app_action_prompt(platform, action_type, session["user_id"])}
        ]
    }
    if action_type == "shopify_product_draft":
        kwargs["response_format"] = {"type": "json_object"}
    try:
        response = safe_openai_chat_completion(**kwargs)
    except Exception:
        return redirect("/app_connection_agent?draft_error=ai_response")
    title = f"{APP_CATALOG[platform]['name']} — {APP_ACTION_TYPES[action_type]}"
    draft_id = create_app_action_draft_record(
        session["user_id"], platform, action_type, title,
        response.choices[0].message.content.strip()
    )
    log_usage(session["user_id"], "app_action_draft")
    return redirect(f"/review_app_action_draft/{draft_id}")


@app.route("/review_app_action_draft/<int:draft_id>")
def review_app_action_draft(draft_id):
    if "user_id" not in session:
        return redirect("/login")
    draft = get_app_action_draft(session["user_id"], draft_id)
    if not draft:
        return redirect("/app_connection_agent")
    return render_template(
        "app_action_draft.html", draft=draft,
        app_info=APP_CATALOG.get(draft[1], {}),
        risk=get_app_action_risk(draft[2]),
        show_confirmation=False
    )


@app.route("/edit_app_action_draft/<int:draft_id>", methods=["POST"])
def edit_app_action_draft(draft_id):
    if "user_id" not in session:
        return redirect("/login")
    content = request.form.get("draft_content", "").strip()
    if not get_app_action_draft(session["user_id"], draft_id) or not content:
        return redirect(f"/review_app_action_draft/{draft_id}?draft_error=empty")
    update_app_action_draft_content(session["user_id"], draft_id, content)
    return redirect(f"/review_app_action_draft/{draft_id}?draft_notice=updated")


@app.route("/improve_app_action_draft/<int:draft_id>", methods=["POST"])
def improve_app_action_draft(draft_id):
    if "user_id" not in session:
        return redirect("/login")
    user_id = session["user_id"]
    draft = get_app_action_draft(user_id, draft_id)
    feedback = request.form.get("feedback", "").strip()
    if not draft or not feedback:
        return redirect(f"/review_app_action_draft/{draft_id}?draft_error=feedback")
    prompt = f"Revise this draft using the feedback. Keep it draft-only and preserve all no-publish, no-send, no-spend safety rules.\nFeedback: {feedback}\nDraft:\n{draft[4]}"
    try:
        response = safe_openai_chat_completion(model="gpt-4.1-mini", messages=[{"role": "system", "content": SYSTEM_PROMPT}, {"role": "user", "content": prompt}])
    except Exception:
        return redirect(f"/review_app_action_draft/{draft_id}?draft_error=ai_response")
    update_app_action_draft_content(user_id, draft_id, response.choices[0].message.content.strip())
    return redirect(f"/review_app_action_draft/{draft_id}?draft_notice=updated")


@app.route("/approve_app_action_draft/<int:draft_id>", methods=["POST"])
def approve_app_action_draft(draft_id):
    if "user_id" not in session:
        return redirect("/login")
    user_id = session["user_id"]
    if not get_app_action_draft(user_id, draft_id):
        return redirect("/app_connection_agent")
    update_app_action_draft_status(user_id, draft_id, "Approved", True, False)
    return redirect(f"/review_app_action_draft/{draft_id}?draft_notice=approved")


@app.route("/apply_app_action_draft/<int:draft_id>", methods=["GET", "POST"])
def apply_app_action_draft(draft_id):
    if "user_id" not in session:
        return redirect("/login")
    user_id = session["user_id"]
    if not user_package_at_least(user_id, "Premium Build"):
        return package_access_redirect("Premium Build", f"/review_app_action_draft/{draft_id}")
    draft = get_app_action_draft(user_id, draft_id)
    if not draft:
        return redirect("/app_connection_agent")
    if not draft[6]:
        return redirect(f"/review_app_action_draft/{draft_id}?draft_error=not_approved")
    risk = get_app_action_risk(draft[2])
    if request.method == "GET":
        return render_template("app_action_draft.html", draft=draft, app_info=APP_CATALOG.get(draft[1], {}), risk=risk, show_confirmation=True)
    if request.form.get("confirm_apply") != "yes":
        return redirect(f"/apply_app_action_draft/{draft_id}?draft_error=confirmation")

    applied = True
    message = "Approved draft saved inside BusinessBuilder AI. No external action was published."
    if draft[2] == "shopify_product_draft":
        task = (None, user_id, "products", draft[3], draft[4])
        applied, message = apply_store_agent_products(user_id, task)
    elif draft[2] == "canva_design_brief":
        save_canva_design_brief(user_id, draft[3], draft[4])
        message = "Canva-ready design brief saved. No design was published."
    elif draft[2] == "email_campaign":
        save_email_campaign(user_id, {
            "business_name": APP_CATALOG[draft[1]]["name"], "business_type": "App Connection Agent",
            "target_customer": "", "campaign_goal": "Draft campaign", "email_platform": APP_CATALOG[draft[1]]["name"],
            "offer": "", "tone": ""
        }, draft[4])
        message = "Email campaign draft saved. No email was sent."

    if not applied:
        return redirect(f"/review_app_action_draft/{draft_id}?draft_error=apply_failed&apply_message={urllib.parse.quote(message)}")
    update_app_action_draft_status(user_id, draft_id, "Applied Safely", True, True)
    return redirect(f"/review_app_action_draft/{draft_id}?draft_notice=applied&apply_message={urllib.parse.quote(message)}")


@app.route("/delete_app_action_draft/<int:draft_id>", methods=["POST"])
def delete_app_action_draft(draft_id):
    if "user_id" not in session:
        return redirect("/login")
    delete_app_action_draft_record(session["user_id"], draft_id)
    return redirect("/app_connection_agent?draft_notice=deleted")


@app.route("/email_marketing")
def email_marketing():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]
    if not user_package_at_least(user_id, "Pro"):
        return package_access_redirect("Pro")
    return render_template(
        "email_marketing.html",
        onboarding=get_user_onboarding(user_id),
        campaigns=get_email_campaigns(user_id),
        current_package=get_user_package(user_id) or "Starter",
        usage_summary=get_usage_summary(user_id)
    )


@app.route("/generate_email_campaign", methods=["GET", "POST"])
def generate_email_campaign():
    if "user_id" not in session:
        return redirect("/login")

    if not user_package_at_least(session["user_id"], "Pro"):
        return package_access_redirect("Pro")
    if request.method == "GET":
        return redirect("/email_marketing")

    user_id = session["user_id"]
    if usage_limit_reached(user_id, "email_campaign"):
        return usage_limit_redirect("email_campaign", "/email_marketing")
    data = {
        "business_name": request.form.get("business_name", "").strip(),
        "business_type": request.form.get("business_type", "").strip(),
        "target_customer": request.form.get("target_customer", "").strip(),
        "campaign_goal": request.form.get("campaign_goal", "").strip(),
        "email_platform": request.form.get("email_platform", "").strip(),
        "offer": request.form.get("offer", "").strip(),
        "tone": request.form.get("tone", "").strip()
    }

    if not data["business_name"] or not data["business_type"] or not data["campaign_goal"]:
        return redirect("/email_marketing?email_error=missing")

    prompt = f"""
Create a practical email marketing campaign draft for this BusinessBuilder AI user.

Safety and approval rules:
- Generate strategy, copy drafts, and setup guidance only.
- The user must review the audience, claims, links, offer terms, consent, unsubscribe handling, and final copy before sending.
- Do not send mass emails automatically. An official integration and explicit user approval would be required for any sending action.
- Do not create Brevo, Mailchimp, Klaviyo, Shopify, or other accounts.
- Do not request passwords, API keys, payment details, or private credentials.
- Remind the user to follow consent, anti-spam, privacy, and unsubscribe rules that apply in their market.

Include these clearly labelled sections:
1. Email campaign strategy
2. At least five subject line options
3. At least three preview text options
4. Full email body
5. Call-to-action button text options
6. Customer segment suggestion
7. Recommended sending time, described as a testable starting point rather than a guarantee
8. Follow-up email idea
9. Platform recommendation based on the user's needs
10. Setup checklist for the selected email platform

Platform guidance:
- Brevo is beginner-friendly and useful for email/SMS marketing.
- Mailchimp is beginner-friendly for email marketing.
- Klaviyo is stronger for ecommerce automation.
- Shopify Email is useful when the user mainly operates in Shopify.

Business name: {data["business_name"]}
Business type: {data["business_type"]}
Target customer: {data["target_customer"]}
Campaign goal: {data["campaign_goal"]}
Preferred platform: {data["email_platform"]}
Offer or discount: {data["offer"] or "None provided"}
Tone: {data["tone"]}

{build_project_context(user_id)}
"""

    try:
        response = safe_openai_chat_completion(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": prompt}
            ]
        )
    except Exception:
        return redirect("/email_marketing?email_error=ai_response")

    campaign_id = save_email_campaign(
        user_id,
        data,
        response.choices[0].message.content.strip()
    )
    log_usage(user_id, "email_campaign")
    return redirect(f"/email_campaign/{campaign_id}?email_notice=created")


@app.route("/email_campaign/<int:campaign_id>")
def email_campaign(campaign_id):
    if "user_id" not in session:
        return redirect("/login")

    campaign = get_email_campaign(session["user_id"], campaign_id)
    if not campaign:
        return redirect("/email_marketing")

    return render_template("email_campaign.html", campaign=campaign)


@app.route("/domain_helper")
def domain_helper():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]
    return render_template(
        "domain_helper.html",
        onboarding=get_user_onboarding(user_id),
        guides=get_domain_guides(user_id),
        current_package=get_user_package(user_id) or "Starter",
        usage_summary=get_usage_summary(user_id)
    )


@app.route("/generate_domain_advice", methods=["GET", "POST"])
def generate_domain_advice():
    if "user_id" not in session:
        return redirect("/login")

    if request.method == "GET":
        return redirect("/domain_helper")

    user_id = session["user_id"]
    if usage_limit_reached(user_id, "domain_guide"):
        return usage_limit_redirect("domain_guide", "/domain_helper")
    data = {
        "business_name": request.form.get("business_name", "").strip(),
        "preferred_domain": request.form.get("preferred_domain", "").strip(),
        "country": request.form.get("country", "").strip(),
        "budget": request.form.get("budget", "").strip(),
        "provider_preference": request.form.get("provider_preference", "").strip(),
        "website_platform": request.form.get("website_platform", "").strip(),
        "owns_domain": request.form.get("owns_domain", "").strip()
    }

    if not data["business_name"] or not data["country"]:
        return redirect("/domain_helper?domain_error=missing")

    prompt = f"""
Create beginner-friendly domain selection and DNS guidance for this BusinessBuilder AI user.

Safety and accuracy rules:
- Do not claim live domain availability or live/current prices. No registrar availability API is configured.
- Do not buy, register, renew, transfer, or charge for a domain automatically.
- Do not create GoDaddy, Namecheap, Cloudflare, Domains.co.za, Shopify, or other accounts.
- Never ask for card details, banking passwords, registrar passwords, API keys, or private credentials.
- Tell the user to verify final availability, first-year price, renewal price, taxes, transfer rules, privacy fees, and terms directly with the provider before paying.
- Explain that the user reviews and approves every real purchase or DNS change.

Include these clearly labelled sections:
1. Domain name suggestions
2. Best extensions, including .com, .co.za, .store, .shop, .site, and other relevant choices
3. Budget-friendly provider suggestions
4. Professional provider suggestions
5. Pros and cons of GoDaddy, Namecheap, Cloudflare, Domains.co.za, and Shopify Domains where relevant
6. What to check before buying
7. General estimated-cost guidance without claiming live prices
8. Beginner DNS connection guide
9. Shopify domain connection guide
10. BusinessBuilder AI custom-domain connection guide that says to confirm the exact hosting target/records in BusinessBuilder AI support before changing DNS
11. Email/domain credibility advice, including professional sender addresses and email authentication concepts
12. Billing, renewal, phishing, account-security, and accidental-expiry warning
13. Step-by-step purchase and connection checklist

Explain these terms in plain language:
- domain
- DNS
- CNAME record
- A record
- nameservers
- SSL certificate
- redirecting the root domain to www
- DNS propagation delay

Business name: {data["business_name"]}
Desired domain: {data["preferred_domain"] or "No preference yet"}
Country/market: {data["country"]}
Budget level: {data["budget"]}
Provider preference: {data["provider_preference"]}
Website/store platform: {data["website_platform"]}
Already owns a domain: {data["owns_domain"]}

{build_project_context(user_id)}
"""

    try:
        response = safe_openai_chat_completion(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": prompt}
            ]
        )
    except Exception:
        return redirect("/domain_helper?domain_error=ai_response")

    guide_id = save_domain_guide(
        user_id,
        data,
        response.choices[0].message.content.strip()
    )
    log_usage(user_id, "domain_guide")
    return redirect(f"/domain_guide/{guide_id}?domain_notice=created")


@app.route("/domain_guide/<int:guide_id>")
def domain_guide(guide_id):
    if "user_id" not in session:
        return redirect("/login")

    guide = get_domain_guide(session["user_id"], guide_id)
    if not guide:
        return redirect("/domain_helper")

    return render_template("domain_guide.html", guide=guide)


@app.route("/launch_readiness")
def launch_readiness():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]
    readiness = get_launch_readiness(user_id)

    return render_template(
        "launch_readiness.html",
        readiness=readiness,
        active_project=get_active_business_project(user_id),
        current_package=get_user_package(user_id),
        premium_build=user_package_at_least(user_id, "Premium Build")
    )


@app.route("/examples")
def examples():
    return render_template("examples.html")


@app.route("/onboarding")
def onboarding():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    return render_template(
        "onboarding.html",
        onboarding=get_user_onboarding(user_id),
        active_project=get_active_business_project(user_id)
    )


@app.route("/save_onboarding", methods=["POST"])
def save_onboarding():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]
    data = {
        "business_idea": request.form.get("business_idea", "").strip(),
        "country": request.form.get("country", "").strip(),
        "budget": request.form.get("budget", "").strip(),
        "has_shopify": request.form.get("has_shopify", "").strip(),
        "has_canva": request.form.get("has_canva", "").strip(),
        "product_type": request.form.get("product_type", "").strip(),
        "target_customer": request.form.get("target_customer", "").strip(),
        "business_goal": request.form.get("business_goal", "").strip()
    }

    if not data["business_idea"]:
        return redirect("/onboarding?onboarding_error=business_idea")

    save_user_onboarding(user_id, data)

    if not get_active_business_project(user_id):
        create_business_project_record(
            user_id,
            data["business_idea"][:80] or "My Business",
            data["business_idea"],
            data["target_customer"],
            data["country"],
            data["budget"],
            f"Goal: {data['business_goal']}. Product type: {data['product_type']}.",
            True
        )

    email = get_user_email(user_id)
    if email:
        send_email_notification(
            email,
            "Your BusinessBuilder AI onboarding is complete",
            "Your onboarding answers were saved. Open your dashboard to continue the recommended next step."
        )

    return redirect("/dashboard?onboarding=complete")


@app.route("/support")
def support():
    email = ""

    if "user_id" in session:
        email = get_user_email(session["user_id"]) or ""

    return render_template(
        "support.html",
        email=email,
        submitted=request.args.get("submitted") == "1"
    )


@app.route("/submit_support", methods=["POST"])
def submit_support():
    user_id = session.get("user_id")
    email = request.form.get("email", "").strip().lower()
    category = request.form.get("category", "").strip()
    subject = request.form.get("subject", "").strip()
    message = request.form.get("message", "").strip()

    if not is_valid_email(email) or not category or not subject or not message:
        return redirect("/support?support_error=missing")

    ticket_id = save_support_ticket(user_id, email, category, subject, message)
    admin_email = os.getenv("ADMIN_EMAIL", "").strip()

    if admin_email:
        send_email_notification(
            admin_email,
            f"BusinessBuilder AI support request #{ticket_id}",
            (
                f"Category: {category}\n"
                f"From: {email}\n"
                f"Subject: {subject}\n\n"
                f"{message}"
            )
        )

    send_email_notification(
        email,
        "BusinessBuilder AI support request received",
        "Your support request was received. The team will review it and respond as soon as possible."
    )

    return redirect("/support?submitted=1")


@app.route("/pricing_advisor")
def pricing_advisor():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    return render_template(
        "pricing_advisor.html",
        onboarding=get_user_onboarding(user_id),
        advice_list=get_pricing_advice_list(user_id),
        current_package=get_user_package(user_id) or "Starter",
        usage_summary=get_usage_summary(user_id)
    )


@app.route("/generate_pricing_advice", methods=["GET", "POST"])
def generate_pricing_advice():
    if "user_id" not in session:
        return redirect("/login")

    if request.method == "GET":
        return redirect("/pricing_advisor")

    user_id = session["user_id"]
    if usage_limit_reached(user_id, "pricing_advice"):
        return usage_limit_redirect("pricing_advice", "/pricing_advisor")
    data = {
        "business_idea": request.form.get("business_idea", "").strip(),
        "product_type": request.form.get("product_type", "").strip(),
        "target_customer": request.form.get("target_customer", "").strip(),
        "country": request.form.get("country", "").strip(),
        "cost_price": request.form.get("cost_price", "").strip(),
        "desired_profit": request.form.get("desired_profit", "").strip(),
        "competitor_price": request.form.get("competitor_price", "").strip(),
        "budget": request.form.get("budget", "").strip(),
        "pricing_style": request.form.get("pricing_style", "").strip()
    }

    if not data["business_idea"]:
        return redirect("/pricing_advisor?pricing_error=missing")

    prompt = f"""
Create practical pricing advice for this BusinessBuilder AI user.

Safety and accuracy rules:
- Use estimates and explain assumptions.
- Do not claim live competitor prices or scrape websites.
- Tell the user to verify cost price, payment fees, taxes, packaging, delivery, refunds, platform fees, and competitor positioning.
- Do not request card numbers, banking passwords, ID numbers, or private payment credentials.

Include:
1. Suggested selling price range
2. Low, standard, and premium pricing options
3. Profit margin explanation
4. Break-even explanation
5. Competitor pricing advice
6. Discount/promotion ideas
7. Warning if price is too low
8. Warning if price may be too high
9. Recommended first price to test
10. How to explain the price to customers

Business/product: {data["business_idea"]}
Product type: {data["product_type"]}
Target customer: {data["target_customer"]}
Country/market: {data["country"]}
Estimated cost price: {data["cost_price"]}
Desired profit: {data["desired_profit"]}
Competitor price if known: {data["competitor_price"]}
Budget level: {data["budget"]}
Pricing style: {data["pricing_style"]}
"""

    try:
        response = safe_openai_chat_completion(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": prompt}
            ]
        )
    except Exception:
        return redirect("/pricing_advisor?pricing_error=ai_response")

    save_pricing_advice(user_id, data, response.choices[0].message.content.strip())
    log_usage(user_id, "pricing_advice")
    return redirect("/pricing_advisor?pricing_notice=created")


@app.route("/payment_guide")
def payment_guide():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    return render_template(
        "payment_guide.html",
        onboarding=get_user_onboarding(user_id),
        guides=get_payment_guides(user_id),
        current_package=get_user_package(user_id) or "Starter",
        usage_summary=get_usage_summary(user_id)
    )


@app.route("/generate_payment_guide", methods=["GET", "POST"])
def generate_payment_guide():
    if "user_id" not in session:
        return redirect("/login")

    if request.method == "GET":
        return redirect("/payment_guide")

    user_id = session["user_id"]
    if usage_limit_reached(user_id, "payment_guide"):
        return usage_limit_redirect("payment_guide", "/payment_guide")
    payment_options = request.form.getlist("payment_options")
    data = {
        "country": request.form.get("country", "").strip(),
        "business_type": request.form.get("business_type", "").strip(),
        "selling_platform": request.form.get("selling_platform", "").strip(),
        "payment_options": ", ".join(payment_options) if payment_options else request.form.get("payment_options", "").strip(),
        "budget": request.form.get("budget", "").strip(),
        "international_payments": request.form.get("international_payments", "").strip()
    }

    if not data["country"] or not data["business_type"]:
        return redirect("/payment_guide?payment_error=missing")

    prompt = f"""
Create a payment setup guide for this BusinessBuilder AI user.

Safety and compliance rules:
- Do not collect or request card numbers, bank passwords, ID numbers, private banking credentials, API keys, or account secrets.
- Do not create PayPal, Paystack, Shopify Payments, bank, or card processing accounts automatically.
- Explain that requirements, fees, availability, verification, settlement times, and documents vary by country and provider.
- Tell the user to verify everything directly with the payment provider, Shopify, bank, accountant, and local regulations.

Include:
1. Best payment methods for the user's country and business type
2. PayPal setup guidance
3. Paystack setup guidance
4. Shopify payment setup guidance
5. EFT/bank transfer guidance
6. Cash on Delivery pros/cons
7. International payment notes
8. Fees/cost considerations in general terms
9. Required documents checklist
10. Fraud prevention tips
11. Best low-budget option
12. Best professional setup option
13. Step-by-step setup checklist

Country: {data["country"]}
Business type: {data["business_type"]}
Selling platform: {data["selling_platform"]}
Preferred payment options: {data["payment_options"]}
Budget level: {data["budget"]}
Need international payments: {data["international_payments"]}
"""

    try:
        response = safe_openai_chat_completion(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": prompt}
            ]
        )
    except Exception:
        return redirect("/payment_guide?payment_error=ai_response")

    save_payment_guide(user_id, data, response.choices[0].message.content.strip())
    log_usage(user_id, "payment_guide")
    return redirect("/payment_guide?payment_notice=created")


@app.route("/supplier_finder")
def supplier_finder():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    return render_template(
        "supplier_finder.html",
        onboarding=get_user_onboarding(user_id),
        recommendations=get_supplier_recommendations(user_id),
        current_package=get_user_package(user_id) or "Starter",
        usage_summary=get_usage_summary(user_id)
    )


@app.route("/generate_supplier_recommendations", methods=["GET", "POST"])
def generate_supplier_recommendations():
    if "user_id" not in session:
        return redirect("/login")

    if request.method == "GET":
        return redirect("/supplier_finder")

    user_id = session["user_id"]
    if usage_limit_reached(user_id, "supplier_guide"):
        return usage_limit_redirect("supplier_guide", "/supplier_finder")
    data = {
        "business_idea": request.form.get("business_idea", "").strip(),
        "product_type": request.form.get("product_type", "").strip(),
        "country": request.form.get("country", "").strip(),
        "budget": request.form.get("budget", "").strip(),
        "sourcing_preference": request.form.get("sourcing_preference", "").strip(),
        "risk_level": request.form.get("risk_level", "").strip(),
        "inventory": request.form.get("inventory", "").strip()
    }

    if not data["business_idea"]:
        return redirect("/supplier_finder?supplier_error=missing")

    prompt = f"""
Create supplier and platform recommendation guidance for this BusinessBuilder AI user.

Safety and research rules:
- Do not scrape Amazon, Alibaba, AliExpress, Takealot, or any website.
- Do not claim live prices, live stock, delivery times, reviews, supplier scores, or availability.
- Generate research guidance and search suggestions only unless official API support is configured.
- Tell the user to verify suppliers, shipping, customs, returns, product quality, platform rules, samples, fees, taxes, and legal restrictions before buying inventory.

Compare these options where relevant:
- Amazon
- Alibaba
- AliExpress
- Local suppliers
- Printful
- Printify
- CJdropshipping
- Zendrop
- Spocket
- Takealot/local marketplace suggestions where relevant
- Digital product platforms
- Handmade/local manufacturing
- Wholesale suppliers
- Direct manufacturer sourcing

For each option include:
1. Best for
2. Budget fit
3. Pros
4. Cons
5. Shipping difficulty
6. Quality-control risk
7. Startup cost level
8. Best product types
9. What the user must verify
10. Recommended option based on budget
11. Cheapest option
12. Safest beginner option
13. Professional long-term option

Business/store idea: {data["business_idea"]}
Product type: {data["product_type"]}
Country/market: {data["country"]}
Budget level: {data["budget"]}
Preference: {data["sourcing_preference"]}
Priority/risk level: {data["risk_level"]}
Willing to hold inventory: {data["inventory"]}
"""

    try:
        response = safe_openai_chat_completion(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": prompt}
            ]
        )
    except Exception:
        return redirect("/supplier_finder?supplier_error=ai_response")

    save_supplier_recommendations(user_id, data, response.choices[0].message.content.strip())
    log_usage(user_id, "supplier_guide")
    return redirect("/supplier_finder?supplier_notice=created")


@app.route("/projects")
def projects_page():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    return render_template(
        "projects.html",
        projects=get_business_projects(user_id),
        active_project=get_active_business_project(user_id)
    )


@app.route("/create_business_project", methods=["POST"])
def create_business_project():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]
    name = request.form.get("name", "").strip()
    business_idea = request.form.get("business_idea", "").strip()
    target_customer = request.form.get("target_customer", "").strip()
    country = request.form.get("country", "").strip()
    budget = request.form.get("budget", "").strip()
    notes = request.form.get("notes", "").strip()

    if not name:
        return redirect("/projects?project_error=name")

    project_id = create_business_project_record(
        user_id,
        name,
        business_idea,
        target_customer,
        country,
        budget,
        notes,
        True
    )

    return redirect(f"/project/{project_id}")


@app.route("/switch_business_project/<int:project_id>")
def switch_business_project(project_id):
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if get_business_project(user_id, project_id):
        set_active_business_project(user_id, project_id)

    return redirect("/dashboard")


@app.route("/project/<int:project_id>")
def project_detail(project_id):
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]
    project = get_business_project(user_id, project_id)

    if not project:
        return redirect("/projects")

    return render_template(
        "project_detail.html",
        project=project,
        business_plans=get_business_plans(user_id),
        product_research_list=get_product_research_list(user_id),
        ai_store_builds=get_ai_store_builds(user_id),
        store_agent_tasks=get_store_agent_tasks(user_id),
        launch_readiness=get_launch_readiness(user_id),
        premium_build=user_package_at_least(user_id, "Premium Build")
    )


@app.route("/pricing")
def pricing():
    current_package = get_user_package(session["user_id"]) if "user_id" in session else None
    return render_template(
        "pricing.html",
        current_package=current_package,
        plan_details=PLAN_DETAILS
    )


@app.route("/service-worker.js")
def service_worker():
    response = send_file(
        os.path.join("static", "service-worker.js"),
        mimetype="application/javascript"
    )
    response.headers["Service-Worker-Allowed"] = "/"
    response.headers["Cache-Control"] = "no-cache"
    return response


@app.route("/terms")
def terms():
    return render_template("terms.html")


@app.route("/privacy")
def privacy():
    return render_template("privacy.html")


@app.route("/refund")
def refund():
    return render_template("refund.html")


# -----------------------------
# PROJECT ROUTES
# -----------------------------

@app.route("/create_project", methods=["POST"])
def create_project():
    if "user_id" not in session:
        return redirect("/login")

    save_project(
        session["user_id"],
        request.form["title"],
        request.form["business_type"],
        request.form["description"]
    )

    return redirect("/dashboard")


@app.route("/upload", methods=["POST"])
def upload_file():
    if "user_id" not in session:
        return redirect("/login")

    if not user_has_paid(session["user_id"]):
        return redirect("/dashboard")

    file = request.files.get("file")

    if not file:
        return redirect("/dashboard")

    filename = secure_filename(file.filename)
    filepath = os.path.join(
        app.config["UPLOAD_FOLDER"],
        filename
    )

    file.save(filepath)

    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            INSERT INTO uploads (user_id, filename, filepath)
            VALUES (?, ?, ?)
        """),
        (session["user_id"], filename, filepath)
    )

    conn.commit()
    conn.close()

    return redirect("/dashboard")


# -----------------------------
# CHAT ROUTES
# -----------------------------

@app.route("/new_chat")
def new_chat():
    if "user_id" not in session:
        return redirect("/login")

    session.pop("chat_id", None)

    return redirect("/")


@app.route("/switch_chat/<int:chat_id>")
def switch_chat(chat_id):
    if "user_id" not in session:
        return redirect("/login")

    session["chat_id"] = chat_id

    return redirect("/")


@app.route("/messages")
def messages():
    if "user_id" not in session:
        return jsonify([])

    chat_id = session.get("chat_id")

    if not chat_id:
        return jsonify([])

    return jsonify(
        get_chat_messages(
            session["user_id"],
            chat_id
        )
    )


# -----------------------------
# PAYMENT ROUTES
# -----------------------------

@app.route("/paystack_checkout")
def paystack_checkout():
    if "user_id" not in session:
        return redirect("/login")

    plan_name = request.args.get("plan", "starter")
    plan = get_paystack_plan(plan_name)

    if not plan:
        return render_error(
            "Choose a valid BusinessBuilder AI package before starting checkout.",
            400,
            "/pricing"
        )

    plan_slug = plan_name.strip().lower()
    paystack_secret = os.getenv("PAYSTACK_SECRET_KEY")

    if not paystack_secret:
        return render_error(
            "Payment checkout is temporarily unavailable. Please try again later.",
            503
        )

    conn = db()
    cur = conn.cursor()

    cur.execute(
        sql("""
            SELECT email
            FROM users
            WHERE id = ?
        """),
        (session["user_id"],)
    )

    user = cur.fetchone()
    conn.close()

    if not user:
        return redirect("/login")

    email = user[0]
    amount = plan["amount"]

    headers = {
        "Authorization": f"Bearer {paystack_secret}",
        "Content-Type": "application/json"
    }

    data = {
        "email": email,
        "amount": amount,
        "currency": plan["currency"],
        "callback_url": (
            request.host_url
            + "payment_success?"
            + urllib.parse.urlencode({"plan": plan_slug})
        ),
        "metadata": {
            "plan": plan_slug,
            "package_name": plan["package_name"],
            "user_id": session["user_id"],
            "email": email
        }
    }

    try:
        response = requests.post(
            "https://api.paystack.co/transaction/initialize",
            headers=headers,
            json=data,
            timeout=20
        )
        response.raise_for_status()
        result = response.json()
    except (requests.RequestException, ValueError):
        return render_error(
            "Payment checkout could not be started. Please try again.",
            503
        )

    authorization_url = (result.get("data") or {}).get("authorization_url")

    if not result.get("status") or not authorization_url:
        return render_error(
            "Payment checkout could not be started. Please try again.",
            503
        )

    return redirect(authorization_url)


@app.route("/payment_success")
def payment_success():
    if "user_id" not in session:
        return redirect("/login")

    reference = request.args.get("reference")

    if not reference:
        return render_error(
            "The payment reference is missing. Please start checkout again.",
            400
        )

    if payment_reference_exists(reference):
        if get_payment_user_id_by_reference(reference) == session["user_id"]:
            ensure_starter_package(session["user_id"])

        return redirect("/dashboard")

    transaction, error = verify_paystack_transaction(reference)

    if error:
        return render_error(
            "Payment verification could not be completed. Please try again.",
            400
        )

    plan = resolve_paystack_plan(
        transaction,
        request.args.get("plan")
    )

    if not plan:
        return render_error(
            "Payment verification returned an invalid package or amount. Please contact support.",
            400
        )

    save_payment(
        session["user_id"],
        "paystack",
        transaction["amount"],
        "success",
        reference
    )
    set_user_package(
        session["user_id"],
        plan["package_name"]
    )

    email = get_user_email(session["user_id"])

    if email:
        send_email(
            email,
            "Your BusinessBuilder AI payment is confirmed",
            (
                "Your Paystack payment has been verified and your BusinessBuilder AI "
                f"{plan['package_name']} package is active. Open your dashboard to complete "
                "the guided workflow."
            )
        )

    return redirect("/dashboard")


@app.route("/paystack_webhook", methods=["POST"])
def paystack_webhook():
    paystack_secret = os.getenv("PAYSTACK_SECRET_KEY")

    if not paystack_secret:
        return "", 503

    raw_body = request.get_data(cache=False)
    received_signature = request.headers.get("X-Paystack-Signature", "")
    expected_signature = hmac.new(
        paystack_secret.encode("utf-8"),
        raw_body,
        hashlib.sha512
    ).hexdigest()

    # Verify the signature before parsing JSON so forged requests cannot
    # mark an account as paid or create fake payment records.
    if (
        not received_signature
        or not hmac.compare_digest(expected_signature, received_signature)
    ):
        return "", 400

    try:
        event = json.loads(raw_body)
    except (TypeError, ValueError, json.JSONDecodeError):
        return "", 400

    if not isinstance(event, dict):
        return "", 400

    if event.get("event") != "charge.success":
        return "", 200

    transaction = event.get("data")

    if not isinstance(transaction, dict):
        return "", 200

    customer = transaction.get("customer")

    if not isinstance(customer, dict):
        return "", 200

    reference = transaction.get("reference")
    amount = transaction.get("amount")
    status = transaction.get("status")
    customer_email = customer.get("email")

    if (
        not isinstance(reference, str)
        or not reference.strip()
        or not isinstance(amount, (int, float))
        or isinstance(amount, bool)
        or status != "success"
        or not isinstance(customer_email, str)
        or not customer_email.strip()
    ):
        return "", 200

    reference = reference.strip()
    user_id = get_user_id_by_email(customer_email)

    if not user_id:
        return "", 200

    plan = resolve_paystack_plan(transaction)

    if not plan:
        return "", 200

    payment_user_id = get_payment_user_id_by_reference(reference)

    if not payment_user_id:
        saved = save_payment(
            user_id,
            "paystack",
            amount,
            "success",
            reference
        )

        if saved:
            set_user_package(
                user_id,
                plan["package_name"]
            )
            send_email(
                customer_email.strip().lower(),
                "Your BusinessBuilder AI payment is confirmed",
                (
                    "Your Paystack payment has been verified and your BusinessBuilder AI "
                    f"{plan['package_name']} package is active. Open your dashboard to complete "
                    "the guided workflow."
                )
            )

    return "", 200


# -----------------------------
# WORKFLOW ROUTES
# -----------------------------

@app.route("/business_workflow")
def business_workflow():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    completed_steps = get_completed_steps(user_id)
    completed_count = len(completed_steps)
    answer_count = len(get_all_workflow_answers(user_id))
    total_steps = 9
    progress_percent = int((completed_count / total_steps) * 100)

    return render_template(
        "business_workflow.html",
        completed_steps=completed_steps,
        completed_count=completed_count,
        answer_count=answer_count,
        total_steps=total_steps,
        progress_percent=progress_percent,
        current_package=get_user_package(user_id),
        pro_package=user_package_at_least(user_id, "Pro")
    )


@app.route("/build_approval", methods=["GET", "POST"])
def build_approval():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_is_paid(user_id):
        return redirect("/dashboard")

    requested_action = request.values.get("action", "full_build")

    if requested_action not in BUILD_APPROVAL_ACTIONS:
        requested_action = "full_build"

    if requested_action == "product_research_products" and request.values.get("research_id"):
        session["product_research_id"] = request.values.get("research_id")

    if not user_package_at_least(user_id, "Pro"):
        return package_access_redirect("Pro")

    if request.method == "POST":
        grant_build_approval(requested_action)

        return redirect(BUILD_APPROVAL_ACTIONS[requested_action])

    workflow_answers = get_nonempty_workflow_answers(user_id)
    shopify_plans = get_shopify_plans(user_id)
    latest_canva_design_brief = get_latest_canva_design_brief(user_id)

    return render_template(
        "build_approval.html",
        requested_action=requested_action,
        workflow_answers=workflow_answers,
        latest_shopify_plan=shopify_plans[0] if shopify_plans else None,
        latest_canva_design_brief=latest_canva_design_brief,
        shopify_product_preview=[
            "3 to 5 draft Shopify products with titles, descriptions, suggested prices, and categories"
        ],
        shopify_store_preview=[
            "3 to 5 draft Shopify products",
            "2 to 4 unpublished Shopify collections",
            "About Us page draft",
            "Contact Us page draft",
            "Shipping Policy page draft",
            "Refund Policy page draft"
        ],
        ai_store_preview=[
            "Complete AI-generated Shopify store package",
            "5 to 10 draft Shopify products",
            "Unpublished Shopify collections",
            "Unpublished About, Contact, FAQ, Shipping, and Refund pages",
            "Canva logo, banner, and social template briefs",
            "Launch checklist"
        ],
        product_research_preview=[
            "Up to 5 Shopify draft products from approved AI Product Finder research",
            "Product titles, descriptions, suggested prices, and categories",
            "Draft products only; nothing is published automatically",
            "Supplier, stock, legal, and pricing details must be verified manually"
        ],
        canva_design_preview=[
            "Logo concept draft",
            "Shopify homepage banner draft",
            "Instagram post template draft",
            "Instagram story template draft",
            "TikTok/Reels cover draft",
            "Business card draft"
        ]
    )


@app.route("/complete_step/<int:step_number>/<step_name>")
def complete_step(step_number, step_name):
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    mark_workflow_step_complete(
        user_id,
        step_number,
        step_name
    )

    return redirect("/business_workflow")


@app.route("/workflow_step/<int:step_number>", methods=["GET", "POST"])
def workflow_step(step_number):
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    step = WORKFLOW_STEPS.get(step_number)

    if not step:
        return redirect("/business_workflow")

    if request.method == "POST":
        answer = request.form.get("answer", "").strip()

        if not answer:
            return render_template(
                "workflow_step.html",
                step_number=step_number,
                step_name=step["name"],
                question=step["question"],
                existing_answer="",
                message="Enter an answer before saving this workflow step."
            ), 400

        save_workflow_answer(
            user_id,
            step_number,
            step["name"],
            answer
        )

        mark_workflow_step_complete(
            user_id,
            step_number,
            step["name"]
        )

        return redirect("/business_workflow")

    existing_answer = get_workflow_answer(
        user_id,
        step_number
    )

    return render_template(
        "workflow_step.html",
        step_number=step_number,
        step_name=step["name"],
        question=step["question"],
        existing_answer=existing_answer
    )


@app.route("/generate_business_plan")
def generate_business_plan():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    if usage_limit_reached(user_id, "business_plan"):
        return usage_limit_redirect("business_plan", "/business_workflow")

    answers = get_nonempty_workflow_answers(user_id)

    if not answers:
        return redirect("/business_workflow?plan_error=no_answers")

    workflow_text = ""

    for step_number, step_name, answer in answers:
        workflow_text += f"""
Step {step_number}: {step_name}
Answer:
{answer}

"""

    prompt = f"""
Create a complete professional business plan using the user's saved workflow answers.

The business plan must include:

1. Executive Summary
2. Business Idea
3. Brand Name and Brand Identity
4. Target Market
5. Products or Services
6. Pricing Strategy
7. Marketing Plan
8. Shopify Setup Plan
9. Canva Branding Plan
10. Launch Checklist
11. Next 30-Day Action Plan

User workflow answers:
{workflow_text}
"""

    response = safe_openai_chat_completion(
        model="gpt-4.1-mini",
        messages=[
            {
                "role": "system",
                "content": SYSTEM_PROMPT
            },
            {
                "role": "user",
                "content": prompt
            }
        ]
    )

    business_plan = response.choices[0].message.content

    plan_id = save_business_plan(
        user_id,
        "Generated Business Plan",
        business_plan
    )

    log_usage(user_id, "business_plan")

    email = get_user_email(user_id)

    if email:
        send_email(
            email,
            "Your BusinessBuilder AI business plan is ready",
            (
                "Your generated business plan has been saved. "
                "Open your BusinessBuilder AI dashboard to review it or download the PDF."
            )
        )

    return redirect(f"/business_plan/{plan_id}")


@app.route("/generate_shopify_plan")
def generate_shopify_plan():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    if usage_limit_reached(user_id, "shopify_plan"):
        return usage_limit_redirect("shopify_plan", "/business_workflow")

    answers = get_nonempty_workflow_answers(user_id)

    if not answers:
        return redirect("/business_workflow?shopify_error=no_answers")

    workflow_text = ""

    for step_number, step_name, answer in answers:
        workflow_text += f"""
Step {step_number}: {step_name}
Answer:
{answer}

"""

    prompt = f"""
Create a complete professional Shopify store setup plan using the user's saved workflow answers.
Provide practical copy and recommendations that the user can apply while building their Shopify store.

The Shopify setup plan must include:

1. Store Name Suggestion
2. Shopify Homepage Structure
3. Product Titles
4. Product Descriptions
5. Collections
6. Navigation Menu
7. About Page Copy
8. Contact Page Copy
9. Shipping Policy Draft
10. Refund Policy Draft
11. Store Launch Checklist

User workflow answers:
{workflow_text}
"""

    response = safe_openai_chat_completion(
        model="gpt-4.1-mini",
        messages=[
            {
                "role": "system",
                "content": SYSTEM_PROMPT
            },
            {
                "role": "user",
                "content": prompt
            }
        ]
    )

    shopify_plan_content = response.choices[0].message.content

    plan_id = save_shopify_plan(
        user_id,
        "Generated Shopify Setup Plan",
        shopify_plan_content
    )

    log_usage(user_id, "shopify_plan")

    return redirect(f"/shopify_plan/{plan_id}")


@app.route("/generate_canva_branding")
def generate_canva_branding():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    if usage_limit_reached(user_id, "canva_branding"):
        return usage_limit_redirect("canva_branding", "/business_workflow")

    answers = get_nonempty_workflow_answers(user_id)

    if not answers:
        return redirect("/business_workflow?canva_error=no_answers")

    workflow_text = ""

    for step_number, step_name, answer in answers:
        workflow_text += f"""
Step {step_number}: {step_name}
Answer:
{answer}

"""

    prompt = f"""
Create a complete professional Canva-ready branding package using the user's saved workflow answers.
Provide practical design direction and instructions the user can apply manually inside Canva.
Do not claim that Canva is connected and do not claim that designs were created automatically.

The Canva-ready branding package must include:

1. Logo Concept
2. Brand Color Palette
3. Font Direction
4. Brand Personality
5. Instagram Post Template Ideas
6. Instagram Story Template Ideas
7. TikTok/Reels Cover Ideas
8. Business Card Design Idea
9. Shopify Homepage Banner Concept
10. Product Mockup Ideas
11. Canva Search Keywords
12. Step-by-Step Canva Creation Checklist

User workflow answers:
{workflow_text}
"""

    response = safe_openai_chat_completion(
        model="gpt-4.1-mini",
        messages=[
            {
                "role": "system",
                "content": SYSTEM_PROMPT
            },
            {
                "role": "user",
                "content": prompt
            }
        ]
    )

    package_content = response.choices[0].message.content

    package_id = save_canva_branding_package(
        user_id,
        "Generated Canva Branding Package",
        package_content
    )

    log_usage(user_id, "canva_branding")

    return redirect(f"/canva_branding/{package_id}")


@app.route("/generate_canva_design_brief")
def generate_canva_design_brief():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    if not user_package_at_least(user_id, "Pro"):
        return package_access_redirect("Pro", "/business_workflow")

    if usage_limit_reached(user_id, "canva_design_brief"):
        return usage_limit_redirect("canva_design_brief", "/business_workflow")

    answers = get_nonempty_workflow_answers(user_id)

    if not answers:
        return redirect("/business_workflow?canva_brief_error=no_answers")

    workflow_text = ""

    for step_number, step_name, answer in answers:
        workflow_text += f"""
Step {step_number}: {step_name}
Answer:
{answer}

"""

    branding_packages = get_canva_branding_packages(user_id)
    branding_package_text = (
        branding_packages[0][2]
        if branding_packages
        else "No saved Canva branding package is available yet."
    )

    prompt = f"""
Create a complete professional set of Canva design briefs using the user's saved workflow answers
and the latest saved Canva branding package when available.
Provide detailed, practical instructions the user can apply manually inside Canva.
Do not claim that Canva designs were created automatically and do not call any Canva API.

The Canva design brief document must include:

1. Logo Design Brief
2. Shopify Homepage Banner Design Brief
3. Instagram Post Template Brief
4. Instagram Story Template Brief
5. TikTok/Reels Cover Brief
6. Business Card Brief
7. Product Mockup Brief
8. Brand Color Usage Instructions
9. Font Usage Instructions
10. Canva Keywords to Search
11. Step-by-Step Canva Creation Checklist

Latest saved Canva branding package:
{branding_package_text}

User workflow answers:
{workflow_text}
"""

    response = safe_openai_chat_completion(
        model="gpt-4.1-mini",
        messages=[
            {
                "role": "system",
                "content": SYSTEM_PROMPT
            },
            {
                "role": "user",
                "content": prompt
            }
        ]
    )

    brief_content = response.choices[0].message.content

    brief_id = save_canva_design_brief(
        user_id,
        "Generated Canva Design Brief",
        brief_content
    )

    log_usage(user_id, "canva_design_brief")

    return redirect(f"/canva_design_brief/{brief_id}")


@app.route("/create_canva_design")
def create_canva_design_from_brief():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    if not user_package_at_least(user_id, "Pro"):
        return package_access_redirect("Pro", "/business_workflow")

    if not consume_build_approval("canva_designs"):
        return build_approval_redirect("canva_designs")

    if usage_limit_reached(user_id, "canva_design"):
        return usage_limit_redirect("canva_design", "/business_workflow")

    connection = get_canva_connection(user_id)

    if not connection or connection[2] != "connected":
        return redirect("/business_workflow?canva_design_error=connection")

    brief = get_latest_canva_design_brief(user_id)

    if not brief:
        return redirect("/business_workflow?canva_design_error=no_brief")

    title = f"{brief[1]} Draft"[:255]
    canva_design_id, edit_url, view_url, status, error = create_canva_design(
        user_id,
        title
    )

    if error:
        return redirect("/business_workflow?canva_design_error=create_failed")

    save_canva_design(
        user_id,
        title,
        canva_design_id,
        edit_url,
        view_url,
        status
    )
    log_usage(user_id, "canva_design")

    return redirect("/dashboard?canva_design=created")


@app.route("/create_canva_designs")
def create_canva_designs_from_brief():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    if not user_package_at_least(user_id, "Pro"):
        return package_access_redirect("Pro", "/business_workflow")

    if not consume_build_approval("canva_designs"):
        return build_approval_redirect("canva_designs")

    if usage_limit_reached(user_id, "canva_design"):
        return usage_limit_redirect("canva_design", "/business_workflow")

    connection = get_canva_connection(user_id)

    if not connection or connection[2] != "connected":
        return redirect("/business_workflow?canva_design_error=connection")

    brief = get_latest_canva_design_brief(user_id)

    if not brief:
        return redirect("/business_workflow?canva_design_error=no_brief")

    base_title = brief[1][:180]
    draft_specs = [
        ("Logo Concept", 1080, 1080),
        ("Shopify Homepage Banner", 1800, 700),
        ("Instagram Post Template", 1080, 1080),
        ("Instagram Story Template", 1080, 1920),
        ("TikTok Reels Cover", 1080, 1920),
        ("Business Card", 1050, 600)
    ]
    created_count = 0
    failed_count = 0

    for label, width, height in draft_specs:
        if usage_limit_reached(user_id, "canva_design"):
            break

        title = f"{base_title} - {label}"[:255]
        canva_design_id, edit_url, view_url, status, error = (
            create_canva_design(
                user_id,
                title,
                width,
                height
            )
        )

        if error:
            failed_count += 1
            continue

        save_canva_design(
            user_id,
            title,
            canva_design_id,
            edit_url,
            view_url,
            status
        )
        log_usage(user_id, "canva_design")
        created_count += 1

    if not created_count:
        return redirect("/business_workflow?canva_design_error=batch_create_failed")

    return redirect(
        f"/dashboard?canva_designs=created"
        f"&created_count={created_count}"
        f"&failed_count={failed_count}"
    )


@app.route("/generate_build_quote")
def generate_build_quote():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    answers = get_nonempty_workflow_answers(user_id)

    if not answers:
        return redirect("/business_workflow?quote_error=no_answers")

    workflow_text = ""

    for step_number, step_name, answer in answers:
        workflow_text += f"""
Step {step_number}: {step_name}
Answer:
{answer}

"""

    shopify_plan = "Recommended Shopify plan based on the store requirements"
    canva_plan = "Recommended Canva plan based on the branding requirements"
    estimated_shopify_cost = "Depends on Shopify plan, domain, apps, and theme"
    estimated_canva_cost = "Depends on Canva plan and premium assets"
    service_fee = "R499"
    total_estimate = "R499 + Shopify/Canva external costs"

    prompt = f"""
Create a professional Store + Branding Build Quote using the user's saved workflow answers.
This is an informational estimate only. Do not claim that Shopify or Canva APIs are connected.
Do not claim that external Shopify or Canva costs will be charged.
Explain that external costs depend on the user's final choices.

The quote must include:

1. Store Build Summary
2. Shopify Setup Items
3. Canva Branding Items
4. Recommended Shopify Plan
5. Recommended Canva Plan
6. Estimated Shopify External Cost
7. Estimated Canva External Cost
8. BusinessBuilder AI Service Fee
9. Total Estimated Setup Cost
10. What Happens After Approval

Use these exact placeholder estimates:
- Estimated Shopify external cost: {estimated_shopify_cost}
- Estimated Canva external cost: {estimated_canva_cost}
- BusinessBuilder AI service fee: {service_fee}
- Total estimated setup cost: {total_estimate}

User workflow answers:
{workflow_text}
"""

    response = safe_openai_chat_completion(
        model="gpt-4.1-mini",
        messages=[
            {
                "role": "system",
                "content": SYSTEM_PROMPT
            },
            {
                "role": "user",
                "content": prompt
            }
        ]
    )

    quote_content = response.choices[0].message.content

    quote_id = save_build_quote(
        user_id,
        "Store + Branding Build Quote",
        shopify_plan,
        canva_plan,
        estimated_shopify_cost,
        estimated_canva_cost,
        service_fee,
        total_estimate,
        quote_content
    )

    return redirect(f"/build_quote/{quote_id}")


def build_workflow_text(answers):
    workflow_text = ""

    for step_number, step_name, answer in answers:
        workflow_text += f"""
Step {step_number}: {step_name}
Answer:
{answer}

"""

    return workflow_text


def normalize_store_package(raw_package):
    if not isinstance(raw_package, dict):
        raise ValueError

    products = raw_package.get("products")
    collections = raw_package.get("product_collections")

    if collections is None:
        collections = raw_package.get("collections")

    if not isinstance(products, list) or not 5 <= len(products) <= 10:
        raise ValueError

    if not isinstance(collections, list) or not collections:
        raise ValueError

    required_text_keys = [
        "store_name",
        "store_positioning",
        "homepage_hero_headline",
        "homepage_subheadline",
        "homepage_section_copy",
        "about_page_copy",
        "contact_page_copy",
        "faq_page_copy",
        "shipping_policy_draft",
        "refund_policy_draft",
        "navigation_menu_plan",
        "canva_logo_brief",
        "canva_banner_brief",
        "canva_social_media_template_brief",
        "launch_checklist"
    ]

    normalized = {}

    for key in required_text_keys:
        value = raw_package.get(key)

        if isinstance(value, (list, dict)):
            value = json.dumps(value, indent=2)

        value = str(value or "").strip()

        if not value:
            raise ValueError

        normalized[key] = value

    normalized_products = []

    for product in products:
        if not isinstance(product, dict):
            raise ValueError

        title = str(product.get("title", "")).strip()
        description = str(product.get("description", "")).strip()
        suggested_price = str(product.get("suggested_price", "")).strip()
        category = str(product.get("category", "")).strip()
        seo_title = str(product.get("seo_title", "")).strip()
        meta_description = str(product.get("meta_description", "")).strip()

        if not title or not description or not suggested_price:
            raise ValueError

        normalized_products.append({
            "title": title,
            "description": description,
            "suggested_price": suggested_price,
            "category": category,
            "seo_title": seo_title,
            "meta_description": meta_description
        })

    normalized_collections = []

    for collection in collections:
        if not isinstance(collection, dict):
            raise ValueError

        title = str(collection.get("title", "")).strip()
        description = str(collection.get("description", "")).strip()

        if not title or not description:
            raise ValueError

        normalized_collections.append({
            "title": title,
            "description": description
        })

    normalized["products"] = normalized_products
    normalized["product_collections"] = normalized_collections

    return normalized


def create_shopify_assets_from_store_package(user_id, store_package):
    created_assets = {
        "products": [],
        "collections": [],
        "pages": []
    }
    failed_actions = []

    for product in store_package["products"]:
        if usage_limit_reached(user_id, "shopify_product"):
            failed_actions.append(
                "Product creation stopped because the Shopify product usage limit was reached."
            )
            break

        title = product["title"]
        description = product["description"]
        suggested_price = product.get("suggested_price", "")
        category = product.get("category", "")
        saved_description = (
            f"{description}\n\n"
            f"Suggested price: {suggested_price}\n"
            f"Collection / category: {category}\n"
            f"SEO title: {product.get('seo_title', '')}\n"
            f"Meta description: {product.get('meta_description', '')}"
        )

        shopify_product_id, status, error = create_shopify_product(
            user_id,
            title,
            description,
            suggested_price,
            category
        )

        if error:
            failed_actions.append(f"Product '{title}': {error}")
            continue

        save_shopify_product(
            user_id,
            title,
            saved_description,
            shopify_product_id,
            status
        )
        log_usage(user_id, "shopify_product")
        created_assets["products"].append({
            "title": title,
            "shopify_id": shopify_product_id,
            "status": status
        })

    for collection in store_package["product_collections"]:
        title = collection["title"]
        description = collection["description"]
        shopify_collection_id, status, error = create_shopify_collection(
            user_id,
            title,
            description
        )

        if error:
            failed_actions.append(f"Collection '{title}': {error}")
            continue

        save_shopify_collection(
            user_id,
            title,
            description,
            shopify_collection_id,
            status
        )
        created_assets["collections"].append({
            "title": title,
            "shopify_id": shopify_collection_id,
            "status": status
        })

    page_map = [
        ("About Us", "about", store_package["about_page_copy"]),
        ("Contact Us", "contact", store_package["contact_page_copy"]),
        ("FAQ", "faq", store_package["faq_page_copy"]),
        ("Shipping Policy", "shipping_policy", store_package["shipping_policy_draft"]),
        ("Refund Policy", "refund_policy", store_package["refund_policy_draft"])
    ]

    for title, page_type, content in page_map:
        shopify_page_id, status, error = create_shopify_page(
            user_id,
            title,
            content
        )

        if error:
            failed_actions.append(f"Page '{title}': {error}")
            continue

        save_shopify_page(
            user_id,
            title,
            page_type,
            content,
            shopify_page_id,
            status
        )
        created_assets["pages"].append({
            "title": title,
            "shopify_id": shopify_page_id,
            "status": status
        })

    return created_assets, failed_actions


@app.route("/generate_full_store")
def generate_full_store():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    if not user_package_at_least(user_id, "Pro"):
        return package_access_redirect("Pro", "/store_builder")

    answers = get_nonempty_workflow_answers(user_id)

    if not answers:
        return redirect("/store_builder?store_error=no_answers")

    shopify_connection = get_shopify_connection(user_id)
    shopify_connected = bool(
        shopify_connection
        and shopify_connection[3] == "connected"
    )

    if shopify_connected and not consume_build_approval("ai_store"):
        return build_approval_redirect("ai_store")

    workflow_text = build_workflow_text(answers)
    prompt = f"""
Generate a complete Shopify store build package from the user's saved workflow answers.
Return JSON only.

Use exactly these top-level keys:
- store_name
- store_positioning
- homepage_hero_headline
- homepage_subheadline
- homepage_section_copy
- product_collections
- products
- about_page_copy
- contact_page_copy
- faq_page_copy
- shipping_policy_draft
- refund_policy_draft
- navigation_menu_plan
- canva_logo_brief
- canva_banner_brief
- canva_social_media_template_brief
- launch_checklist

product_collections must be an array of collection objects with title and description.
products must be an array of 5 to 10 product objects.
Each product object must include:
- title
- description
- suggested_price
- category
- seo_title
- meta_description

Do not include HTML. Do not claim the store is live or published.
Use practical customer-facing copy suitable for a new Shopify store.

User workflow answers:
{workflow_text}
"""

    try:
        response = safe_openai_chat_completion(
            model="gpt-4.1-mini",
            response_format={"type": "json_object"},
            messages=[
                {
                    "role": "system",
                    "content": SYSTEM_PROMPT
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ]
        )
        raw_package = json.loads(response.choices[0].message.content)
        store_package = normalize_store_package(raw_package)
    except (ExternalServiceError, KeyError, TypeError, ValueError, json.JSONDecodeError):
        return redirect("/store_builder?store_error=ai_response")

    created_assets = {
        "products": [],
        "collections": [],
        "pages": []
    }
    failed_actions = []
    notices = []

    if shopify_connected:
        created_assets, failed_actions = create_shopify_assets_from_store_package(
            user_id,
            store_package
        )
    else:
        notices.append("Connect Shopify to create products automatically.")

    status = "generated"

    if failed_actions:
        status = "partially_created"

    content = json.dumps(
        {
            "store_package": store_package,
            "created_shopify_assets": created_assets,
            "failed_shopify_actions": failed_actions,
            "notices": notices
        },
        indent=2
    )

    build_id = save_ai_store_build(
        user_id,
        "AI Generated Shopify Store Package",
        store_package["store_name"],
        content,
        status
    )

    email = get_user_email(user_id)
    if email:
        send_email_notification(
            email,
            "Your BusinessBuilder AI store package is ready",
            "Your AI-generated store package has been saved. Open your dashboard or Build Center to review it."
        )

    query = ""

    if notices:
        query = "?store_notice=shopify_not_connected"
    elif failed_actions:
        query = "?store_notice=partial_shopify"
    else:
        query = "?store_notice=created"

    return redirect(f"/store_build/{build_id}{query}")


@app.route("/store_build/<int:build_id>")
def store_build(build_id):
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    build = get_ai_store_build(user_id, build_id)

    if not build:
        return redirect("/store_builder")

    build_data = None

    try:
        build_data = json.loads(build[3])
    except (TypeError, ValueError, json.JSONDecodeError):
        build_data = None

    return render_template(
        "store_build.html",
        build=build,
        build_data=build_data
    )


@app.route("/download_store_build/<int:build_id>")
def download_store_build(build_id):
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    build = get_ai_store_build(user_id, build_id)

    if not build:
        return redirect("/store_builder")

    if usage_limit_reached(user_id, "pdf_export"):
        return usage_limit_redirect("pdf_export", f"/store_build/{build_id}")

    filename = re.sub(r"[_-]+", "-", secure_filename(build[2])).strip("-")
    filename = filename or "ai-store-build"

    response = send_file(
        create_business_plan_pdf(build[1], build[3]),
        mimetype="application/pdf",
        as_attachment=True,
        download_name=f"{filename}-store-build.pdf"
    )

    log_usage(user_id, "pdf_export")

    return response


@app.route("/create_shopify_product")
def create_shopify_product_from_workflow():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    if not user_package_at_least(user_id, "Pro"):
        return package_access_redirect("Pro", "/business_workflow")

    if not consume_build_approval("shopify_products"):
        return build_approval_redirect("shopify_products")

    if usage_limit_reached(user_id, "shopify_product"):
        return usage_limit_redirect("shopify_product", "/business_workflow")

    connection = get_shopify_connection(user_id)

    if not connection or connection[3] != "connected":
        return redirect("/business_workflow?product_error=shopify_connection")

    answers = get_nonempty_workflow_answers(user_id)

    if not answers:
        return redirect("/business_workflow?product_error=no_answers")

    workflow_text = ""

    for step_number, step_name, answer in answers:
        workflow_text += f"""
Step {step_number}: {step_name}
Answer:
{answer}

"""

    prompt = f"""
Create one Shopify product from the user's saved workflow answers.
Return JSON only with exactly these keys:
- title
- description

The title must be concise and suitable for an online store.
The description must be a clear customer-facing product description in plain text.

User workflow answers:
{workflow_text}
"""

    try:
        response = safe_openai_chat_completion(
            model="gpt-4.1-mini",
            response_format={"type": "json_object"},
            messages=[
                {
                    "role": "system",
                    "content": SYSTEM_PROMPT
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ]
        )
    except Exception:
        return redirect("/business_workflow?product_error=ai_response")

    try:
        product_details = json.loads(response.choices[0].message.content)
        title = product_details["title"].strip()
        description = product_details["description"].strip()
    except (KeyError, TypeError, ValueError, json.JSONDecodeError):
        return redirect("/business_workflow?product_error=ai_response")

    if not title or not description:
        return redirect("/business_workflow?product_error=ai_response")

    shopify_product_id, status, error = create_shopify_product(
        user_id,
        title,
        description
    )

    if error:
        return redirect("/business_workflow?product_error=create_failed")

    save_shopify_product(
        user_id,
        title,
        description,
        shopify_product_id,
        status
    )
    log_usage(user_id, "shopify_product")

    return redirect("/dashboard?shopify_product=created")


@app.route("/build_shopify_store_draft")
def build_shopify_store_draft():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    if not user_package_at_least(user_id, "Pro"):
        return package_access_redirect("Pro", "/business_workflow")

    if not consume_build_approval("shopify_store"):
        return build_approval_redirect("shopify_store")

    continue_full_build = session.pop("continue_full_build", False)

    if usage_limit_reached(user_id, "shopify_product"):
        return usage_limit_redirect("shopify_product", "/business_workflow")

    connection = get_shopify_connection(user_id)

    if not connection or connection[3] != "connected":
        return redirect("/business_workflow?store_draft_error=shopify_connection")

    answers = get_nonempty_workflow_answers(user_id)

    if not answers:
        return redirect("/business_workflow?store_draft_error=no_answers")

    workflow_text = ""

    for step_number, step_name, answer in answers:
        workflow_text += f"""
Step {step_number}: {step_name}
Answer:
{answer}

"""

    prompt = f"""
Create a Shopify draft store build package from the user's saved workflow answers.
Return JSON only with exactly these top-level keys:
- products
- collections
- pages

The "products" value must be an array of 3 to 5 objects.
Each product object must contain exactly:
- title
- description
- suggested_price
- category

The "collections" value must be an array of 2 to 4 objects.
Each collection object must contain exactly:
- title
- description

The "pages" value must be an object with exactly:
- about
- contact
- shipping_policy
- refund_policy

Each page value must contain clear customer-facing plain-text copy.
Do not include HTML. Do not claim the store is published or live.

User workflow answers:
{workflow_text}
"""

    try:
        response = safe_openai_chat_completion(
            model="gpt-4.1-mini",
            response_format={"type": "json_object"},
            messages=[
                {
                    "role": "system",
                    "content": SYSTEM_PROMPT
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ]
        )
    except Exception:
        return redirect("/business_workflow?store_draft_error=ai_response")

    try:
        draft_details = json.loads(response.choices[0].message.content)
        products = draft_details["products"]
        collections = draft_details["collections"]
        pages = draft_details["pages"]

        if not isinstance(products, list) or not 3 <= len(products) <= 5:
            raise ValueError

        if not isinstance(collections, list) or not 2 <= len(collections) <= 4:
            raise ValueError

        if not isinstance(pages, dict):
            raise ValueError
    except (KeyError, TypeError, ValueError, json.JSONDecodeError):
        return redirect("/business_workflow?store_draft_error=ai_response")

    created_product_count = 0
    created_collection_count = 0
    created_page_count = 0
    failed_count = 0

    for product in products:
        if usage_limit_reached(user_id, "shopify_product"):
            break

        try:
            title = product["title"].strip()
            description = product["description"].strip()
            suggested_price = product["suggested_price"].strip()
            category = product["category"].strip()
        except (AttributeError, KeyError, TypeError):
            failed_count += 1
            continue

        if not title or not description or not suggested_price or not category:
            failed_count += 1
            continue

        saved_description = (
            f"{description}\n\n"
            f"Suggested price: {suggested_price}\n"
            f"Collection / category: {category}"
        )
        shopify_product_id, status, error = create_shopify_product(
            user_id,
            title,
            description,
            suggested_price,
            category
        )

        if error:
            failed_count += 1
            continue

        save_shopify_product(
            user_id,
            title,
            saved_description,
            shopify_product_id,
            status
        )
        log_usage(user_id, "shopify_product")
        created_product_count += 1

    for collection in collections:
        try:
            title = collection["title"].strip()
            description = collection["description"].strip()
        except (AttributeError, KeyError, TypeError):
            failed_count += 1
            continue

        if not title or not description:
            failed_count += 1
            continue

        shopify_collection_id, status, error = create_shopify_collection(
            user_id,
            title,
            description
        )

        if error:
            failed_count += 1
            continue

        save_shopify_collection(
            user_id,
            title,
            description,
            shopify_collection_id,
            status
        )
        created_collection_count += 1

    page_titles = {
        "about": "About Us",
        "contact": "Contact Us",
        "shipping_policy": "Shipping Policy",
        "refund_policy": "Refund Policy"
    }

    for page_type, title in page_titles.items():
        try:
            content = pages[page_type].strip()
        except (AttributeError, KeyError, TypeError):
            failed_count += 1
            continue

        if not content:
            failed_count += 1
            continue

        shopify_page_id, status, error = create_shopify_page(
            user_id,
            title,
            content
        )

        if error:
            failed_count += 1
            continue

        save_shopify_page(
            user_id,
            title,
            page_type,
            content,
            shopify_page_id,
            status
        )
        created_page_count += 1

    if not (
        created_product_count
        or created_collection_count
        or created_page_count
    ):
        return redirect("/business_workflow?store_draft_error=create_failed")

    if continue_full_build:
        return redirect("/create_canva_designs")

    return redirect(
        f"/shopify_build_summary?draft_store=created"
        f"&product_count={created_product_count}"
        f"&collection_count={created_collection_count}"
        f"&page_count={created_page_count}"
        f"&failed_count={failed_count}"
    )


@app.route("/create_shopify_products")
def create_shopify_products_from_workflow():
    if "user_id" not in session:
        return redirect("/login")

    user_id = session["user_id"]

    if not user_has_paid(user_id):
        return redirect("/dashboard")

    if not user_package_at_least(user_id, "Pro"):
        return package_access_redirect("Pro", "/business_workflow")

    if not consume_build_approval("shopify_products"):
        return build_approval_redirect("shopify_products")

    if usage_limit_reached(user_id, "shopify_product"):
        return usage_limit_redirect("shopify_product", "/business_workflow")

    connection = get_shopify_connection(user_id)

    if not connection or connection[3] != "connected":
        return redirect("/business_workflow?product_error=shopify_connection")

    answers = get_nonempty_workflow_answers(user_id)

    if not answers:
        return redirect("/business_workflow?product_error=no_answers")

    workflow_text = ""

    for step_number, step_name, answer in answers:
        workflow_text += f"""
Step {step_number}: {step_name}
Answer:
{answer}

"""

    prompt = f"""
Create 3 to 5 Shopify products from the user's saved workflow answers.
Return JSON only with one top-level key named "products".
The value of "products" must be an array of 3 to 5 objects.
Each object must contain exactly these keys:
- title
- description
- suggested_price
- category

Each title must be concise and suitable for an online store.
Each description must be clear customer-facing product copy in plain text.
Each suggested_price must be a plain-text price suggestion.
Each category must be a concise Shopify product type or collection suggestion.

User workflow answers:
{workflow_text}
"""

    response = safe_openai_chat_completion(
        model="gpt-4.1-mini",
        response_format={"type": "json_object"},
        messages=[
            {
                "role": "system",
                "content": SYSTEM_PROMPT
            },
            {
                "role": "user",
                "content": prompt
            }
        ]
    )

    try:
        product_details = json.loads(response.choices[0].message.content)
        products = product_details["products"]

        if not isinstance(products, list) or not 3 <= len(products) <= 5:
            raise ValueError
    except (KeyError, TypeError, ValueError, json.JSONDecodeError):
        return redirect("/business_workflow?product_error=batch_ai_response")

    created_count = 0
    failed_count = 0

    for product in products:
        if usage_limit_reached(user_id, "shopify_product"):
            break

        try:
            title = product["title"].strip()
            description = product["description"].strip()
            suggested_price = product["suggested_price"].strip()
            category = product["category"].strip()
        except (AttributeError, KeyError, TypeError):
            failed_count += 1
            continue

        if not title or not description or not suggested_price or not category:
            failed_count += 1
            continue

        saved_description = (
            f"{description}\n\n"
            f"Suggested price: {suggested_price}\n"
            f"Collection / category: {category}"
        )
        shopify_product_id, status, error = create_shopify_product(
            user_id,
            title,
            description,
            suggested_price,
            category
        )

        if error:
            failed_count += 1
            continue

        save_shopify_product(
            user_id,
            title,
            saved_description,
            shopify_product_id,
            status
        )
        log_usage(user_id, "shopify_product")
        created_count += 1

    if not created_count:
        return redirect("/business_workflow?product_error=batch_create_failed")

    return redirect(
        f"/dashboard?shopify_products=created"
        f"&created_count={created_count}"
        f"&failed_count={failed_count}"
    )


# -----------------------------
# MAIN AI CHAT ROUTE
# -----------------------------

@app.route("/chat", methods=["POST"])
def chat():
    if "user_id" not in session:
        return jsonify({
            "reply": "Please log in first."
        })

    user_id = session["user_id"]

    data = request.get_json(silent=True) or {}
    user_message = str(data.get("message", "")).strip()

    if not user_message:
        return jsonify({
            "reply": "Enter a message before sending."
        }), 400

    if usage_limit_reached(user_id, "chat_message"):
        return jsonify({
            "reply": get_usage_limit_message("chat_message", user_id)
        }), 429

    chat_id = session.get("chat_id")
    new_chat_created = False

    if not chat_id:
        chat_id = create_chat(
            user_id,
            "New Chat"
        )

        session["chat_id"] = chat_id
        new_chat_created = True

    save_message(
        user_id,
        chat_id,
        "user",
        user_message
    )

    if new_chat_created:
        title_response = safe_openai_chat_completion(
            model="gpt-4.1-mini",
            messages=[
                {
                    "role": "system",
                    "content": "Create a short chat title. Maximum 5 words. No quotation marks."
                },
                {
                    "role": "user",
                    "content": user_message
                }
            ]
        )

        title = title_response.choices[0].message.content.strip()

        update_chat_title(
            chat_id,
            title
        )

    uploaded_file = get_latest_uploaded_file(user_id)
    file_context = ""

    if uploaded_file:
        if (
            uploaded_file.endswith(".txt")
            or uploaded_file.endswith(".pdf")
            or uploaded_file.endswith(".docx")
        ):
            file_context = extract_file_text(uploaded_file)[:4000]

    messages = [
        {
            "role": "system",
            "content": SYSTEM_PROMPT
        }
    ]

    project_context = build_project_context(user_id)

    if project_context:
        messages.append({
            "role": "system",
            "content": project_context
        })

    if file_context:
        messages.append({
            "role": "system",
            "content":
                "The user uploaded this file content:\n\n"
                + file_context
        })

    messages.extend(
        get_memory(
            user_id,
            chat_id
        )
    )

    image_extensions = [
        ".png",
        ".jpg",
        ".jpeg",
        ".webp"
    ]

    if uploaded_file and any(
        uploaded_file.lower().endswith(ext)
        for ext in image_extensions
    ):
        image_base64 = image_to_base64(uploaded_file)

        response = safe_openai_chat_completion(
            model="gpt-4.1-mini",
            messages=[
                {
                    "role": "system",
                    "content": SYSTEM_PROMPT
                },
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": user_message
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url":
                                "data:image/jpeg;base64,"
                                + image_base64
                            }
                        }
                    ]
                }
            ]
        )

    else:
        response = safe_openai_chat_completion(
            model="gpt-4.1-mini",
            messages=messages
        )

    reply = response.choices[0].message.content

    save_message(
        user_id,
        chat_id,
        "assistant",
        reply
    )
    log_usage(user_id, "chat_message")

    return jsonify({
        "reply": reply
    })


# -----------------------------
# START APP
# -----------------------------

init_db()

if __name__ == "__main__":
    app.run(
        host="0.0.0.0",
        port=int(os.environ.get("PORT", 5000))
    )
